import pytest
import os
from docx_mcp_server.server import docx_save, docx_read_content, docx_insert_paragraph, docx_find_paragraphs
from docx import Document
import json

# Add parent directory to path for helpers import
import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    is_success,
    is_error,
    create_session_with_file,
)
from tests.helpers.session_helpers import setup_active_session, teardown_active_session

class TestLoadExisting:
    def test_create_with_existing_file(self, tmp_path):
        # Setup: Create a real docx file
        file_path = tmp_path / "test_existing.docx"
        doc = Document()
        doc.add_paragraph("Existing Content")
        doc.save(str(file_path))

        # Test: Load it via docx_create
        session_response = create_session_with_file(str(file_path))
        assert session_id is not None

        # Verify content can be read (using our new tool later, or just save and check)
        # For now, let's verify we can add to it and save
        docx_insert_paragraph("New Content", position="end:document_body")

        output_path = tmp_path / "test_existing_modified.docx"
        docx_save(str(output_path))

        # Verify output
        doc2 = Document(str(output_path))
        texts = [p.text for p in doc2.paragraphs]
        assert "Existing Content" in texts
        assert "New Content" in texts

    def test_create_with_nonexistent_file(self, tmp_path):
        # Test that creating a session with a non-existent file creates a new document
        # intended to be saved to that path later.
        target_path = tmp_path / "new_doc.docx"
        session_response = create_session_with_file(str(target_path))
        assert session_id is not None

        # Verify it's an empty document (or default styles)
        content = docx_read_content()
        assert content == "[Empty Document]"

    def test_read_content(self, tmp_path):
        # Setup
        file_path = tmp_path / "test_read.docx"
        doc = Document()
        doc.add_paragraph("Line 1")
        doc.add_paragraph("Line 2")
        doc.save(str(file_path))

        session_response = create_session_with_file(str(file_path))
        # Test read
        content = docx_read_content()
        assert "Line 1" in content
        assert "Line 2" in content

    def test_find_paragraphs(self, tmp_path):
        # Setup
        file_path = tmp_path / "test_find.docx"
        doc = Document()
        doc.add_paragraph("First paragraph")
        doc.add_paragraph("Target paragraph here")
        doc.add_paragraph("Last paragraph")
        doc.save(str(file_path))

        session_response = create_session_with_file(str(file_path))
        # Test find
        result_json = docx_find_paragraphs("Target")
        matches = json.loads(result_json)

        assert len(matches) == 1
        assert matches[0]["text"] == "Target paragraph here"
        assert matches[0]["id"].startswith("para_")

        # Test find - no match
        result_json = docx_find_paragraphs("Nonexistent")
        matches = json.loads(result_json)
        assert len(matches) == 0

        # Test case insensitivity
        result_json = docx_find_paragraphs("target")
        matches = json.loads(result_json)
        assert len(matches) == 1
