"""Unit tests for special position IDs functionality.

Tests the Session class enhancements for special ID resolution:
- last_insert
- last_update
- cursor/current
- document_body
"""

import pytest
from docx import Document
from docx_mcp_server.core.session import Session


@pytest.fixture
def session():
    """Create a test session."""
    doc = Document()
    return Session(
        session_id="test_session",
        document=doc
    )


class TestResolveSpecialID:
    """Test Session.resolve_special_id() method."""

    def test_resolve_last_insert_success(self, session):
        """Test resolving last_insert when available."""
        session.last_insert_id = "para_abc123"
        result = session.resolve_special_id("last_insert")
        assert result == "para_abc123"

    def test_resolve_last_insert_case_insensitive(self, session):
        """Test last_insert is case-insensitive."""
        session.last_insert_id = "para_abc123"
        assert session.resolve_special_id("LAST_INSERT") == "para_abc123"
        assert session.resolve_special_id("Last_Insert") == "para_abc123"
        assert session.resolve_special_id("last_INSERT") == "para_abc123"

    def test_resolve_last_insert_not_available(self, session):
        """Test error when last_insert is not available."""
        with pytest.raises(ValueError) as exc_info:
            session.resolve_special_id("last_insert")
        assert "last_insert" in str(exc_info.value).lower()
        assert "not available" in str(exc_info.value).lower()

    def test_resolve_last_update_success(self, session):
        """Test resolving last_update when available."""
        session.last_update_id = "para_xyz789"
        result = session.resolve_special_id("last_update")
        assert result == "para_xyz789"

    def test_resolve_last_update_case_insensitive(self, session):
        """Test last_update is case-insensitive."""
        session.last_update_id = "para_xyz789"
        assert session.resolve_special_id("LAST_UPDATE") == "para_xyz789"
        assert session.resolve_special_id("Last_Update") == "para_xyz789"

    def test_resolve_last_update_not_available(self, session):
        """Test error when last_update is not available."""
        with pytest.raises(ValueError) as exc_info:
            session.resolve_special_id("last_update")
        assert "last_update" in str(exc_info.value).lower()
        assert "not available" in str(exc_info.value).lower()

    def test_resolve_cursor_success(self, session):
        """Test resolving cursor when available."""
        session.cursor.element_id = "para_cursor123"
        result = session.resolve_special_id("cursor")
        assert result == "para_cursor123"

    def test_resolve_current_alias(self, session):
        """Test 'current' as alias for 'cursor'."""
        session.cursor.element_id = "para_cursor123"
        result = session.resolve_special_id("current")
        assert result == "para_cursor123"

    def test_resolve_cursor_case_insensitive(self, session):
        """Test cursor is case-insensitive."""
        session.cursor.element_id = "para_cursor123"
        assert session.resolve_special_id("CURSOR") == "para_cursor123"
        assert session.resolve_special_id("Cursor") == "para_cursor123"
        assert session.resolve_special_id("CURRENT") == "para_cursor123"

    def test_resolve_cursor_not_available(self, session):
        """Test error when cursor is not available."""
        with pytest.raises(ValueError) as exc_info:
            session.resolve_special_id("cursor")
        assert "cursor" in str(exc_info.value).lower()
        assert "not available" in str(exc_info.value).lower()

    def test_resolve_document_body(self, session):
        """Test resolving document_body special ID."""
        result = session.resolve_special_id("document_body")
        assert result == "document_body"

    def test_resolve_document_body_case_insensitive(self, session):
        """Test document_body is case-insensitive."""
        assert session.resolve_special_id("DOCUMENT_BODY") == "document_body"
        assert session.resolve_special_id("Document_Body") == "document_body"

    def test_resolve_non_special_id_passthrough(self, session):
        """Test non-special IDs are returned as-is."""
        assert session.resolve_special_id("para_abc123") == "para_abc123"
        assert session.resolve_special_id("run_xyz789") == "run_xyz789"
        assert session.resolve_special_id("table_123") == "table_123"

    def test_resolve_with_whitespace(self, session):
        """Test IDs with whitespace are trimmed."""
        session.last_insert_id = "para_abc123"
        assert session.resolve_special_id("  last_insert  ") == "para_abc123"
        assert session.resolve_special_id("  para_xyz  ") == "para_xyz"


class TestUpdateContext:
    """Test Session.update_context() method."""

    def test_update_context_create_action(self, session):
        """Test update_context with action='create'."""
        session.update_context("para_new123", action="create")
        assert session.last_insert_id == "para_new123"
        assert session.last_accessed_id == "para_new123"

    def test_update_context_update_action(self, session):
        """Test update_context with action='update'."""
        session.update_context("para_modified456", action="update")
        assert session.last_update_id == "para_modified456"
        assert session.last_accessed_id == "para_modified456"

    def test_update_context_access_action(self, session):
        """Test update_context with action='access' (default)."""
        session.update_context("para_accessed789")
        assert session.last_accessed_id == "para_accessed789"
        # Should not update last_insert_id or last_update_id
        assert session.last_insert_id is None
        assert session.last_update_id is None

    def test_update_context_sequence(self, session):
        """Test sequence of context updates."""
        # Create
        session.update_context("para_1", action="create")
        assert session.last_insert_id == "para_1"
        assert session.last_update_id is None

        # Update
        session.update_context("para_2", action="update")
        assert session.last_insert_id == "para_1"  # Unchanged
        assert session.last_update_id == "para_2"

        # Another create
        session.update_context("para_3", action="create")
        assert session.last_insert_id == "para_3"  # Updated
        assert session.last_update_id == "para_2"  # Unchanged


class TestGetObjectIntegration:
    """Test Session.get_object() integration with special IDs."""

    def test_get_object_with_special_id(self, session):
        """Test get_object resolves special IDs."""
        # Register an object
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")
        session.last_insert_id = para_id

        # Get object using special ID
        result = session.get_object("last_insert")
        assert result is para

    def test_get_object_special_id_not_available(self, session):
        """Test get_object raises ValueError for unavailable special ID."""
        with pytest.raises(ValueError) as exc_info:
            session.get_object("last_insert")
        assert "last_insert" in str(exc_info.value).lower()

    def test_get_object_with_regular_id(self, session):
        """Test get_object still works with regular IDs."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")

        result = session.get_object(para_id)
        assert result is para

    def test_get_object_case_insensitive_special_id(self, session):
        """Test get_object with case-insensitive special ID."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")
        session.last_insert_id = para_id

        result = session.get_object("LAST_INSERT")
        assert result is para


class TestEdgeCases:
    """Test edge cases and boundary conditions."""

    def test_element_deleted_after_last_insert(self, session):
        """Test using last_insert after element is deleted."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")
        session.last_insert_id = para_id

        # Delete from registry
        del session.object_registry[para_id]

        # Special ID resolves but get_object returns None
        resolved = session.resolve_special_id("last_insert")
        assert resolved == para_id
        result = session.get_object("last_insert")
        assert result is None

    def test_session_just_created(self, session):
        """Test special IDs on newly created session."""
        # All special IDs should be unavailable
        with pytest.raises(ValueError):
            session.resolve_special_id("last_insert")
        with pytest.raises(ValueError):
            session.resolve_special_id("last_update")
        with pytest.raises(ValueError):
            session.resolve_special_id("cursor")

    def test_multiple_operations_tracking(self, session):
        """Test tracking across multiple operations."""
        # Create first element
        para1 = session.document.add_paragraph("First")
        para1_id = session.register_object(para1, "para")
        session.update_context(para1_id, action="create")

        # Create second element
        para2 = session.document.add_paragraph("Second")
        para2_id = session.register_object(para2, "para")
        session.update_context(para2_id, action="create")

        # last_insert should point to most recent
        assert session.last_insert_id == para2_id
        assert session.get_object("last_insert") is para2

        # Update first element
        session.update_context(para1_id, action="update")
        assert session.last_update_id == para1_id
        assert session.get_object("last_update") is para1

    def test_empty_string_id(self, session):
        """Test handling of empty string ID."""
        result = session.get_object("")
        assert result is None

    def test_whitespace_only_id(self, session):
        """Test handling of whitespace-only ID."""
        result = session.get_object("   ")
        assert result is None

