import pytest
from docx import Document
from docx.text.paragraph import Paragraph
from docx_mcp_server.utils.copy_engine import CopyEngine

def test_get_elements_between_basic():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    p2 = doc.add_paragraph("P2")
    p3 = doc.add_paragraph("P3")
    p4 = doc.add_paragraph("P4")

    engine = CopyEngine()
    elements = engine.get_elements_between(p2, p3)

    assert len(elements) == 2
    assert elements[0].text == "P2"
    assert elements[1].text == "P3"

def test_get_elements_between_validation():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    p2 = doc.add_paragraph("P2")

    doc2 = Document()
    p_other = doc2.add_paragraph("Other")

    engine = CopyEngine()

    # Different parents
    with pytest.raises(ValueError, match="same parent"):
        engine.get_elements_between(p1, p_other)

    # Reverse order
    with pytest.raises(ValueError, match="Start element must appear before"):
        engine.get_elements_between(p2, p1)

def test_copy_range_append():
    doc = Document()
    p1 = doc.add_paragraph("Source 1")
    p2 = doc.add_paragraph("Source 2")

    engine = CopyEngine()

    # Copy range P1-P2 and append to same doc
    new_objects = engine.copy_range(p1, p2, doc)

    assert len(new_objects) == 2
    assert new_objects[0].text == "Source 1"
    assert new_objects[1].text == "Source 2"

    # Verify they are at the end
    assert doc.paragraphs[-2].text == "Source 1"
    assert doc.paragraphs[-1].text == "Source 2"
    assert len(doc.paragraphs) == 4

def test_copy_range_insert_middle():
    doc = Document()
    p1 = doc.add_paragraph("P1") # Source Start
    p2 = doc.add_paragraph("P2") # Source End
    target = doc.add_paragraph("Target") # Insertion point
    footer = doc.add_paragraph("Footer")

    engine = CopyEngine()

    # Insert P1-P2 range AFTER 'Target'
    new_objects = engine.copy_range(p1, p2, doc, ref_element_xml=target._element)

    # Expected: P1, P2, Target, P1(Copy), P2(Copy), Footer
    texts = [p.text for p in doc.paragraphs]
    assert texts == ["P1", "P2", "Target", "P1", "P2", "Footer"]
