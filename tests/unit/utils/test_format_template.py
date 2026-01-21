import pytest
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_mcp_server.utils.format_template import TemplateManager, FormatTemplate

def test_extract_apply_run_template():
    doc = Document()
    p = doc.add_paragraph()
    run_src = p.add_run("Source")
    run_src.font.bold = True
    run_src.font.size = Pt(12)
    run_src.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    manager = TemplateManager()
    template = manager.extract_template(run_src)

    assert template.element_type == "run"
    assert template.font_properties['bold'] == True
    assert template.font_properties['color_rgb'] == "FF0000"

    # Apply to new run
    run_tgt = p.add_run("Target")
    manager.apply_template(run_tgt, template)

    assert run_tgt.font.bold == True
    assert run_tgt.font.size == Pt(12)
    assert str(run_tgt.font.color.rgb) == "FF0000"

def test_extract_apply_paragraph_template():
    doc = Document()
    p_src = doc.add_paragraph("Source Para")
    p_src.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_src.paragraph_format.space_after = Pt(10)

    manager = TemplateManager()
    template = manager.extract_template(p_src)

    assert template.element_type == "paragraph"
    assert template.properties['alignment'] == WD_ALIGN_PARAGRAPH.CENTER

    p_tgt = doc.add_paragraph("Target Para")
    manager.apply_template(p_tgt, template)

    assert p_tgt.alignment == WD_ALIGN_PARAGRAPH.CENTER
    assert p_tgt.paragraph_format.space_after == Pt(10)

def test_json_serialization():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("Text")
    run.font.bold = True

    manager = TemplateManager()
    template = manager.extract_template(run)

    json_str = manager.to_json(template)
    assert "bold" in json_str
    assert "true" in json_str.lower()

    template_restored = manager.from_json(json_str)
    assert template_restored.font_properties['bold'] == True
