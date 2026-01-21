import pytest
from docx import Document
from docx.text.paragraph import Paragraph
from docx_mcp_server.utils.copy_engine import CopyEngine

def test_copy_paragraph_content():
    doc = Document()
    p1 = doc.add_paragraph("Original")
    run = p1.runs[0]
    run.bold = True

    engine = CopyEngine()

    # 1. Copy Element
    new_xml = engine.copy_element(p1)

    # 2. Insert back into document
    # For testing, we insert into the same document
    new_para = engine.insert_element_after(doc, new_xml)

    assert isinstance(new_para, Paragraph)
    assert new_para.text == "Original"
    assert new_para.runs[0].bold == True

    # Verify independence
    p1.text = "Modified"
    assert new_para.text == "Original"

def test_insert_position():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    p2 = doc.add_paragraph("P2")
    p3 = doc.add_paragraph("P3")

    engine = CopyEngine()
    new_xml = engine.copy_element(p1) # Copy P1

    # Insert after P2
    new_para = engine.insert_element_after(doc, new_xml, ref_element_xml=p2._element)

    # Expected order: P1, P2, P1(Copy), P3
    assert doc.paragraphs[0].text == "P1"
    assert doc.paragraphs[1].text == "P2"
    assert doc.paragraphs[2].text == "P1"
    assert doc.paragraphs[3].text == "P3"

def test_unsupported_type():
    engine = CopyEngine()
    with pytest.raises(ValueError):
        engine.copy_element("Not an element")
