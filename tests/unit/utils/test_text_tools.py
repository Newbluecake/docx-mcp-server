import pytest
from docx import Document
from docx_mcp_server.utils.text_tools import TextTools

def test_batch_replace_runs_simple():
    doc = Document()
    p = doc.add_paragraph("Hello World")
    p.add_run(" This is a test.")

    tools = TextTools()
    replacements = {"World": "Universe", "test": "demo"}

    count = tools.batch_replace_text([p], replacements)

    assert count == 2
    assert p.text == "Hello Universe This is a demo."

def test_batch_replace_multiple_paragraphs():
    doc = Document()
    p1 = doc.add_paragraph("Para 1: {VAR}")
    p2 = doc.add_paragraph("Para 2: {VAR}")

    tools = TextTools()
    replacements = {"{VAR}": "Value"}

    count = tools.batch_replace_text([p1, p2], replacements)

    assert count == 2
    assert p1.text == "Para 1: Value"
    assert p2.text == "Para 2: Value"

def test_batch_replace_skip_split_run():
    """
    Verifies that if text is split across runs, it is NOT replaced by this tool
    (as per strict run-level MVP constraint).
    """
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("He")
    p.add_run("llo") # "Hello" is split

    tools = TextTools()
    replacements = {"Hello": "Hi"}

    count = tools.batch_replace_text([p], replacements)

    assert count == 0
    assert p.text == "Hello" # Unchanged
