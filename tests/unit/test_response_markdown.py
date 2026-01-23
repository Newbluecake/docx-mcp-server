"""Unit tests for Markdown response formatting."""

import pytest
from docx import Document
from docx_mcp_server.core.response import (
    create_markdown_response,
    create_error_response
)
from docx_mcp_server.core.session import Session


@pytest.fixture
def session():
    """Create a test session with a document."""
    doc = Document()
    session = Session(
        session_id="test_session",
        document=doc
    )
    return session


class TestMarkdownResponse:
    """Tests for create_markdown_response function."""

    def test_create_markdown_response_success(self, session):
        """Test basic success response."""
        para = session.document.add_paragraph("Test paragraph")
        para_id = session.register_object(para, "para")

        result = create_markdown_response(
            session=session,
            message="Paragraph created successfully",
            element_id=para_id,
            operation="Insert Paragraph"
        )

        assert "# 操作结果: Insert Paragraph" in result
        assert "**Status**: ✅ Success" in result
        assert f"**Element ID**: {para_id}" in result
        assert "**Operation**: Insert Paragraph" in result
        assert "📄 Document Context" in result

    def test_create_markdown_response_no_context(self, session):
        """Test response without context."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")

        result = create_markdown_response(
            session=session,
            message="Paragraph created",
            element_id=para_id,
            operation="Insert Paragraph",
            show_context=False
        )

        assert "**Status**: ✅ Success" in result
        assert "📄 Document Context" not in result

    def test_create_markdown_response_with_diff(self, session):
        """Test response with diff."""
        para = session.document.add_paragraph("Old text")
        para_id = session.register_object(para, "para")

        result = create_markdown_response(
            session=session,
            message="Paragraph updated",
            element_id=para_id,
            operation="Update Paragraph Text",
            show_diff=True,
            old_content="Old text",
            new_content="New text"
        )

        assert "🔄 Changes" in result
        assert "- │ Old text" in result
        assert "+ │ New text" in result

    def test_create_markdown_response_with_extra_metadata(self, session):
        """Test response with extra metadata fields."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")

        result = create_markdown_response(
            session=session,
            message="Paragraph created",
            element_id=para_id,
            operation="Insert Paragraph",
            show_context=False,
            position="end:document_body",
            style="Normal"
        )

        assert "**Position**: end:document_body" in result
        assert "**Style**: Normal" in result

    def test_create_markdown_response_no_session(self):
        """Test response without session (error case)."""
        result = create_markdown_response(
            session=None,
            message="Operation completed",
            element_id="para_123",
            operation="Test Operation",
            show_context=False
        )

        assert "**Status**: ✅ Success" in result
        assert "**Element ID**: para_123" in result
        assert "📄 Document Context" not in result

    def test_create_markdown_response_error_status(self):
        """Test error status response."""
        result = create_markdown_response(
            session=None,
            message="Operation failed",
            status="error",
            show_context=False,
            error_type="ValidationError"
        )

        assert "**Status**: ❌ Error" in result
        assert "**Error Type**: ValidationError" in result


class TestErrorResponse:
    """Tests for create_error_response function."""

    def test_create_error_response_basic(self):
        """Test basic error response."""
        result = create_error_response("Session not found")

        assert "# 操作结果: Error" in result
        assert "**Status**: ❌ Error" in result
        assert "**Message**: Session not found" in result

    def test_create_error_response_with_type(self):
        """Test error response with error type."""
        result = create_error_response(
            "Session not found",
            error_type="SessionNotFound"
        )

        assert "**Status**: ❌ Error" in result
        assert "**Error Type**: SessionNotFound" in result
        assert "**Message**: Session not found" in result


class TestIntegration:
    """Integration tests for response formatting."""

    def test_full_workflow_create_paragraph(self, session):
        """Test full workflow: create paragraph with context."""
        # Add some context paragraphs
        for i in range(5):
            para = session.document.add_paragraph(f"Context paragraph {i}")
            session.register_object(para, "para")

        # Add target paragraph
        target_para = session.document.add_paragraph("Target paragraph")
        target_id = session.register_object(target_para, "para")

        # Create response
        result = create_markdown_response(
            session=session,
            message="Paragraph created successfully",
            element_id=target_id,
            operation="Insert Paragraph",
            position="end:document_body"
        )

        # Verify structure
        assert "# 操作结果: Insert Paragraph" in result
        assert "**Status**: ✅ Success" in result
        assert f"**Element ID**: {target_id}" in result
        assert "**Position**: end:document_body" in result
        assert "📄 Document Context" in result
        assert "Target paragraph" in result
        assert "⭐ CURRENT" in result

    def test_full_workflow_update_paragraph(self, session):
        """Test full workflow: update paragraph with diff."""
        # Create paragraph
        para = session.document.add_paragraph("Original text")
        para_id = session.register_object(para, "para")

        # Simulate update
        old_text = para.text
        para.text = "Updated text"

        # Create response with diff
        result = create_markdown_response(
            session=session,
            message="Paragraph text updated",
            element_id=para_id,
            operation="Update Paragraph Text",
            show_diff=True,
            old_content=old_text,
            new_content=para.text
        )

        # Verify structure
        assert "**Status**: ✅ Success" in result
        assert "🔄 Changes" in result
        assert "- │ Original text" in result
        assert "+ │ Updated text" in result
        assert "📄 Document Context" in result
