import pytest
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx_mcp_server.core.template_parser import TemplateParser


def test_template_parser_init():
    """Test TemplateParser initialization."""
    parser = TemplateParser()
    assert parser is not None


def test_extract_structure_empty_document():
    """Test extracting structure from empty document."""
    doc = Document()
    parser = TemplateParser()
    result = parser.extract_structure(doc)

    assert "metadata" in result
    assert "document_structure" in result
    assert isinstance(result["document_structure"], list)
    assert len(result["document_structure"]) == 0


# T-002: Header detection tests
def test_detect_header_row_bold():
    """Test header detection with bold first row."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)

    # Make first row bold
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"Header {i + 1}"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    parser = TemplateParser()
    header_row = parser.detect_header_row(table)
    assert header_row == 0


def test_detect_header_row_background():
    """Test header detection with background color."""
    from lxml import etree
    doc = Document()
    table = doc.add_table(rows=2, cols=3)

    # Add background color to first row using lxml
    for cell in table.rows[0].cells:
        cell.text = "Header"
        tc_pr = cell._element.get_or_add_tcPr()
        # Create shading element manually
        shd = etree.SubElement(tc_pr, '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd')
        shd.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill', 'CCCCCC')

    parser = TemplateParser()
    header_row = parser.detect_header_row(table)
    assert header_row == 0


def test_detect_header_row_fail():
    """Test header detection failure when no bold or background."""
    doc = Document()
    table = doc.add_table(rows=2, cols=3)

    # First row without bold or background
    for cell in table.rows[0].cells:
        cell.text = "Not Header"

    parser = TemplateParser()
    with pytest.raises(ValueError, match="无法检测表头行"):
        parser.detect_header_row(table)


# T-004: Heading extraction tests
def test_extract_heading_structure():
    """Test extracting heading structure."""
    doc = Document()
    heading = doc.add_heading("Test Heading", level=1)

    parser = TemplateParser()
    result = parser.extract_heading_structure(heading)

    assert result["type"] == "heading"
    assert result["level"] == 1
    assert result["text"] == "Test Heading"
    assert "style" in result


# T-005: Paragraph extraction tests
def test_extract_paragraph_structure():
    """Test extracting paragraph structure."""
    doc = Document()
    para = doc.add_paragraph("Test paragraph")
    run = para.runs[0]
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.bold = True

    parser = TemplateParser()
    result = parser.extract_paragraph_structure(para)

    assert result["type"] == "paragraph"
    assert result["text"] == "Test paragraph"
    assert result["style"]["font"] == "Arial"
    assert result["style"]["size"] == 12
    assert result["style"]["bold"] is True


# T-006: Image extraction tests
def test_extract_image_structure():
    """Test extracting image structure from paragraph."""
    # This test will be skipped for now as it requires actual image file
    pytest.skip("Image extraction requires actual image file")
