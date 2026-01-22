from docx import Document

from docx_mcp_server.core.replacer import replace_text_in_paragraph


def test_replace_across_runs_preserves_run_formatting():
    doc = Document()
    p = doc.add_paragraph()

    run_plain = p.add_run("Hi ")
    run_bold = p.add_run("World")
    run_bold.bold = True
    run_italic = p.add_run("!")
    run_italic.italic = True

    replaced = replace_text_in_paragraph(p, "World!", "Alice?")

    assert replaced is True
    assert p.text == "Hi Alice?"

    # Bold formatting from the start run should stay on the new text
    assert p.runs[1].text == "Alice?"
    assert p.runs[1].bold is True

    # The trailing run should be cleared, keeping its formatting untouched
    assert p.runs[2].text == ""
    assert p.runs[2].italic is True


def test_replace_inside_single_run_keeps_formatting():
    doc = Document()
    p = doc.add_paragraph()
    r = p.add_run("Token")
    r.bold = True

    replaced = replace_text_in_paragraph(p, "Token", "Value")

    assert replaced is True
    assert p.text == "Value"
    assert p.runs[0].text == "Value"
    assert p.runs[0].bold is True
