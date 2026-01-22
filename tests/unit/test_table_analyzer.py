"""
Unit tests for TableStructureAnalyzer.
"""

import pytest
from docx import Document
from docx_mcp_server.core.table_analyzer import TableStructureAnalyzer


def test_detect_regular_table():
    """Test detection of regular table structure."""
    doc = Document()
    table = doc.add_table(rows=3, cols=3)

    result = TableStructureAnalyzer.detect_irregular_structure(table)

    assert result["is_irregular"] == False
    assert result["has_merged_cells"] == False
    assert result["has_nested_tables"] == False
    assert len(result["irregular_regions"]) == 0


def test_detect_merged_cells():
    """Test detection of merged cells."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    # Merge first row cells
    cell_a = table.rows[0].cells[0]
    cell_b = table.rows[0].cells[1]
    cell_a.merge(cell_b)

    result = TableStructureAnalyzer.detect_irregular_structure(table)

    assert result["is_irregular"] == True
    assert result["has_merged_cells"] == True


def test_generate_ascii_visualization():
    """Test ASCII visualization generation."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Header 1"
    table.rows[0].cells[1].text = "Header 2"
    table.rows[1].cells[0].text = "Data 1"
    table.rows[1].cells[1].text = "Data 2"

    viz = TableStructureAnalyzer.generate_ascii_visualization(table)

    assert "Table: 2 rows x 2 cols" in viz
    assert "Header 1" in viz
    assert "Data 2" in viz
    assert "+" in viz  # Border characters
    assert "|" in viz  # Cell separators


def test_get_fillable_cells_regular():
    """Test getting fillable cells from regular table."""
    doc = Document()
    table = doc.add_table(rows=3, cols=2)

    fillable = TableStructureAnalyzer.get_fillable_cells(table)

    assert len(fillable) == 6  # 3 rows x 2 cols
    assert (0, 0) in fillable
    assert (2, 1) in fillable


def test_get_fillable_cells_with_merged():
    """Test getting fillable cells from table with merged cells."""
    doc = Document()
    table = doc.add_table(rows=2, cols=2)

    # Merge first row
    cell_a = table.rows[0].cells[0]
    cell_b = table.rows[0].cells[1]
    cell_a.merge(cell_b)

    structure_info = TableStructureAnalyzer.detect_irregular_structure(table)
    fillable = TableStructureAnalyzer.get_fillable_cells(table, structure_info)

    # Should exclude merged cells
    assert len(fillable) < 4
