"""
Unit tests for element_id enhancement feature.

Tests the following functionality:
- Session._get_element_id() caching mechanism
- TemplateParser methods with session parameter
- Tools returning element_id in responses
- Element ID reusability across tools
"""

import json
import pytest
from docx import Document
from docx_mcp_server.core.session import SessionManager
from docx_mcp_server.core.template_parser import TemplateParser
from docx_mcp_server.tools.content_tools import (
    docx_extract_template_structure,
    docx_read_content
)
from docx_mcp_server.tools.composite_tools import docx_get_structure_summary

# Add parent directory to path for helpers import
import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    is_success,
    is_error,
    extract_error_message
)
from tests.helpers.session_helpers import setup_active_session, teardown_active_session


@pytest.fixture
def session_manager():
    """Create a fresh SessionManager for each test."""
    return SessionManager()


@pytest.fixture
def session_id(session_manager):
    """Create a test session."""
    return session_manager.create_session()


class TestSessionElementIdCache:
    """Test Session._get_element_id() caching mechanism."""

    def test_cache_hit_returns_same_id(self, session_manager, session_id):
        """Test that cache hit returns the same element_id."""
        session = session_manager.get_session(session_id)
        para = session.document.add_paragraph("Test paragraph")

        # First call - cache miss, should register
        id1 = session._get_element_id(para, auto_register=True)
        assert id1 is not None
        assert id1.startswith("para_")

        # Second call - cache hit, should return same ID
        id2 = session._get_element_id(para, auto_register=True)
        assert id1 == id2

    def test_cache_miss_with_auto_register(self, session_manager, session_id):
        """Test that cache miss with auto_register=True registers element."""
        session = session_manager.get_session(session_id)
        para = session.document.add_paragraph("Test paragraph")

        # Cache miss, auto_register=True
        element_id = session._get_element_id(para, auto_register=True)
        assert element_id is not None
        assert element_id.startswith("para_")

        # Verify element is in registry
        assert element_id in session.object_registry
        assert session.object_registry[element_id] == para

    def test_cache_miss_without_auto_register(self, session_manager, session_id):
        """Test that cache miss with auto_register=False returns None."""
        session = session_manager.get_session(session_id)
        para = session.document.add_paragraph("Test paragraph")

        # Cache miss, auto_register=False
        element_id = session._get_element_id(para, auto_register=False)
        assert element_id is None

    def test_different_element_types(self, session_manager, session_id):
        """Test that different element types get correct prefixes."""
        session = session_manager.get_session(session_id)

        # Paragraph
        para = session.document.add_paragraph("Test")
        para_id = session._get_element_id(para, auto_register=True)
        assert para_id.startswith("para_")

        # Table
        table = session.document.add_table(rows=2, cols=2)
        table_id = session._get_element_id(table, auto_register=True)
        assert table_id.startswith("table_")

        # Cell
        cell = table.rows[0].cells[0]
        cell_id = session._get_element_id(cell, auto_register=True)
        assert cell_id.startswith("cell_")


class TestTemplateParserWithSession:
    """Test TemplateParser methods with session parameter."""

    def test_extract_heading_structure_with_session(self, session_manager, session_id):
        """Test extract_heading_structure() with session parameter."""
        session = session_manager.get_session(session_id)
        para = session.document.add_heading("Test Heading", level=1)

        parser = TemplateParser()
        result = parser.extract_heading_structure(para, session=session)

        assert "element_id" in result
        assert result["element_id"].startswith("para_")
        assert result["type"] == "heading"
        assert result["level"] == 1
        assert result["text"] == "Test Heading"

    def test_extract_heading_structure_without_session(self):
        """Test extract_heading_structure() without session (backward compatibility)."""
        doc = Document()
        para = doc.add_heading("Test Heading", level=1)

        parser = TemplateParser()
        result = parser.extract_heading_structure(para, session=None)

        assert "element_id" not in result  # Backward compatibility
        assert result["type"] == "heading"
        assert result["level"] == 1

    def test_extract_paragraph_structure_with_session(self, session_manager, session_id):
        """Test extract_paragraph_structure() with session parameter."""
        session = session_manager.get_session(session_id)
        para = session.document.add_paragraph("Test paragraph")

        parser = TemplateParser()
        result = parser.extract_paragraph_structure(para, session=session)

        assert "element_id" in result
        assert result["element_id"].startswith("para_")
        assert result["type"] == "paragraph"
        assert result["text"] == "Test paragraph"

    def test_extract_paragraph_structure_without_session(self):
        """Test extract_paragraph_structure() without session (backward compatibility)."""
        doc = Document()
        para = doc.add_paragraph("Test paragraph")

        parser = TemplateParser()
        result = parser.extract_paragraph_structure(para, session=None)

        assert "element_id" not in result  # Backward compatibility
        assert result["type"] == "paragraph"

    def test_extract_table_structure_with_session(self, session_manager, session_id):
        """Test extract_table_structure() with session parameter."""
        session = session_manager.get_session(session_id)
        table = session.document.add_table(rows=2, cols=2)

        # Set header row (bold)
        for cell in table.rows[0].cells:
            cell.paragraphs[0].add_run("Header").bold = True

        parser = TemplateParser()
        result = parser.extract_table_structure(table, session=session)

        assert "element_id" in result
        assert result["element_id"].startswith("table_")
        assert result["type"] == "table"
        assert result["rows"] == 2
        assert result["cols"] == 2

    def test_extract_table_structure_without_session(self):
        """Test extract_table_structure() without session (backward compatibility)."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)

        # Set header row (bold)
        for cell in table.rows[0].cells:
            cell.paragraphs[0].add_run("Header").bold = True

        parser = TemplateParser()
        result = parser.extract_table_structure(table, session=None)

        assert "element_id" not in result  # Backward compatibility
        assert result["type"] == "table"


class TestToolsWithElementId:
    """Test tools returning element_id in responses."""

    def test_extract_template_structure_returns_element_ids(self):
        """Test docx_extract_template_structure() returns element_id."""
        setup_active_session()
        try:
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            from tests.helpers import extract_template_structure
            session = session_manager.get_session(global_state.active_session_id)

            # Add content
            session.document.add_heading("Test Heading", level=1)
            session.document.add_paragraph("Test paragraph")
            table = session.document.add_table(rows=2, cols=2)
            for cell in table.rows[0].cells:
                cell.paragraphs[0].add_run("Header").bold = True

            # Extract structure
            result_markdown = docx_extract_template_structure()
            result = extract_template_structure(result_markdown)

            # Verify structure
            assert "document_structure" in result
            doc_structure = result["document_structure"]
            assert len(doc_structure) == 3

            # Verify heading has element_id (now 'id' in Markdown format)
            heading = doc_structure[0]
            assert heading["type"] == "heading"
            assert "id" in heading
            assert heading["id"].startswith("para_")

            # Verify paragraph has element_id
            para = doc_structure[1]
            assert para["type"] == "paragraph"
            assert "id" in para
            assert para["id"].startswith("para_")

            # Verify table has element_id
            table_elem = doc_structure[2]
            assert table_elem["type"] == "table"
            assert "id" in table_elem
            assert table_elem["id"].startswith("table_")
        finally:
            teardown_active_session()

    def test_get_structure_summary_returns_element_ids(self):
        """Test docx_get_structure_summary() returns element_id."""
        setup_active_session()
        try:
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            from tests.helpers import extract_json_from_markdown
            session = session_manager.get_session(global_state.active_session_id)

            # Add content
            session.document.add_heading("Heading 1", level=1)
            session.document.add_heading("Heading 2", level=2)
            session.document.add_paragraph("Paragraph 1")
            table = session.document.add_table(rows=2, cols=2)
            for cell in table.rows[0].cells:
                cell.paragraphs[0].add_run("Header").bold = True

            # Get structure summary
            result_markdown = docx_get_structure_summary(
                max_headings=10,
                max_tables=5,
                max_paragraphs=5
            )
            result = extract_json_from_markdown(result_markdown)

            # Verify headings have element_id
            assert len(result["headings"]) == 2
            for heading in result["headings"]:
                assert "element_id" in heading
                assert heading["element_id"].startswith("para_")

            # Verify tables have element_id
            assert len(result["tables"]) == 1
            for table in result["tables"]:
                assert "element_id" in table
                assert table["element_id"].startswith("table_")

            # Verify paragraphs have element_id
            assert len(result["paragraphs"]) == 1
            for para in result["paragraphs"]:
                assert "element_id" in para
                assert para["element_id"].startswith("para_")
        finally:
            teardown_active_session()

    def test_read_content_default_includes_ids(self):
        """Test docx_read_content() defaults to include_ids=True."""
        setup_active_session()
        try:
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            session = session_manager.get_session(global_state.active_session_id)
            session.document.add_paragraph("Test paragraph 1")
            session.document.add_paragraph("Test paragraph 2")

            # Call without include_ids parameter (should default to True)
            result = docx_read_content(return_json=True)

            # Verify element_ids are included by default
            assert is_success(result)
            data = json.loads(result)
            assert "data" in data
            assert len(data["data"]) == 2
            for para in data["data"]:
                assert "id" in para
                assert para["id"].startswith("para_")
        finally:
            teardown_active_session()

    def test_read_content_explicit_exclude_ids(self):
        """Test docx_read_content() with include_ids=False."""
        setup_active_session()
        try:
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            session = session_manager.get_session(global_state.active_session_id)
            session.document.add_paragraph("Test paragraph 1")
            session.document.add_paragraph("Test paragraph 2")

            # Call with include_ids=False
            result = docx_read_content(return_json=True, include_ids=False)

            # Verify element_ids are NOT included
            assert is_success(result)
            data = json.loads(result)
            assert "data" in data
            assert len(data["data"]) == 2
            for para in data["data"]:
                assert "id" not in para
        finally:
            teardown_active_session()


class TestElementIdReusability:
    """Test that returned element_ids can be used with other tools."""

    def test_element_id_usable_across_tools(self):
        """Test that element_id from extract_structure can be used to modify elements."""
        setup_active_session()
        try:
            from docx_mcp_server.tools.paragraph_tools import docx_update_paragraph_text
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            from tests.helpers import extract_template_structure

            session = session_manager.get_session(global_state.active_session_id)

            # Add content
            session.document.add_heading("Test Heading", level=1)
            session.document.add_paragraph("Original text")

            # Extract structure to get element_ids
            result_markdown = docx_extract_template_structure()
            result = extract_template_structure(result_markdown)

            # Get paragraph element_id (now 'id' in Markdown format)
            para = result["document_structure"][1]
            para_id = para["id"]

            # Use element_id to modify paragraph
            update_result = docx_update_paragraph_text(para_id, "Modified text")

            # Verify modification succeeded (check for success in markdown response)
            assert is_success(update_result)

            # Verify text was actually modified
            modified_para = session.get_object(para_id)
            assert modified_para.text == "Modified text"
        finally:
            teardown_active_session()

    def test_table_element_id_usable(self):
        """Test that table element_id can be used with table tools."""
        setup_active_session()
        try:
            from docx_mcp_server.tools.table_tools import docx_get_table_structure
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            from tests.helpers import extract_template_structure, extract_all_metadata

            session = session_manager.get_session(global_state.active_session_id)

            # Add table
            table = session.document.add_table(rows=2, cols=2)
            for cell in table.rows[0].cells:
                cell.paragraphs[0].add_run("Header").bold = True

            # Extract structure to get table_id
            result_markdown = docx_extract_template_structure()
            result = extract_template_structure(result_markdown)

            # Get table element_id (now 'id' in Markdown format)
            table_elem = result["document_structure"][0]
            table_id = table_elem["id"]

            # Use table_id with docx_get_table_structure
            table_structure_markdown = docx_get_table_structure(table_id)

            # Extract metadata from Markdown response
            assert is_success(table_structure_markdown)
            metadata = extract_all_metadata(table_structure_markdown)

            # Verify it worked
            assert metadata["rows"] == 2
            assert metadata["cols"] == 2
        finally:
            teardown_active_session()


class TestBackwardCompatibility:
    """Test backward compatibility - existing code should still work."""

    def test_existing_tests_still_pass(self):
        """Test that existing functionality is not broken."""
        setup_active_session()
        try:
            from docx_mcp_server.server import session_manager
            from docx_mcp_server.core.global_state import global_state
            from tests.helpers import extract_template_structure
            session = session_manager.get_session(global_state.active_session_id)

            # Old way of using tools (without relying on element_id)
            session.document.add_paragraph("Test paragraph")

            # Extract structure (old code doesn't check for element_id)
            result_markdown = docx_extract_template_structure()
            result = extract_template_structure(result_markdown)

            # Old code just checks structure exists
            assert "document_structure" in result
            assert len(result["document_structure"]) > 0

            # This should still work even though we added element_id
        finally:
            teardown_active_session()

    def test_parser_without_session_still_works(self):
        """Test that TemplateParser can still be used without session."""
        doc = Document()
        doc.add_heading("Test", level=1)
        doc.add_paragraph("Test paragraph")

        parser = TemplateParser()

        # Old way - no session parameter
        structure = parser.extract_structure(doc)

        # Should still work, just without element_ids
        assert "document_structure" in structure
        assert len(structure["document_structure"]) == 2


if __name__ == "__main__":
    pytest.main([__file__, "-v"])
