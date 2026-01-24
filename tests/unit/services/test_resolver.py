import pytest
from unittest.mock import MagicMock
from docx import Document
from docx_mcp_server.services.navigation import PositionResolver, ContextBuilder

class MockSession:
    def __init__(self, document):
        self.document = document
        self.objects = {} # id -> obj
        self.id_map = {} # obj -> id

    def get_object(self, element_id):
        return self.objects.get(element_id)

    def _get_element_id(self, obj, auto_register=False):
        return self.id_map.get(obj, "?")

    def resolve_special_id(self, element_id):
        """Resolve special IDs (last_insert, last_update, cursor) to actual IDs."""
        # For testing, just return the ID as-is (no special IDs in these tests)
        return element_id

def test_resolve_append_default():
    doc = Document()
    session = MockSession(doc)
    resolver = PositionResolver(session)

    # Test None -> Default Append
    parent, ref, mode = resolver.resolve(None)
    assert parent == doc
    assert ref is None
    assert mode == "append"

def test_resolve_after_paragraph():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    # docx.Document does not have _parent set usually, but elements do have _parent pointing to it
    # We need to ensure p1._parent works for the resolver logic
    # In python-docx, p1._parent is the Document object (or Body)

    session = MockSession(doc)
    session.objects["para_1"] = p1
    session.id_map[p1] = "para_1"

    resolver = PositionResolver(session)

    # Test "after:para_1"
    parent, ref, mode = resolver.resolve("after:para_1")

    # Parent should be the document body (or document depending on mock)
    # python-docx: p1._parent is usually the Body object, but ElementNavigator.get_docx_parent attempts to resolve
    assert parent is not None
    assert ref == p1
    assert mode == "after"

def test_resolve_inside_container():
    doc = Document()
    table = doc.add_table(1, 1)
    cell = table.cell(0, 0)

    session = MockSession(doc)
    session.objects["cell_1"] = cell

    resolver = PositionResolver(session)

    # Test "inside:cell_1"
    parent, ref, mode = resolver.resolve("inside:cell_1")

    assert parent == cell
    assert ref is None
    assert mode == "append"

def test_context_builder_structure():
    doc = Document()
    p1 = doc.add_paragraph("P1")

    session = MockSession(doc)
    session.id_map[p1] = "para_1"

    builder = ContextBuilder(session)

    data = builder.build_response_data(p1, "para_1")

    assert data["element_id"] == "para_1"
    assert "cursor" in data
    assert data["cursor"]["element_id"] == "para_1"
    assert "visual" in data["cursor"]
    assert "path" in data["cursor"]
    # Visual should contain "P1"
    assert "P1" in data["cursor"]["visual"]
