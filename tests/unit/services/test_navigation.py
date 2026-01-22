import pytest
from unittest.mock import MagicMock
from docx import Document
from docx_mcp_server.services.navigation import ContextVisualizer

class MockSession:
    def __init__(self, document):
        self.document = document
        self.text_id_map = {} # text -> id

    def _get_element_id(self, obj, auto_register=False):
        # Map by text content since python-docx objects are transient wrappers
        if hasattr(obj, 'text') and obj.text in self.text_id_map:
            return self.text_id_map[obj.text]
        return "id_unknown"

def test_visualizer_basic_structure():
    doc = Document()
    h1 = doc.add_heading("Introduction", level=1)
    p1 = doc.add_paragraph("This is the first paragraph.")
    p2 = doc.add_paragraph("This is the target paragraph.")
    p3 = doc.add_paragraph("This is the last paragraph.")

    session = MockSession(doc)
    # Map by text
    session.text_id_map = {
        "Introduction": "h_1",
        "This is the first paragraph.": "p_1",
        "This is the target paragraph.": "p_target",
        "This is the last paragraph.": "p_3"
    }

    viz = ContextVisualizer(session)

    # Test focusing on p2 (middle). Index 2.
    # Range 1 -> Shows Indices 1, 2, 3.
    # Heading is Index 0, so it should be skipped.
    tree = viz.generate_tree_view(p2, sibling_range=1)

    print("\n--- Generated Tree (Basic) ---")
    print(tree)

    assert "Document Body" in tree
    # Heading (Index 0) should be hidden by "..."
    assert 'Heading 1: "Introduction"' not in tree
    assert "..." in tree or "(...)" in tree

    # Immediate siblings should be present
    assert 'Para: "This is the first paragraph."' in tree
    assert 'p_1' in tree # Check ID mapping works
    assert '[Para: "This is the target paragraph." (p_target)] <--- Current' in tree
    assert 'Para: "This is the last paragraph."' in tree

def test_visualizer_wide_range():
    # Test allowing wider range to see the header
    doc = Document()
    h1 = doc.add_heading("Introduction", level=1)
    p1 = doc.add_paragraph("P1")
    p2 = doc.add_paragraph("P2")

    session = MockSession(doc)
    viz = ContextVisualizer(session)

    # Focus on P2 (Index 2). Range 2 -> Shows 0, 1, 2, 3(if exists)
    tree = viz.generate_tree_view(p2, sibling_range=2)
    assert 'Heading 1: "Introduction"' in tree

def test_visualizer_truncation_and_omission():
    doc = Document()
    # Create 5 paragraphs
    p_list = []
    for i in range(5):
        text = f"Paragraph {i} " * 10 # Long text
        p_list.append(doc.add_paragraph(text))

    target = p_list[2] # Middle one (index 2)
    session = MockSession(doc)
    viz = ContextVisualizer(session)

    # Range=0 -> Only show target (and maybe parent line)
    tree_concise = viz.generate_tree_view(target, sibling_range=0)
    print("\n--- Generated Tree (Concise) ---")
    print(tree_concise)

    assert "..." in tree_concise  # Text truncation
    assert "Paragraph 2" in tree_concise
    assert "Paragraph 0" not in tree_concise # Sibling hidden
    assert "(...)" in tree_concise # Indicates siblings exist but hidden

def test_visualizer_nested_table():
    doc = Document()
    doc.add_paragraph("Outside")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p_in_cell = cell.add_paragraph("Inside Cell")

    session = MockSession(doc)
    viz = ContextVisualizer(session)

    tree = viz.generate_tree_view(p_in_cell)
    print("\n--- Generated Tree (Nested) ---")
    print(tree)

    # Parent should be identified roughly (Cell or XML Container)
    # The heuristic in code tries to identify "Cell (Parent)" or "XML Container"
    assert "Cell" in tree or "tc" in tree
    assert "Inside Cell" in tree
