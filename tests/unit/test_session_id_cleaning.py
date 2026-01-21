
import pytest
from docx_mcp_server.core.session import Session
from docx import Document

def test_get_object_with_rich_context():
    # Setup
    doc = Document()
    session = Session(session_id="test_sess", document=doc)

    # Register an object
    para = doc.add_paragraph("Test")
    para_id = session.register_object(para, "para")

    # Test cases for dirty IDs
    dirty_ids = [
        f"{para_id}\n\nCursor: after...",
        f"{para_id}   ",
        f"{para_id}\tSome other info",
        f"  {para_id}  \n"
    ]

    for dirty_id in dirty_ids:
        # Should return the object, not None
        obj = session.get_object(dirty_id)
        assert obj == para, f"Failed to retrieve object with dirty ID: {repr(dirty_id)}"

def test_get_object_normal():
    doc = Document()
    session = Session(session_id="test_sess", document=doc)
    para = doc.add_paragraph("Test")
    para_id = session.register_object(para, "para")

    assert session.get_object(para_id) == para

def test_get_object_invalid():
    doc = Document()
    session = Session(session_id="test_sess", document=doc)
    assert session.get_object("invalid_id") is None
    assert session.get_object(None) is None
    assert session.get_object("") is None
