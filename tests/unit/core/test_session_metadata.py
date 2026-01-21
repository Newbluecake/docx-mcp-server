import pytest
from docx import Document
from docx_mcp_server.core.session import Session

def test_session_metadata_registration():
    """Test registering objects with metadata."""
    doc = Document()
    session = Session(session_id="test-123", document=doc)

    # Create a paragraph
    para = doc.add_paragraph("Test Paragraph")

    # Register with metadata
    metadata = {"source": "template.docx", "original_id": "p1"}
    para_id = session.register_object(para, "para", metadata=metadata)

    # Verify object is registered
    assert session.get_object(para_id) == para

    # Verify metadata is retrieved
    retrieved_meta = session.get_metadata(para_id)
    assert retrieved_meta == metadata
    assert retrieved_meta["source"] == "template.docx"

def test_session_metadata_backward_compatibility():
    """Test that existing registration without metadata still works."""
    doc = Document()
    session = Session(session_id="test-456", document=doc)

    # Create a paragraph
    para = doc.add_paragraph("Test Paragraph 2")

    # Register WITHOUT metadata (old signature usage)
    para_id = session.register_object(para, "para")

    # Verify object is registered
    assert session.get_object(para_id) == para

    # Verify get_metadata returns None
    assert session.get_metadata(para_id) is None

def test_session_metadata_isolation():
    """Test that metadata is isolated between objects."""
    doc = Document()
    session = Session(session_id="test-789", document=doc)

    para1 = doc.add_paragraph("P1")
    para2 = doc.add_paragraph("P2")

    id1 = session.register_object(para1, "para", metadata={"id": 1})
    id2 = session.register_object(para2, "para", metadata={"id": 2})

    assert session.get_metadata(id1)["id"] == 1
    assert session.get_metadata(id2)["id"] == 2
