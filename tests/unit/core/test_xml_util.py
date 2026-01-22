import pytest
from docx import Document
from docx_mcp_server.core.xml_util import ElementNavigator, ElementManipulator

def test_element_navigator_path():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    p2 = cell.add_paragraph("P2")

    # Path for top-level paragraph
    # Document/Body[0]/Para[0]
    path1 = ElementNavigator.get_path(p1)
    assert "Body" in path1
    assert "Para[0]" in path1

    # Path for nested paragraph
    # Document/Body[0]/Table[0]/Row[0]/Cell[0]/Para[1]
    # (Note: Cell usually has an empty paragraph created by default, so P2 might be index 1)
    path2 = ElementNavigator.get_path(p2)
    assert "Table[0]" in path2
    assert "Row[0]" in path2
    assert "Cell[0]" in path2

def test_element_manipulator_insert():
    doc = Document()
    p1 = doc.add_paragraph("P1")
    p3 = doc.add_paragraph("P3")

    # Create a new paragraph but don't attach it yet (hard to do purely with python-docx API without attaching)
    # Instead, let's create one attached to end, then move it.
    p2 = doc.add_paragraph("P2")

    # Move P2 to be before P3
    # Initial state: P1, P3, P2
    # Note: When moving an element that is already in the tree using lxml methods like addprevious/addnext,
    # lxml automatically handles removing it from its old location.
    ElementManipulator.insert_xml_before(p3._element, p2._element)

    # The XML structure should now be P1, P2, P3.
    # Note: add_paragraph appends, so p2 was last.

    # We need to reload paragraphs from document to check order because doc.paragraphs might be cached or dynamic
    paras = list(doc.paragraphs)
    assert paras[0].text == "P1"
    assert paras[1].text == "P2"
    assert paras[2].text == "P3"

def test_iter_children():
    doc = Document()
    doc.add_paragraph("A")
    doc.add_table(1, 1)
    doc.add_paragraph("B")

    children = list(ElementNavigator.iter_children_xml(doc.element.body))
    # Should contain P, Tbl, P, and maybe Section properties
    tags = [c.tag.split('}')[-1] for c in children]

    assert 'p' in tags
    assert 'tbl' in tags
    assert tags.count('p') >= 2
