"""
Unit tests for Session history tracking functionality.
"""

import pytest
from docx import Document
from docx_mcp_server.core.session import Session


def test_create_commit():
    """Test commit creation."""
    session = Session(session_id="test", document=Document())

    commit_id = session.create_commit(
        operation="test_op",
        changes={"before": {"text": "old"}, "after": {"text": "new"}},
        affected_elements=["elem_1"]
    )

    assert len(session.history_stack) == 1
    assert session.current_commit_index == 0
    assert session.history_stack[0].commit_id == commit_id
    assert session.history_stack[0].operation == "test_op"


def test_get_commit_log():
    """Test retrieving commit log."""
    session = Session(session_id="test", document=Document())

    # Create multiple commits
    for i in range(5):
        session.create_commit(
            operation=f"op_{i}",
            changes={"before": {}, "after": {}},
            affected_elements=[f"elem_{i}"]
        )

    # Get last 3 commits
    log = session.get_commit_log(limit=3)

    assert len(log) == 3
    assert log[0]["operation"] == "op_4"  # Most recent first
    assert log[1]["operation"] == "op_3"
    assert log[2]["operation"] == "op_2"


def test_rollback_to_previous():
    """Test rollback to previous commit."""
    session = Session(session_id="test", document=Document())

    # Create paragraph and register it
    para = session.document.add_paragraph("initial text")
    para_id = session.register_object(para, "para")

    # Create commits
    session.create_commit(
        operation="update_1",
        changes={"before": {"text": "initial"}, "after": {"text": "updated"}},
        affected_elements=[para_id]
    )

    session.create_commit(
        operation="update_2",
        changes={"before": {"text": "updated"}, "after": {"text": "final"}},
        affected_elements=[para_id]
    )

    assert session.current_commit_index == 1

    # Rollback
    result = session.rollback()

    assert session.current_commit_index == 0
    assert len(result["rolled_back_commits"]) == 1


def test_checkout_commit():
    """Test checkout to specific commit."""
    session = Session(session_id="test", document=Document())

    # Create multiple commits
    commit_ids = []
    for i in range(3):
        commit_id = session.create_commit(
            operation=f"op_{i}",
            changes={"before": {}, "after": {}},
            affected_elements=[]
        )
        commit_ids.append(commit_id)

    assert session.current_commit_index == 2

    # Checkout to first commit
    result = session.checkout(commit_ids[0])

    assert session.current_commit_index == 0
    assert result["target_commit"] == commit_ids[0]


def test_discard_future_history():
    """Test that creating commit after rollback discards future history."""
    session = Session(session_id="test", document=Document())

    # Create 3 commits
    for i in range(3):
        session.create_commit(
            operation=f"op_{i}",
            changes={"before": {}, "after": {}},
            affected_elements=[]
        )

    assert len(session.history_stack) == 3

    # Rollback to first commit
    session.rollback(session.history_stack[0].commit_id)
    assert session.current_commit_index == 0

    # Create new commit
    session.create_commit(
        operation="new_op",
        changes={"before": {}, "after": {}},
        affected_elements=[]
    )

    # Future history should be discarded
    assert len(session.history_stack) == 2
    assert session.history_stack[1].operation == "new_op"

