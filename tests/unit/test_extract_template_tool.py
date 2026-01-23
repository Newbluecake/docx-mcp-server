"""Integration test for docx_extract_template_structure tool."""
import json
from docx import Document
from docx_mcp_server.server import docx_create, docx_extract_template_structure, docx_close

# Add parent directory to path for helpers import
import sys
import os
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

from helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    is_success,
    is_error
)


def test_extract_template_structure_integration():
    """Test the MCP tool integration."""
    # Create a session with a simple template
    session_response = docx_create()

    session_id = extract_session_id(session_response)

    # Create the document using the docx API directly
    from docx_mcp_server.server import session_manager
    session = session_manager.get_session(session_id)
    doc = session.document

    # Add elements
    doc.add_heading("Test Heading", level=1)
    doc.add_paragraph("Test paragraph")

    table = doc.add_table(rows=2, cols=2)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"Header {i + 1}"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Extract structure
    result_json = docx_extract_template_structure(session_id)
    result = json.loads(result_json)

    # Verify structure
    assert "metadata" in result
    assert "document_structure" in result
    assert len(result["document_structure"]) == 3

    # Verify element types
    assert result["document_structure"][0]["type"] == "heading"
    assert result["document_structure"][1]["type"] == "paragraph"
    assert result["document_structure"][2]["type"] == "table"

    # Verify table details
    table_data = result["document_structure"][2]
    assert table_data["rows"] == 2
    assert table_data["cols"] == 2
    assert table_data["headers"] == ["Header 1", "Header 2"]

    # Clean up
    docx_close(session_id)


def test_extract_template_structure_error_handling():
    """Test error handling for invalid session."""
    try:
        docx_extract_template_structure("invalid_session_id")
        assert False, "Should have raised ValueError"
    except ValueError as e:
        assert "not found" in str(e)
