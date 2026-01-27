"""Unit tests for FileController.switch_file auto-create session feature."""

import pytest
import tempfile
import os
from docx import Document
from docx_mcp_server.api.file_controller import FileController, set_session_manager
from docx_mcp_server.server import session_manager
from docx_mcp_server.core.global_state import global_state


class TestFileControllerAutoCreateSession:
    """Test suite for FileController.switch_file auto-create session feature (T-002)."""

    @classmethod
    def setup_class(cls):
        """Setup test environment once for all tests."""
        # Initialize file controller with session manager
        set_session_manager(session_manager, "4.0.0-test")

    def setup_method(self):
        """Setup before each test."""
        # Clear global state
        global_state.clear()
        # Clean up sessions
        for sid in list(session_manager.sessions.keys()):
            session_manager.close_session(sid)

    def teardown_method(self):
        """Cleanup after each test."""
        global_state.clear()
        for sid in list(session_manager.sessions.keys()):
            session_manager.close_session(sid)

    def test_switch_file_creates_session(self):
        """Test that switch_file automatically creates a session."""
        # Create a temporary docx file
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            # Switch to the file
            result = FileController.switch_file(tmp_path)

            # Verify result contains sessionId
            assert "sessionId" in result
            assert result["sessionId"] is not None
            assert result["currentFile"] == tmp_path

            # Verify global state is updated
            assert global_state.active_file == tmp_path
            assert global_state.active_session_id == result["sessionId"]

            # Verify session exists in session manager
            session = session_manager.get_session(result["sessionId"])
            assert session is not None
            assert session.file_path == tmp_path

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_switch_file_closes_old_session(self):
        """Test that switch_file closes the old session before creating new one."""
        # Create two temporary files
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp1:
            doc1 = Document()
            doc1.save(tmp1.name)
            tmp_path1 = tmp1.name

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp2:
            doc2 = Document()
            doc2.save(tmp2.name)
            tmp_path2 = tmp2.name

        try:
            # Switch to first file
            result1 = FileController.switch_file(tmp_path1)
            session_id1 = result1["sessionId"]

            # Verify first session exists
            assert session_manager.get_session(session_id1) is not None

            # Switch to second file
            result2 = FileController.switch_file(tmp_path2)
            session_id2 = result2["sessionId"]

            # Verify second session exists
            assert session_manager.get_session(session_id2) is not None

            # Verify first session is closed
            assert session_manager.get_session(session_id1) is None

            # Verify global state points to second session
            assert global_state.active_session_id == session_id2
            assert global_state.active_file == tmp_path2

        finally:
            if os.path.exists(tmp_path1):
                os.unlink(tmp_path1)
            if os.path.exists(tmp_path2):
                os.unlink(tmp_path2)

    def test_switch_file_session_has_correct_file_path(self):
        """Test that the created session has the correct file_path."""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            doc = Document()
            doc.save(tmp.name)
            tmp_path = tmp.name

        try:
            result = FileController.switch_file(tmp_path)
            session_id = result["sessionId"]

            session = session_manager.get_session(session_id)
            assert session.file_path == tmp_path

        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)

    def test_switch_file_multiple_times(self):
        """Test switching files multiple times."""
        files = []
        try:
            # Create 3 temporary files
            for i in range(3):
                with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
                    doc = Document()
                    doc.save(tmp.name)
                    files.append(tmp.name)

            session_ids = []
            for file_path in files:
                result = FileController.switch_file(file_path)
                session_ids.append(result["sessionId"])

                # Verify current session
                assert global_state.active_session_id == result["sessionId"]
                assert global_state.active_file == file_path

            # Verify only the last session exists
            for i, sid in enumerate(session_ids[:-1]):
                assert session_manager.get_session(sid) is None, f"Session {i} should be closed"

            # Last session should exist
            assert session_manager.get_session(session_ids[-1]) is not None

        finally:
            for file_path in files:
                if os.path.exists(file_path):
                    os.unlink(file_path)
