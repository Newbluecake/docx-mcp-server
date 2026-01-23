"""Unit tests for visualizer module."""

import pytest
from docx import Document
from docx_mcp_server.core.visualizer import DocumentVisualizer, DiffRenderer
from docx_mcp_server.core.session import Session


@pytest.fixture
def session():
    """Create a test session with a document."""
    doc = Document()
    session = Session(
        session_id="test_session",
        document=doc
    )
    return session


@pytest.fixture
def visualizer(session):
    """Create a visualizer instance."""
    return DocumentVisualizer(session)


@pytest.fixture
def diff_renderer():
    """Create a diff renderer instance."""
    return DiffRenderer()


class TestDocumentVisualizer:
    """Tests for DocumentVisualizer class."""

    def test_render_paragraph_basic(self, visualizer, session):
        """Test basic paragraph rendering."""
        para = session.document.add_paragraph("Hello World")
        para_id = session.register_object(para, "para")

        result = visualizer.render_paragraph(para, para_id)

        assert "┌" in result
        assert "└" in result
        assert "Paragraph (para_" in result
        assert "Hello World" in result

    def test_render_paragraph_with_format(self, visualizer, session):
        """Test paragraph rendering with formatting."""
        para = session.document.add_paragraph()
        run1 = para.add_run("Bold text")
        run1.bold = True
        run2 = para.add_run(" and ")
        run3 = para.add_run("italic text")
        run3.italic = True

        para_id = session.register_object(para, "para")

        result = visualizer.render_paragraph(para, para_id)

        assert "**Bold text**" in result
        assert "*italic text*" in result

    def test_render_paragraph_truncation(self, visualizer, session):
        """Test paragraph text truncation."""
        long_text = "A" * 100
        para = session.document.add_paragraph(long_text)
        para_id = session.register_object(para, "para")

        result = visualizer.render_paragraph(para, para_id)

        assert "..." in result
        assert len(result.split('\n')[2]) <= 85  # Box width + padding

    def test_render_paragraph_highlight(self, visualizer, session):
        """Test paragraph highlighting."""
        para = session.document.add_paragraph("Test")
        para_id = session.register_object(para, "para")

        result = visualizer.render_paragraph(para, para_id, highlight=True)

        assert "⭐ CURRENT" in result

    def test_render_table_basic(self, visualizer, session):
        """Test basic table rendering."""
        table = session.document.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "A"
        table.rows[0].cells[1].text = "B"
        table.rows[1].cells[0].text = "C"
        table.rows[1].cells[1].text = "D"

        table_id = session.register_object(table, "table")

        result = visualizer.render_table(table, table_id)

        assert "Table (table_" in result
        assert "A" in result
        assert "B" in result
        assert "C" in result
        assert "D" in result
        assert "├" in result
        assert "┼" in result

    def test_render_table_large(self, visualizer, session):
        """Test large table truncation."""
        table = session.document.add_table(rows=25, cols=15)
        table_id = session.register_object(table, "table")

        result = visualizer.render_table(table, table_id)

        assert "showing 20/25 rows" in result
        assert "10/15 cols" in result

    def test_render_context_middle(self, visualizer, session):
        """Test context rendering for middle element."""
        # Add multiple paragraphs
        para_ids = []
        for i in range(15):
            para = session.document.add_paragraph(f"Paragraph {i}")
            para_id = session.register_object(para, "para")
            para_ids.append(para_id)

        # Render context for middle element
        result = visualizer.render_context(para_ids[7], context_range=3)

        assert "📄 Document Context" in result
        assert "Paragraph 4" in result  # 7 - 3
        assert "Paragraph 7" in result  # Current
        assert "Paragraph 10" in result  # 7 + 3
        assert "⭐ CURRENT" in result
        assert ">>> [CURSOR] <<<" in result

    def test_render_context_start(self, visualizer, session):
        """Test context rendering for element at start."""
        para_ids = []
        for i in range(10):
            para = session.document.add_paragraph(f"Paragraph {i}")
            para_id = session.register_object(para, "para")
            para_ids.append(para_id)

        result = visualizer.render_context(para_ids[0], context_range=3)

        assert "Paragraph 0" in result
        assert "more elements above" not in result
        assert "more elements below" in result

    def test_render_context_end(self, visualizer, session):
        """Test context rendering for element at end."""
        para_ids = []
        for i in range(10):
            para = session.document.add_paragraph(f"Paragraph {i}")
            para_id = session.register_object(para, "para")
            para_ids.append(para_id)

        result = visualizer.render_context(para_ids[9], context_range=3)

        assert "Paragraph 9" in result
        assert "more elements above" in result
        assert "more elements below" not in result

    def test_render_image(self, visualizer):
        """Test image placeholder rendering."""
        result = visualizer.render_image("/path/to/image.png", "img_123")

        assert "[IMG: image.png]" in result

    def test_render_cursor(self, visualizer):
        """Test cursor marker rendering."""
        result = visualizer.render_cursor()

        assert ">>> [CURSOR] <<<" == result

    def test_truncate_text(self, visualizer):
        """Test text truncation."""
        long_text = "A" * 100
        result = visualizer._truncate_text(long_text, 50)

        assert len(result) == 50
        assert result.endswith("...")

    def test_truncate_text_short(self, visualizer):
        """Test text truncation with short text."""
        short_text = "Hello"
        result = visualizer._truncate_text(short_text, 50)

        assert result == "Hello"


class TestDiffRenderer:
    """Tests for DiffRenderer class."""

    def test_render_diff_basic(self, diff_renderer):
        """Test basic diff rendering."""
        old_content = "Hello World"
        new_content = "Hello Claude"

        result = diff_renderer.render_diff(
            old_content, new_content,
            "para_123", "Paragraph"
        )

        assert "🔄 Changes" in result
        assert "Paragraph (para_123)" in result
        assert "-" in result
        assert "+" in result

    def test_render_diff_multiline(self, diff_renderer):
        """Test multiline diff rendering."""
        old_content = "Line 1\nLine 2\nLine 3"
        new_content = "Line 1\nLine 2 modified\nLine 3"

        result = diff_renderer.render_diff(
            old_content, new_content,
            "para_123", "Paragraph"
        )

        assert "Line 1" in result
        assert "- │ Line 2" in result
        assert "+ │ Line 2 modified" in result
        assert "Line 3" in result

    def test_compute_line_diff_equal(self, diff_renderer):
        """Test line diff computation for equal content."""
        old_lines = ["Line 1", "Line 2"]
        new_lines = ["Line 1", "Line 2"]

        result = diff_renderer._compute_line_diff(old_lines, new_lines)

        assert len(result) == 2
        assert all(prefix == ' ' for prefix, _ in result)

    def test_compute_line_diff_delete(self, diff_renderer):
        """Test line diff computation for deletion."""
        old_lines = ["Line 1", "Line 2", "Line 3"]
        new_lines = ["Line 1", "Line 3"]

        result = diff_renderer._compute_line_diff(old_lines, new_lines)

        # Should have: Line 1 (equal), Line 2 (delete), Line 3 (equal)
        assert any(prefix == '-' and line == "Line 2" for prefix, line in result)

    def test_compute_line_diff_insert(self, diff_renderer):
        """Test line diff computation for insertion."""
        old_lines = ["Line 1", "Line 3"]
        new_lines = ["Line 1", "Line 2", "Line 3"]

        result = diff_renderer._compute_line_diff(old_lines, new_lines)

        # Should have: Line 1 (equal), Line 2 (insert), Line 3 (equal)
        assert any(prefix == '+' and line == "Line 2" for prefix, line in result)

    def test_compute_line_diff_replace(self, diff_renderer):
        """Test line diff computation for replacement."""
        old_lines = ["Line 1", "Old Line", "Line 3"]
        new_lines = ["Line 1", "New Line", "Line 3"]

        result = diff_renderer._compute_line_diff(old_lines, new_lines)

        # Should have: Line 1 (equal), Old Line (delete), New Line (insert), Line 3 (equal)
        assert any(prefix == '-' and line == "Old Line" for prefix, line in result)
        assert any(prefix == '+' and line == "New Line" for prefix, line in result)
