"""End-to-end tests for template extraction feature."""
import json
import pytest
from tests.helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    is_success,
    is_error
)
from tests.helpers.session_helpers import setup_active_session, teardown_active_session
from docx import Document
from docx_mcp_server.tools.content_tools import docx_extract_template_structure
from docx_mcp_server.tools.session_tools import docx_save
def test_extract_complete_template():
    """Test extracting a complete template with all element types."""
    session_id = setup_active_session()
    # Access the document
    from docx_mcp_server.server import session_manager
    session = session_manager.get_session(session_id)
    doc = session.document

    # Build a complete template
    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_paragraph("This is the introduction paragraph with some content.")

    doc.add_heading("Section 1.1: Details", level=2)
    doc.add_paragraph("More details here in a regular paragraph.")

    # Add table with bold headers
    table = doc.add_table(rows=4, cols=3)
    headers = ["Name", "Age", "Department"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header_text
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Add data rows
    for row_idx in range(1, 4):
        for col_idx in range(3):
            table.rows[row_idx].cells[col_idx].text = f"Data-{row_idx}-{col_idx}"

    # Extract structure
    result_json = docx_extract_template_structure()
    result = json.loads(result_json)

    # Verify metadata
    assert "metadata" in result
    assert "extracted_at" in result["metadata"]
    assert "docx_version" in result["metadata"]

    # Verify structure
    structure = result["document_structure"]
    assert len(structure) == 5

    # Verify element types and order
    assert structure[0]["type"] == "heading"
    assert structure[0]["level"] == 1
    assert "Introduction" in structure[0]["text"]

    assert structure[1]["type"] == "paragraph"
    assert "introduction paragraph" in structure[1]["text"]

    assert structure[2]["type"] == "heading"
    assert structure[2]["level"] == 2
    assert "Details" in structure[2]["text"]

    assert structure[3]["type"] == "paragraph"

    assert structure[4]["type"] == "table"
    assert structure[4]["rows"] == 4
    assert structure[4]["cols"] == 3
    assert structure[4]["header_row"] == 0
    assert structure[4]["headers"] == ["Name", "Age", "Department"]

    # Clean up
    teardown_active_session()


def test_extract_with_header_detection():
    """Test header detection with bold formatting."""
    session_id = setup_active_session()
    from docx_mcp_server.server import session_manager
    session = session_manager.get_session(session_id)
    doc = session.document

    # Create table with bold headers
    table = doc.add_table(rows=2, cols=2)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"Column{i + 1}"
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Extract
    result_json = docx_extract_template_structure(session_id)
    result = json.loads(result_json)

    # Verify header detection
    table_data = result["document_structure"][0]
    assert table_data["type"] == "table"
    assert table_data["header_row"] == 0
    assert table_data["headers"] == ["Column1", "Column2"]

    teardown_active_session()


def test_extract_header_detection_fail():
    """Test that tables without detectable headers are skipped."""
    session_id = setup_active_session()
    from docx_mcp_server.server import session_manager
    session = session_manager.get_session(session_id)
    doc = session.document

    # Create table WITHOUT bold headers or background
    table = doc.add_table(rows=2, cols=2)
    for i, cell in enumerate(table.rows[0].cells):
        cell.text = f"NotHeader{i + 1}"
        # Deliberately not making it bold

    # Extract - should skip the table
    result_json = docx_extract_template_structure(session_id)
    result = json.loads(result_json)

    # Table should be skipped due to header detection failure
    assert len(result["document_structure"]) == 0

    teardown_active_session()
