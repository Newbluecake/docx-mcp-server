import pytest
import ast
import re
from tests.helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    extract_all_metadata,
    is_success,
    is_error,
    create_session_with_file,
)
import os
import json
from docx import Document
from docx_mcp_server.tools.session_tools import docx_save
from docx_mcp_server.tools.content_tools import docx_read_content, docx_find_paragraphs
from docx_mcp_server.tools.run_tools import docx_insert_run, docx_set_font

class TestLoadEditE2E:
    def test_full_workflow(self, tmp_path):
        # 1. Prepare initial document
        initial_doc_path = tmp_path / "initial.docx"
        doc = Document()
        doc.add_paragraph("Title Paragraph")
        doc.add_paragraph("Target: This needs editing.")
        doc.add_paragraph("Footer Paragraph")
        doc.save(str(initial_doc_path))

        # 2. Load the document
        session_id = create_session_with_file(str(initial_doc_path))
        assert session_id is not None

        # 3. Read content to confirm
        content = docx_read_content()
        assert "Target: This needs editing." in content

        # 4. Find specific paragraph
        results_resp = docx_find_paragraphs("Target:")
        results = []
        if is_success(results_resp):
            meta = extract_all_metadata(results_resp)
            if 'paragraphs' in meta:
                val = meta['paragraphs']
                if isinstance(val, str):
                    try:
                        results = json.loads(val)
                    except Exception:
                        try:
                            results = ast.literal_eval(val)
                        except Exception:
                            results = []
                else:
                    results = val or []
        else:
            # Legacy behavior: tool may return a raw JSON list
            try:
                results = json.loads(results_resp)
            except Exception:
                match = re.search(r'\[.*\]', results_resp, re.DOTALL)
                if match:
                    try:
                        results = json.loads(match.group(0))
                    except Exception:
                        results = []

        assert len(results) == 1
        para_id = results[0]["id"]

        # 5. Edit (Append text)
        result = docx_insert_run(" [EDITED]", position=f"inside:{para_id}")
        assert is_success(result)
        run_id = extract_element_id(result)
        assert run_id is not None

        # 6. Style the new run
        docx_set_font(run_id, bold=True, color_hex="FF0000")

        # 7. Save to new file
        final_doc_path = tmp_path / "final.docx"
        docx_save(str(final_doc_path))

        # 8. Verification
        doc_final = Document(str(final_doc_path))
        full_text = "\n".join([p.text for p in doc_final.paragraphs])

        assert "Target: This needs editing. [EDITED]" in full_text

        # Verify formatting
        # Find the paragraph again in the loaded doc object
        target_para = None
        for p in doc_final.paragraphs:
            if "Target:" in p.text:
                target_para = p
                break

        assert target_para is not None
        # Check the last run for bold/color
        last_run = target_para.runs[-1]
        assert last_run.text == " [EDITED]"
        assert last_run.bold is True
        # Color check might be complex depending on how python-docx stores it,
        # but we know we called the tool.
        assert str(last_run.font.color.rgb) == "FF0000"
