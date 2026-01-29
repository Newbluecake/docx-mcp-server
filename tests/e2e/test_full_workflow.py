import pytest
from tests.helpers import (
    extract_session_id,
    extract_element_id,
    extract_metadata_field,
    extract_all_metadata,
    is_success,
    is_error
)
from tests.helpers.session_helpers import setup_active_session, teardown_active_session
import os
import json
from docx_mcp_server.tools.session_tools import docx_save
from docx_mcp_server.tools.paragraph_tools import docx_insert_paragraph
from docx_mcp_server.tools.run_tools import docx_insert_run
from docx_mcp_server.tools.format_tools import docx_set_properties
from docx_mcp_server.tools.table_tools import (
    docx_insert_table,
    docx_fill_table,
    docx_find_table,
    docx_copy_table
)
from docx_mcp_server.tools.advanced_tools import docx_replace_text
from docx_mcp_server.server import session_manager

def _extract(response):
    if not is_success(response):
        raise ValueError(f"Tool failed: {response}")
    return extract_all_metadata(response)

# E2E test simulating a user workflow
def test_full_workflow(tmp_path):
    output_file = tmp_path / "report.docx"
    output_path = str(output_file)

    # 1. Create Document
    sid = setup_active_session()

    # 2. Add Title
    title_resp = docx_insert_paragraph("Monthly Report", position="end:document_body", style="Heading 1")
    title_id = _extract(title_resp)["element_id"]
    docx_set_properties(json.dumps({"paragraph_format": {"alignment": "center"}}), title_id)

    # 3. Add Intro text
    para_resp = docx_insert_paragraph("This report covers the period: {{period}}.", position="end:document_body")
    para_id = _extract(para_resp)["element_id"]

    # 4. Create Data Table
    table_resp = docx_insert_table(rows=1, cols=3, position="end:document_body")
    table_id = _extract(table_resp)["element_id"]

    # 5. Fill Table Headers & Data
    data = [
        ["Metric", "Value", "Status"],
        ["Revenue", "$1M", "Green"],
        ["Users", "50k", "Yellow"]
    ]
    docx_fill_table(json.dumps(data), table_id)

    # 6. Find Table (Simulation of navigation)
    found_resp = docx_find_table("Revenue")
    found_id = _extract(found_resp)["element_id"]
    # Verify it found the correct table (should be same underlying object as table_id)
    # IDs might differ if re-registered, but logic holds.

    # 7. Copy Table (Template usage)
    copy_resp = docx_copy_table(found_id, position="end:document_body")
    copy_id = _extract(copy_resp)["element_id"]

    # 8. Modify Copy (Fill with different data)
    new_data = [
        ["Metric", "Value", "Status"],
        ["Cost", "$500k", "Red"]
    ]
    docx_fill_table(json.dumps(new_data), copy_id)

    # 9. Replace Placeholders
    docx_replace_text("{{period}}", "January 2026")

    # 10. Save
    docx_save(output_path)

    # 11. Close
    teardown_active_session()

    # Verification
    assert os.path.exists(output_path)

    # Verify content using python-docx directly
    from docx import Document
    doc = Document(output_path)

    # Check text replacement
    full_text = "\n".join([p.text for p in doc.paragraphs])
    assert "January 2026" in full_text
    assert "{{period}}" not in full_text

    # Check tables
    assert len(doc.tables) == 2

    t1 = doc.tables[0]
    assert t1.cell(1, 0).text == "Revenue"

    t2 = doc.tables[1]
    assert t2.cell(1, 0).text == "Cost"
