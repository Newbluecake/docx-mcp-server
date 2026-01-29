"""E2E test for paragraph copying functionality."""

import sys
import os
import tempfile
import json
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from tests.helpers import extract_session_id, extract_element_id
from tests.helpers.session_helpers import setup_active_session, teardown_active_session

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../src')))

from docx_mcp_server.tools.session_tools import docx_save
from docx_mcp_server.tools.paragraph_tools import docx_insert_paragraph, docx_copy_paragraph
from docx_mcp_server.tools.run_tools import docx_insert_run, docx_set_font
from docx_mcp_server.tools.format_tools import docx_set_alignment

def _extract(response):
    # Updated to handle Markdown response
    eid = extract_element_id(response)
    if not eid:
         # Fallback for checking success if no element_id expected
         if "✅ Success" not in response and "status" not in response:
             raise ValueError(f"Tool failed: {response}")
    return {"element_id": eid}

def test_copy_paragraph_e2e():
    """Test copying paragraphs and verify in saved document."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        # Create session
        session_id = setup_active_session()
        # Create original paragraph with formatting
        p1_resp = docx_insert_paragraph("", position="end:document_body")
        para1_id = _extract(p1_resp)["element_id"]

        docx_insert_run("This is ", position=f"inside:{para1_id}")

        r2_resp = docx_insert_run("formatted ", position=f"inside:{para1_id}")
        run2_id = _extract(r2_resp)["element_id"]

        docx_set_font(run2_id, bold=True, color_hex="FF0000")
        docx_insert_run("text", position=f"inside:{para1_id}")

        docx_set_alignment(para1_id, "center")

        # Copy the paragraph
        cp_resp = docx_copy_paragraph(para1_id, position="end:document_body")
        _extract(cp_resp) # verify success

        # Add another paragraph
        p3_resp = docx_insert_paragraph("Normal paragraph", position="end:document_body")
        para3_id = _extract(p3_resp)["element_id"]

        # Copy it too
        cp2_resp = docx_copy_paragraph(para3_id, position="end:document_body")
        _extract(cp2_resp)

        # Save
        docx_save(output_path)
        teardown_active_session()

        # Verify the saved document
        doc = Document(output_path)
        assert len(doc.paragraphs) == 4

        # Check first two paragraphs are identical
        assert doc.paragraphs[0].text == doc.paragraphs[1].text
        assert doc.paragraphs[0].text == "This is formatted text"
        assert doc.paragraphs[0].alignment == doc.paragraphs[1].alignment

        # Check last two paragraphs are identical
        assert doc.paragraphs[2].text == doc.paragraphs[3].text
        assert doc.paragraphs[2].text == "Normal paragraph"

        print(f"✓ E2E test passed! Document saved to {output_path}")

    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)

if __name__ == "__main__":
    test_copy_paragraph_e2e()
