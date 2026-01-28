"""E2E test for text update functionality."""

import sys
import os
import tempfile
import json
from docx import Document
from tests.helpers import extract_session_id, extract_element_id
from tests.helpers.session_helpers import setup_active_session, teardown_active_session

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), '../../src')))

from docx_mcp_server.tools.session_tools import docx_save, docx_close
from docx_mcp_server.tools.paragraph_tools import docx_insert_paragraph, docx_update_paragraph_text
from docx_mcp_server.tools.run_tools import docx_insert_run, docx_set_font, docx_update_run_text

def test_update_text_e2e():
    """Test text updates and verify in saved document."""
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        output_path = tmp.name

    try:
        # Create session
        session_id = setup_active_session()
        # Create original content
        result = docx_insert_paragraph("Original paragraph 1", position="end:document_body")
        para1_id = extract_element_id(result)

        result = docx_insert_paragraph("", position="end:document_body")
        para2_id = extract_element_id(result)

        result = docx_insert_run("Original run", position=f"inside:{para2_id}")
        run1_id = extract_element_id(result)

        docx_set_font(run1_id, bold=True, size=14)

        # Update paragraph text
        docx_update_paragraph_text(para1_id, "Updated paragraph 1")

        # Update run text (should preserve formatting)
        docx_update_run_text(run1_id, "Updated run")

        # Save
        docx_save(output_path)
        teardown_active_session()

        # Verify the saved document
        doc = Document(output_path)
        assert len(doc.paragraphs) == 2

        # Check updated paragraph
        assert doc.paragraphs[0].text == "Updated paragraph 1"
        assert "Original paragraph 1" not in doc.paragraphs[0].text

        # Check updated run
        assert doc.paragraphs[1].text == "Updated run"
        assert "Original run" not in doc.paragraphs[1].text

        # Verify formatting is preserved
        assert doc.paragraphs[1].runs[0].font.bold == True
        assert doc.paragraphs[1].runs[0].font.size.pt == 14

        print(f"✓ E2E test passed! Document saved to {output_path}")

    finally:
        if os.path.exists(output_path):
            os.unlink(output_path)

if __name__ == "__main__":
    test_update_text_e2e()
