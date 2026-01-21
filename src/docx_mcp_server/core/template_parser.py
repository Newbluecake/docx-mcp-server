"""Template structure extraction module for docx-mcp-server."""

import time
from typing import Dict, List, Any, Optional
from docx.document import Document as DocumentType
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TemplateParser:
    """Parser for extracting structured information from Word documents."""

    def __init__(self):
        """Initialize the TemplateParser."""
        pass

    def extract_structure(self, document: DocumentType) -> Dict[str, Any]:
        """
        Extract the complete structure of a Word document.

        Args:
            document: The python-docx Document object to parse.

        Returns:
            dict: Structured representation with metadata and document_structure.
        """
        from docx_mcp_server.server import VERSION

        result = {
            "metadata": {
                "extracted_at": time.strftime("%Y-%m-%dT%H:%M:%SZ", time.gmtime()),
                "docx_version": VERSION
            },
            "document_structure": []
        }

        # Traverse document elements in order
        for element in document.element.body:
            tag = element.tag.split('}')[-1] if '}' in element.tag else element.tag

            if tag == 'p':  # Paragraph
                para = Paragraph(element, document)
                # Check if it's a heading
                if para.style and para.style.name and 'Heading' in para.style.name:
                    result["document_structure"].append(self.extract_heading_structure(para))
                elif para.text.strip():  # Only add non-empty paragraphs
                    result["document_structure"].append(self.extract_paragraph_structure(para))

            elif tag == 'tbl':  # Table
                table = Table(element, document)
                try:
                    result["document_structure"].append(self.extract_table_structure(table))
                except ValueError:
                    # Skip tables without detectable headers
                    pass

        return result

    def detect_header_row(self, table: Table) -> int:
        """
        Detect header row in a table (intelligent detection).

        Detection rules:
        1. First row all cells bold -> header
        2. First row all cells have background color -> header
        3. Both conditions met -> header
        4. Neither condition met -> raise ValueError

        Args:
            table: The table to analyze.

        Returns:
            int: Header row index (usually 0).

        Raises:
            ValueError: If header row cannot be detected.
        """
        first_row = table.rows[0]

        # Check if all cells are bold
        all_bold = all(
            any(run.bold for run in cell.paragraphs[0].runs if run.bold)
            for cell in first_row.cells
            if cell.paragraphs and cell.paragraphs[0].runs
        )

        # Check if all cells have background color
        all_has_bg = all(
            self._has_background_color(cell)
            for cell in first_row.cells
        )

        if all_bold or all_has_bg:
            return 0
        else:
            raise ValueError("无法检测表头行：第一行既无加粗也无背景色")

    def _has_background_color(self, cell) -> bool:
        """Check if cell has non-white background color."""
        try:
            tc_pr = cell._element.tcPr
            if tc_pr is None:
                return False

            # Try to get shading element using xpath
            shading_list = tc_pr.xpath('.//w:shd')
            if not shading_list:
                return False

            shading = shading_list[0]
            fill = shading.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')

            if fill is None:
                return False

            return fill.upper() not in ["FFFFFF", "NONE", "AUTO", ""]
        except Exception:
            return False

    def extract_heading_structure(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Extract heading structure from a paragraph.

        Args:
            paragraph: Paragraph with Heading style.

        Returns:
            dict: Heading structure with type, level, text, and style.
        """
        # Extract level from style name (e.g., "Heading 1" -> 1)
        style_name = paragraph.style.name
        level = 1
        if "Heading" in style_name:
            parts = style_name.split()
            if len(parts) > 1 and parts[1].isdigit():
                level = int(parts[1])

        # Extract text and style from first run
        text = paragraph.text
        style = {}

        if paragraph.runs:
            run = paragraph.runs[0]
            color_hex = "000000"
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                color_hex = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
            style = {
                "font": run.font.name or "默认",
                "size": run.font.size.pt if run.font.size else 16,
                "bold": run.font.bold or False,
                "color": color_hex
            }

        return {
            "type": "heading",
            "level": level,
            "text": text,
            "style": style
        }

    def extract_paragraph_structure(self, paragraph: Paragraph) -> Dict[str, Any]:
        """
        Extract paragraph structure.

        Args:
            paragraph: Paragraph to extract.

        Returns:
            dict: Paragraph structure with type, text, and style.
        """
        text = paragraph.text

        # Extract style from first run
        style = {
            "font": "默认",
            "size": 12,
            "bold": False,
            "italic": False,
            "underline": False,
            "color": "000000",
            "alignment": "left",
            "left_indent": 0.0,
            "right_indent": 0.0,
            "first_line_indent": 0.0
        }

        if paragraph.runs:
            run = paragraph.runs[0]
            style["font"] = run.font.name or "默认"
            style["size"] = run.font.size.pt if run.font.size else 12
            style["bold"] = run.font.bold or False
            style["italic"] = run.font.italic or False
            style["underline"] = run.font.underline or False
            if run.font.color and run.font.color.rgb:
                rgb = run.font.color.rgb
                style["color"] = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
            else:
                style["color"] = "000000"

        # Extract alignment
        alignment_map = {
            WD_ALIGN_PARAGRAPH.LEFT: "left",
            WD_ALIGN_PARAGRAPH.CENTER: "center",
            WD_ALIGN_PARAGRAPH.RIGHT: "right",
            WD_ALIGN_PARAGRAPH.JUSTIFY: "justify"
        }
        style["alignment"] = alignment_map.get(paragraph.alignment, "left")

        # Extract indents
        if paragraph.paragraph_format.left_indent:
            style["left_indent"] = paragraph.paragraph_format.left_indent.inches
        if paragraph.paragraph_format.right_indent:
            style["right_indent"] = paragraph.paragraph_format.right_indent.inches
        if paragraph.paragraph_format.first_line_indent:
            style["first_line_indent"] = paragraph.paragraph_format.first_line_indent.inches

        return {
            "type": "paragraph",
            "text": text,
            "style": style
        }

    def extract_table_structure(self, table: Table) -> Dict[str, Any]:
        """
        Extract table structure including headers and styling.

        Args:
            table: Table to extract.

        Returns:
            dict: Table structure with type, dimensions, headers, and style.

        Raises:
            ValueError: If header row cannot be detected.
        """
        rows = len(table.rows)
        cols = len(table.columns)

        # Detect header row
        header_row = self.detect_header_row(table)

        # Extract headers
        headers = [cell.text for cell in table.rows[header_row].cells]

        # Extract basic style (simplified for minimal implementation)
        style = {
            "border": {},
            "cell_style": {
                "font": "默认",
                "size": 12,
                "alignment": "left"
            }
        }

        # Try to extract cell style from first data cell
        if rows > 1 and table.rows[1].cells:
            first_cell = table.rows[1].cells[0]
            if first_cell.paragraphs and first_cell.paragraphs[0].runs:
                run = first_cell.paragraphs[0].runs[0]
                style["cell_style"]["font"] = run.font.name or "默认"
                style["cell_style"]["size"] = run.font.size.pt if run.font.size else 12

        return {
            "type": "table",
            "rows": rows,
            "cols": cols,
            "header_row": header_row,
            "headers": headers,
            "style": style
        }
