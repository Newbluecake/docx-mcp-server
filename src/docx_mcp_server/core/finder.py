import os
import time
from typing import List, Dict, Any, Optional, Union
from docx.document import Document
from docx.text.paragraph import Paragraph
from docx.table import Table

class Finder:
    """Helper class for finding elements in a document."""

    def __init__(self, document: Document):
        self.document = document

    def get_table_by_index(self, index: int) -> Optional[Table]:
        """Get table by 0-based index."""
        try:
            return self.document.tables[index]
        except IndexError:
            return None

    def find_paragraphs_by_text(self, text: str, case_sensitive: bool = False) -> List[Paragraph]:
        """Find paragraphs containing specific text."""
        matches = []
        for p in self.document.paragraphs:
            if case_sensitive:
                if text in p.text:
                    matches.append(p)
            else:
                if text.lower() in p.text.lower():
                    matches.append(p)
        return matches

    def find_tables_by_text(self, text: str, case_sensitive: bool = False) -> List[Table]:
        """Find tables that contain specific text in any cell."""
        matches = []
        for table in self.document.tables:
            found = False
            for row in table.rows:
                for cell in row.cells:
                    cell_text = cell.text
                    if not case_sensitive:
                        cell_text = cell_text.lower()
                        query = text.lower()
                    else:
                        query = text

                    if query in cell_text:
                        matches.append(table)
                        found = True
                        break
                if found:
                    break
        return matches

def list_docx_files(
    root_path: str = ".",
    recursive: bool = False,
    include_meta: bool = False
) -> List[Any]:
    """List .docx files with optional recursion and metadata."""
    if not os.path.exists(root_path):
        raise ValueError(f"Directory not found: {root_path}")

    collected = []

    def _add(path: str):
        rel_path = os.path.relpath(path, start=root_path)
        rel_path = f"./{rel_path}" if not rel_path.startswith(".") else rel_path
        if include_meta:
            stat = os.stat(path)
            collected.append({
                "path": rel_path.replace(os.sep, "/"),
                "size": stat.st_size,
                "modified": time.strftime("%Y-%m-%dT%H:%M:%S", time.localtime(stat.st_mtime)),
            })
        else:
            collected.append(rel_path.replace(os.sep, "/"))

    try:
        if recursive:
            for dirpath, _, filenames in os.walk(root_path):
                for entry in filenames:
                    if entry.lower().endswith(".docx") and not entry.startswith("~$"):
                        _add(os.path.join(dirpath, entry))
        else:
            for entry in os.listdir(root_path):
                if entry.lower().endswith(".docx") and not entry.startswith("~$"):
                    _add(os.path.join(root_path, entry))
    except Exception as e:
        raise ValueError(f"Failed to list files: {str(e)}")

    return sorted(collected, key=lambda x: x["path"] if include_meta else x)
