import uuid
import time
import os
import logging
from typing import Dict, Any, Optional, List
import shutil
from dataclasses import dataclass, field
from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx_mcp_server.core.validators import validate_path_safety
from docx_mcp_server.core.cursor import Cursor
from docx_mcp_server.core.commit import Commit
from docx_mcp_server.preview.manager import PreviewManager

logger = logging.getLogger(__name__)

@dataclass
class Session:
    session_id: str
    document: DocumentType
    created_at: float = field(default_factory=time.time)
    last_accessed: float = field(default_factory=time.time)
    file_path: Optional[str] = None
    # Registry to map string IDs to internal python-docx objects
    # structure: { "para_123": <docx.text.paragraph.Paragraph>, "run_456": ... }
    object_registry: Dict[str, Any] = field(default_factory=dict)
    # Metadata registry to store additional info about objects (e.g., source info)
    element_metadata: Dict[str, Dict[str, Any]] = field(default_factory=dict)
    last_created_id: Optional[str] = None
    last_accessed_id: Optional[str] = None
    auto_save: bool = False
    backup_on_save: bool = False
    backup_dir: Optional[str] = None
    backup_suffix: Optional[str] = None
    cursor: Cursor = field(default_factory=Cursor)

    # Context stack for nested operations
    context_stack: List[str] = field(default_factory=list)

    # Reverse ID mapping cache: id(element._element) -> element_id
    _element_id_cache: Dict[int, str] = field(default_factory=dict)

    # Preview controller for handling live updates
    preview_controller: Any = field(init=False)

    # History tracking for change management
    history_stack: List[Commit] = field(default_factory=list)
    current_commit_index: int = -1

    def __post_init__(self):
        self.preview_controller = PreviewManager.get_controller()

    def touch(self):
        self.last_accessed = time.time()

    def update_context(self, element_id: str, action: str = "access"):
        """Update context pointers based on action type."""
        self.last_accessed_id = element_id
        if action == "create":
            self.last_created_id = element_id
        logger.debug(f"Context updated: element_id={element_id}, action={action}")

        # Trigger auto-save if enabled
        if self.auto_save and self.file_path:
            try:
                self._save_with_optional_backup(
                    self.file_path,
                    backup=self.backup_on_save,
                    backup_dir=self.backup_dir,
                    backup_suffix=self.backup_suffix,
                )
                logger.debug(f"Auto-save successful: {self.file_path}")
            except Exception as e:
                logger.warning(f"Auto-save failed for {self.file_path}: {e}")

    def register_object(self, obj: Any, prefix: str = "obj", metadata: Optional[Dict[str, Any]] = None) -> str:
        """Register a docx object and return its ID, optionally storing metadata."""
        obj_id = f"{prefix}_{uuid.uuid4().hex[:8]}"
        self.object_registry[obj_id] = obj

        # Update reverse cache for context lookup
        if hasattr(obj, '_element'):
            self._element_id_cache[id(obj._element)] = obj_id

        if metadata:
            self.element_metadata[obj_id] = metadata

        logger.debug(f"Object registered: {obj_id} (type={type(obj).__name__})")
        return obj_id

    def get_metadata(self, obj_id: str) -> Optional[Dict[str, Any]]:
        """Retrieve metadata for a registered object."""
        return self.element_metadata.get(obj_id)

    def push_context(self, element_id: str):
        """Push element ID onto context stack for nested operations."""
        self.context_stack.append(element_id)
        logger.debug(f"Context pushed: {element_id}, stack depth: {len(self.context_stack)}")

    def pop_context(self) -> Optional[str]:
        """Pop element ID from context stack."""
        if self.context_stack:
            element_id = self.context_stack.pop()
            logger.debug(f"Context popped: {element_id}, stack depth: {len(self.context_stack)}")
            return element_id
        return None

    def get_current_context(self) -> Optional[str]:
        """Get current context without popping."""
        return self.context_stack[-1] if self.context_stack else None

    def get_object(self, obj_id: str) -> Optional[Any]:
        if not obj_id or not isinstance(obj_id, str):
            return None

        # Clean ID: remove whitespace and extra context (take first token)
        # This handles cases where LLM returns "id\n\nContext..."
        clean_id = obj_id.strip().split()[0] if obj_id.strip() else ""

        return self.object_registry.get(clean_id)

    def _get_element_id(self, element: Any, auto_register: bool = True) -> Optional[str]:
        """Get element ID from cache, optionally auto-register if not found.

        Args:
            element: The docx element (Paragraph, Table, Cell, Run, etc.)
            auto_register: If True, automatically register element if not in cache

        Returns:
            Element ID string if found/registered, None otherwise
        """
        if not hasattr(element, '_element'):
            logger.debug(f"Element {type(element).__name__} has no _element attribute")
            return None

        element_key = id(element._element)

        # Check cache first
        if element_key in self._element_id_cache:
            element_id = self._element_id_cache[element_key]
            logger.debug(f"Element ID cache hit: {element_id} (type={type(element).__name__})")
            return element_id

        # Cache miss
        if auto_register:
            prefix = "para" if isinstance(element, Paragraph) else \
                     "table" if isinstance(element, Table) else \
                     "cell" if isinstance(element, _Cell) else \
                     "run" if isinstance(element, Run) else "obj"
            element_id = self.register_object(element, prefix)
            logger.debug(f"Element ID cache miss, registered: {element_id} (type={type(element).__name__})")
            return element_id

        logger.debug(f"Element ID cache miss, auto_register=False (type={type(element).__name__})")
        return None

    def _get_siblings(self, parent: Any) -> List[Any]:
        """Get all child elements from parent container."""
        elements = []

        # For Document, use paragraphs and tables properties
        if hasattr(parent, 'paragraphs') and hasattr(parent, 'tables') and not isinstance(parent, _Cell):
            # Merge paragraphs and tables in document order
            para_list = list(parent.paragraphs)
            table_list = list(parent.tables)

            # Get all body elements in order
            for child in parent._element.body:
                if child.tag.endswith('p'):
                    # Find matching paragraph
                    for p in para_list:
                        if p._element is child:
                            elements.append(p)
                            break
                elif child.tag.endswith('tbl'):
                    # Find matching table
                    for t in table_list:
                        if t._element is child:
                            elements.append(t)
                            break
        # For Cell, use paragraphs and tables
        elif isinstance(parent, _Cell):
            para_list = list(parent.paragraphs)
            table_list = list(parent.tables)

            # Get all cell elements in order
            for child in parent._element:
                if child.tag.endswith('p'):
                    for p in para_list:
                        if p._element is child:
                            elements.append(p)
                            break
                elif child.tag.endswith('tbl'):
                    for t in table_list:
                        if t._element is child:
                            elements.append(t)
                            break

        # For Paragraph, use runs
        elif isinstance(parent, Paragraph):
            return list(parent.runs)

        return elements

    def _save_with_optional_backup(
        self,
        target_path: str,
        backup: bool = False,
        backup_dir: Optional[str] = None,
        backup_suffix: Optional[str] = None,
    ) -> Optional[str]:
        """
        Save document to target_path. If backup is True and file exists, create a timestamped backup copy.

        Returns backup path if created.
        """
        abs_target = os.path.abspath(target_path)
        parent_dir = os.path.dirname(abs_target) or "."
        if not os.path.exists(parent_dir):
            raise ValueError(f"Parent directory does not exist: {parent_dir}")

        backup_path = None
        if backup and os.path.exists(abs_target):
            ts = time.strftime("%Y%m%d-%H%M%S")
            base, ext = os.path.splitext(os.path.basename(abs_target))
            suffix = backup_suffix if backup_suffix else "-backup"
            backup_name = f"{base}{suffix}-{ts}{ext}"
            target_dir = os.path.abspath(backup_dir) if backup_dir else parent_dir
            os.makedirs(target_dir, exist_ok=True)
            backup_path = os.path.join(target_dir, backup_name)
            shutil.copy2(abs_target, backup_path)

        self.document.save(abs_target)
        return backup_path

    def create_commit(
        self,
        operation: str,
        changes: Dict[str, Any],
        affected_elements: List[str],
        description: str = ""
    ) -> str:
        """Create a new commit record."""
        commit = Commit.create(
            operation=operation,
            changes=changes,
            affected_elements=affected_elements,
            description=description
        )

        # If not at latest, discard future history
        if self.current_commit_index < len(self.history_stack) - 1:
            self.history_stack = self.history_stack[:self.current_commit_index + 1]

        self.history_stack.append(commit)
        self.current_commit_index = len(self.history_stack) - 1

        logger.info(f"Commit created: {commit.commit_id} ({operation})")
        return commit.commit_id

    def get_commit_log(self, limit: int = 10) -> List[Dict[str, Any]]:
        """Get recent commit history."""
        commits = self.history_stack[-limit:]
        commits.reverse()
        return [commit.to_dict() for commit in commits]

    def rollback(self, commit_id: Optional[str] = None) -> Dict[str, Any]:
        """Rollback to specified commit or previous commit."""
        if not self.history_stack:
            raise ValueError("No commits to rollback")

        # Determine target index
        if commit_id is None:
            target_index = self.current_commit_index - 1
        else:
            target_index = None
            for i, commit in enumerate(self.history_stack):
                if commit.commit_id == commit_id:
                    target_index = i
                    break
            if target_index is None:
                raise ValueError(f"Commit {commit_id} not found")

        if target_index < -1:
            raise ValueError("Cannot rollback beyond initial state")

        # Apply reverse changes
        rollback_info = {
            "rolled_back_commits": [],
            "restored_elements": []
        }

        for i in range(self.current_commit_index, target_index, -1):
            commit = self.history_stack[i]
            self._apply_reverse_changes(commit)
            rollback_info["rolled_back_commits"].append(commit.commit_id)
            rollback_info["restored_elements"].extend(commit.affected_elements)

        self.current_commit_index = target_index
        logger.info(f"Rolled back to index {target_index}")
        return rollback_info

    def checkout(self, commit_id: str) -> Dict[str, Any]:
        """Checkout to specified commit state."""
        target_index = None
        for i, commit in enumerate(self.history_stack):
            if commit.commit_id == commit_id:
                target_index = i
                break

        if target_index is None:
            raise ValueError(f"Commit {commit_id} not found")

        # Apply changes to reach target state
        checkout_info = {
            "target_commit": commit_id,
            "applied_commits": []
        }

        if target_index < self.current_commit_index:
            # Rollback
            for i in range(self.current_commit_index, target_index, -1):
                commit = self.history_stack[i]
                self._apply_reverse_changes(commit)
                checkout_info["applied_commits"].append(commit.commit_id)
        elif target_index > self.current_commit_index:
            # Forward
            for i in range(self.current_commit_index + 1, target_index + 1):
                commit = self.history_stack[i]
                self._apply_forward_changes(commit)
                checkout_info["applied_commits"].append(commit.commit_id)

        self.current_commit_index = target_index
        logger.info(f"Checked out to commit {commit_id}")
        return checkout_info

    def _apply_reverse_changes(self, commit: Commit):
        """Apply reverse changes from a commit."""
        changes = commit.changes
        before_state = changes.get("before", {})

        for element_id in commit.affected_elements:
            element = self.get_object(element_id)
            if not element:
                logger.warning(f"Element {element_id} not found, skipping")
                continue

            # Restore based on element type
            if isinstance(element, Paragraph):
                if "text" in before_state:
                    element.clear()
                    element.add_run(before_state["text"])
            elif isinstance(element, Run):
                if "text" in before_state:
                    element.text = before_state["text"]
                if "bold" in before_state:
                    element.bold = before_state["bold"]
                if "italic" in before_state:
                    element.italic = before_state["italic"]
            elif isinstance(element, Table):
                if "cells" in before_state:
                    self._restore_table_cells(element, before_state["cells"])

    def _apply_forward_changes(self, commit: Commit):
        """Apply forward changes from a commit."""
        changes = commit.changes
        after_state = changes.get("after", {})

        for element_id in commit.affected_elements:
            element = self.get_object(element_id)
            if not element:
                logger.warning(f"Element {element_id} not found, skipping")
                continue

            # Apply based on element type
            if isinstance(element, Paragraph):
                if "text" in after_state:
                    element.clear()
                    element.add_run(after_state["text"])
            elif isinstance(element, Run):
                if "text" in after_state:
                    element.text = after_state["text"]
                if "bold" in after_state:
                    element.bold = after_state["bold"]
                if "italic" in after_state:
                    element.italic = after_state["italic"]
            elif isinstance(element, Table):
                if "cells" in after_state:
                    self._restore_table_cells(element, after_state["cells"])

    def _restore_table_cells(self, table: Table, cells_data: List[Dict[str, Any]]):
        """Restore table cells from saved data."""
        for cell_data in cells_data:
            row = cell_data.get("row")
            col = cell_data.get("col")
            text = cell_data.get("text", "")
            if row is not None and col is not None:
                try:
                    cell = table.rows[row].cells[col]
                    cell.text = text
                except IndexError:
                    logger.warning(f"Cell ({row}, {col}) out of range")

    def _format_element_summary(self, element: Any) -> str:
        """Format element summary with truncation."""
        if isinstance(element, Paragraph):
            text = element.text.replace('\n', ' ')
            return f'"{text[:50]}{"..." if len(text) > 50 else ""}"'
        elif isinstance(element, Run):
            text = element.text.replace('\n', ' ')
            return f'Run: "{text[:50]}{"..." if len(text) > 50 else ""}"'
        elif isinstance(element, Table):
            rows = len(element.rows)
            cols = len(element.columns) if element.rows else 0
            return f"Table ({rows}x{cols})"
        return str(type(element).__name__)

    def get_cursor_context(self, num_before: int = 2, num_after: int = 2) -> str:
        """Generate cursor position context description."""
        try:
            # Handle empty document
            if not self.cursor.element_id:
                return "Cursor: at empty document start"

            current_element = self.get_object(self.cursor.element_id)
            if not current_element:
                return f"Cursor: after {self.cursor.element_id} (element not found)"

            # Get parent container
            if self.cursor.parent_id and self.cursor.parent_id != "document_body":
                parent = self.get_object(self.cursor.parent_id)
            else:
                parent = self.document

            # Get siblings
            siblings = self._get_siblings(parent)
            if not siblings:
                return f"Cursor: after {self.cursor.element_id}"

            # Find current index
            current_idx = None
            for i, sib in enumerate(siblings):
                if hasattr(sib, '_element') and hasattr(current_element, '_element'):
                    if sib._element is current_element._element:
                        current_idx = i
                        break

            if current_idx is None:
                return f"Cursor: after {self.cursor.element_id}"

            # Build context
            lines = [f"Cursor: after {type(current_element).__name__} {self.cursor.element_id}"]

            if self.cursor.parent_id and self.cursor.parent_id != "document_body":
                parent_obj = self.get_object(self.cursor.parent_id)
                if parent_obj:
                    lines.append(f"Parent: {type(parent_obj).__name__} {self.cursor.parent_id}")

            lines.append("Context:")

            # Before elements
            start = max(0, current_idx - num_before)
            for i in range(start, current_idx):
                elem = siblings[i]
                elem_id = self._get_element_id(elem, auto_register=True)
                summary = self._format_element_summary(elem)
                lines.append(f"  [{i - current_idx}] {type(elem).__name__} {elem_id}: {summary}")

            # Current element
            summary = self._format_element_summary(current_element)
            lines.append(f"  [Current] {type(current_element).__name__} {self.cursor.element_id}: {summary}")

            # After elements
            end = min(len(siblings), current_idx + 1 + num_after)
            for i in range(current_idx + 1, end):
                elem = siblings[i]
                elem_id = self._get_element_id(elem, auto_register=True)
                summary = self._format_element_summary(elem)
                lines.append(f"  [+{i - current_idx}] {type(elem).__name__} {elem_id}: {summary}")

            if current_idx + 1 >= len(siblings):
                lines.append("  [Document End]")

            return "\n".join(lines)

        except Exception as e:
            logger.warning(f"Failed to generate cursor context: {e}")
            return f"Cursor: after {self.cursor.element_id if self.cursor.element_id else 'unknown'}"

class SessionManager:
    def __init__(self, ttl_seconds: int = 3600):
        self.sessions: Dict[str, Session] = {}
        self.ttl_seconds = ttl_seconds

    def create_session(
        self,
        file_path: Optional[str] = None,
        auto_save: bool = False,
        backup_on_save: bool = False,
        backup_dir: Optional[str] = None,
        backup_suffix: Optional[str] = None,
    ) -> str:
        """Create a new session, optionally loading a file."""
        # Shortened session id for usability while retaining randomness
        session_id = uuid.uuid4().hex[:12]

        if file_path:
            # 1. Validate cross-OS path safety
            validate_path_safety(file_path)

            if os.path.exists(file_path):
                try:
                    doc = Document(file_path)
                    logger.info(f"Session created: {session_id}, loaded file: {file_path}, auto_save={auto_save}")
                except Exception as e:
                    # If file exists but fails to load (e.g. locked, corrupt), raise error
                    # so user knows why instead of silently returning empty doc
                    logger.exception(f"Failed to load file '{file_path}': {e}")
                    raise RuntimeError(f"Failed to load existing file '{file_path}': {str(e)}")
            else:
                # File doesn't exist.
                # Validate that the parent directory exists so we can save later.
                # We use abspath to ensure we check the correct directory even for relative paths.
                parent_dir = os.path.dirname(os.path.abspath(file_path))
                if not os.path.exists(parent_dir):
                    logger.error(f"Parent directory does not exist: {parent_dir}")
                    raise ValueError(f"Parent directory does not exist: {parent_dir}")

                # Create new empty doc (intended for new file creation)
                doc = Document()
                logger.info(f"Session created: {session_id}, new file: {file_path}, auto_save={auto_save}")
        else:
            doc = Document()
            logger.info(f"Session created: {session_id}, in-memory document")

        session = Session(
            session_id=session_id,
            document=doc,
            file_path=file_path,
            auto_save=auto_save,
            backup_on_save=backup_on_save,
            backup_dir=backup_dir,
            backup_suffix=backup_suffix,
        )
        self.sessions[session_id] = session
        return session_id

    def get_session(self, session_id: str) -> Optional[Session]:
        session = self.sessions.get(session_id)
        if not session:
            logger.debug(f"Session not found: {session_id}")
            return None

        # Check expiry
        if time.time() - session.last_accessed > self.ttl_seconds:
            logger.warning(f"Session expired: {session_id}")
            del self.sessions[session_id]
            return None

        session.touch()
        return session

    def close_session(self, session_id: str) -> bool:
        if session_id in self.sessions:
            del self.sessions[session_id]
            logger.info(f"Session closed: {session_id}")
            return True
        logger.debug(f"Close session failed - not found: {session_id}")
        return False

    def cleanup_expired(self, max_idle_seconds: Optional[int] = None) -> int:
        now = time.time()
        ttl = max_idle_seconds if max_idle_seconds is not None else self.ttl_seconds
        expired = [
            sid for sid, s in self.sessions.items()
            if now - s.last_accessed > ttl
        ]
        for sid in expired:
            del self.sessions[sid]
        if expired:
            logger.info(f"Cleaned up {len(expired)} expired sessions (ttl={ttl}s)")
        return len(expired)

    def list_sessions(self) -> List[Dict[str, Any]]:
        now = time.time()
        items: List[Dict[str, Any]] = []
        for sid, s in self.sessions.items():
            items.append({
                "session_id": sid,
                "file_path": s.file_path,
                "auto_save": s.auto_save,
                "backup_on_save": s.backup_on_save,
                "last_accessed": s.last_accessed,
                "age_seconds": now - s.created_at,
                "idle_seconds": now - s.last_accessed,
            })
        return items
