"""Advanced document operations"""
import os
import json
import logging
from mcp.server.fastmcp import FastMCP
from docx.shared import Inches
from docx_mcp_server.core.replacer import replace_text_in_paragraph
from docx_mcp_server.utils.text_tools import TextTools
from docx_mcp_server.core.response import (
    create_markdown_response,
    create_error_response
)

from docx_mcp_server.services.navigation import PositionResolver, ContextBuilder
from docx_mcp_server.core.xml_util import ElementManipulator

logger = logging.getLogger(__name__)


def docx_replace_text(session_id: str, old_text: str, new_text: str, scope_id: str = None) -> str:
    """
    Replace all occurrences of text in the document or a specific scope.

    Searches and replaces text while preserving formatting. Can operate on the entire
    document or be limited to a specific element (paragraph, table, etc.).

    Typical Use Cases:
        - Replace placeholders in templates
        - Update repeated text globally
        - Batch text modifications

    Args:
        session_id (str): Active session ID returned by docx_create().
        old_text (str): Text to find and replace (exact match).
        new_text (str): Replacement text.
        scope_id (str, optional): ID of element to limit scope (paragraph, table, cell).
            If None, searches entire document body.

    Returns:
        str: JSON response with count of replacements made.

    Raises:
        ValueError: If session_id or scope_id is invalid.

    Examples:
        Replace globally:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> result = docx_replace_text(session_id, "{{COMPANY}}", "Acme Corp")
        >>> print(result)
        'Replaced 5 occurrences of {{COMPANY}}'

        Replace in specific paragraph:
        >>> para_id = docx_find_paragraphs(session_id, "{{NAME}}")[0]["id"]
        >>> docx_replace_text(session_id, "{{NAME}}", "John", scope_id=para_id)

        Replace in table:
        >>> table_id = docx_find_table(session_id, "Invoice")
        >>> docx_replace_text(session_id, "TBD", "Completed", scope_id=table_id)

    Notes:
        - Preserves text formatting (bold, italic, etc.)
        - Exact text match (case-sensitive)
        - Searches document body, tables, and cells
        - Does not search headers/footers

    See Also:
        - docx_find_paragraphs: Find specific paragraphs
        - docx_update_paragraph_text: Replace entire paragraph
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_replace_text called: session_id={session_id}, old_text='{old_text}', new_text='{new_text}', scope_id={scope_id}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_replace_text failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    # Use shared scope resolution logic
    tools = TextTools()
    if scope_id:
        obj = session.get_object(scope_id)
        if not obj:
            logger.error(f"docx_replace_text failed: Scope object {scope_id} not found")
            return create_error_response(f"Scope object {scope_id} not found", error_type="ElementNotFound")
        targets = tools.collect_paragraphs_from_scope(obj)
    else:
        targets = tools.collect_paragraphs_from_scope(None, session.document)

    count = 0
    affected_paragraphs = []
    for p in targets:
        if replace_text_in_paragraph(p, old_text, new_text):
            count += 1
            # Track affected paragraphs for context
            p_id = session._get_element_id(p, auto_register=True)
            affected_paragraphs.append(p_id)

    logger.debug(f"docx_replace_text success: replaced {count} occurrences")

    # For batch operations, show context of first affected paragraph if any
    element_id = affected_paragraphs[0] if affected_paragraphs else None

    return create_markdown_response(
        session=session,
        message=f"Replaced {count} occurrences of '{old_text}'",
        element_id=element_id,
        operation="Replace Text",
        show_context=True,
        show_diff=True,
        old_content=old_text,
        new_content=new_text,
        replacements=count,
        affected_paragraphs=affected_paragraphs,
        scope_id=scope_id,
        old_text=old_text,
        new_text=new_text
    )

def docx_batch_replace_text(session_id: str, replacements_json: str, scope_id: str = None) -> str:
    """
    Perform batch text replacement across the document or a specific scope.

    Efficiently replaces multiple text patterns in a single pass. Strictly preserves
    formatting by only replacing text within individual runs. Does NOT match text
    that spans across multiple runs (e.g., if half a word is bold and half is not).

    Typical Use Cases:
        - Bulk fill templates ({NAME} -> "John", {DATE} -> "2023")
        - Anonymize documents (replace names with Redacted)
        - Update terminology globally

    Args:
        session_id (str): Active session ID.
        replacements_json (str): JSON object mapping old text to new text.
            Example: '{"{{NAME}}": "John Doe", "{{DATE}}": "2023-01-01"}'
        scope_id (str, optional): ID of element to limit scope (paragraph, table).
            If None, applies to entire document body.

    Returns:
        str: JSON response with total replacement count.

    Raises:
        ValueError: If JSON is invalid or session not found.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_batch_replace_text called: session_id={session_id}, replacements_len={len(replacements_json)}, scope_id={scope_id}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_batch_replace_text failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        replacements = json.loads(replacements_json)
    except json.JSONDecodeError as e:
        logger.error(f"docx_batch_replace_text failed: Invalid JSON - {e}")
        return create_error_response("Invalid JSON for replacements", error_type="ValidationError")

    # Use shared scope resolution logic
    tools = TextTools()
    if scope_id:
        obj = session.get_object(scope_id)
        if not obj:
            logger.error(f"docx_batch_replace_text failed: Scope object {scope_id} not found")
            return create_error_response(f"Scope object {scope_id} not found", error_type="ElementNotFound")
        targets = tools.collect_paragraphs_from_scope(obj)
    else:
        targets = tools.collect_paragraphs_from_scope(None, session.document)

    count = tools.batch_replace_text(targets, replacements)

    logger.debug(f"docx_batch_replace_text success: replaced {count} occurrences")

    # For batch operations, show first target paragraph if available
    element_id = None
    if targets:
        element_id = session._get_element_id(targets[0], auto_register=True)

    # Create a summary of replacements for diff display
    replacements_summary = "\n".join([f"{old} → {new}" for old, new in replacements.items()])

    return create_markdown_response(
        session=session,
        message=f"Batch replacement completed. Replaced {count} occurrences.",
        element_id=element_id,
        operation="Batch Replace Text",
        show_context=True,
        show_diff=True,
        old_content=f"Multiple patterns ({len(replacements)} patterns)",
        new_content=replacements_summary,
        replacements=count,
        patterns=len(replacements),
        scope_id=scope_id
    )

def docx_insert_image(session_id: str, image_path: str, position: str, width: float = None, height: float = None) -> str:
    """
    Insert an image into the document.

    Adds an image file to the document with optional size constraints. Images can be
    added to the document body or within specific paragraphs.

    Typical Use Cases:
        - Add logos or diagrams to documents
        - Insert charts or screenshots
        - Embed visual content in reports

    Args:
        session_id (str): Active session ID returned by docx_create().
        image_path (str): Absolute or relative path to the image file (PNG, JPG, etc.).
        position (str): Insertion position string (e.g., "after:para_123").
            If used, creates a new paragraph for the image at that location.
        width (float, optional): Image width in inches. If None, uses original size.
        height (float, optional): Image height in inches. If None, uses original size.

    Returns:
        str: JSON response with element ID of the container paragraph or run.

    Raises:
        ValueError: If session_id is invalid or image_path not found.
        FileNotFoundError: If image file does not exist.

    Examples:
        Insert image with default size:
        >>> session_id = docx_create()
        >>> para_id = docx_insert_image(session_id, "./logo.png", position="end:document_body")

        Insert with specific dimensions:
        >>> para_id = docx_insert_image(session_id, "./chart.png", width=4.0, height=3.0, position="end:document_body")

        Insert into existing paragraph:
        >>> para_id = docx_insert_paragraph(session_id, "See image: ", position="end:document_body")
        >>> docx_insert_image(session_id, "./diagram.png", width=2.0, position=f"inside:{para_id}")

        Insert at specific position:
        >>> docx_insert_image(session_id, "./logo.png", position="start:document_body")

    Notes:
        - Supported formats: PNG, JPG, GIF, BMP
        - Size is in inches (1 inch = 2.54 cm)
        - If only width or height specified, aspect ratio is preserved
        - Image is embedded in document (increases file size)

    See Also:
        - docx_insert_paragraph: Create container paragraph
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_image called: session_id={session_id}, image_path={image_path}, width={width}, height={height}, position={position}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_insert_image failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    if not os.path.exists(image_path):
        logger.error(f"docx_insert_image failed: Image file not found - {image_path}")
        return create_error_response(f"Image file not found: {image_path}", error_type="FileNotFound")

    # Resolve Target Parent and Position
    target_parent = session.document
    ref_element = None
    mode = "append"

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    # Case 1: Target is a Paragraph -> Add Run with Picture
    if hasattr(target_parent, "add_run"):
        try:
            run = target_parent.add_run()
            run.add_picture(image_path, width=Inches(width) if width else None, height=Inches(height) if height else None)

            if mode != "append":
                if mode == "before" and ref_element:
                    ElementManipulator.insert_xml_before(ref_element._element, run._element)
                elif mode == "after" and ref_element:
                    ElementManipulator.insert_xml_after(ref_element._element, run._element)
                elif mode == "start":
                    ElementManipulator.insert_at_index(target_parent._element, run._element, 0)

            r_id = session.register_object(run, "run")
            session.update_context(r_id, action="create")

            session.cursor.element_id = r_id
            session.cursor.position = "after"

            # Context for Run
            builder = ContextBuilder(session)
            data = builder.build_response_data(run, r_id)

            return create_markdown_response(
            session=session,
            message="Image inserted into paragraph successfully",
            operation="Operation",
            show_context=True,
                image_path=image_path,
                **data
            
        )
        except Exception as e:
            logger.exception(f"docx_insert_image failed (add run): {e}")
            return create_error_response(f"Failed to add image to paragraph: {str(e)}", error_type="CreationError")

    # Case 2: Target is a Document/Body -> Create New Paragraph with Picture
    elif hasattr(target_parent, "add_paragraph"):
        try:
            # Create paragraph
            paragraph = target_parent.add_paragraph()
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(width) if width else None, height=Inches(height) if height else None)

            # Move if necessary
            if mode != "append":
                if mode == "before" and ref_element:
                    ElementManipulator.insert_xml_before(ref_element._element, paragraph._element)
                elif mode == "after" and ref_element:
                    ElementManipulator.insert_xml_after(ref_element._element, paragraph._element)
                elif mode == "start":
                    container_xml = target_parent._element
                    if hasattr(target_parent, '_body'):
                        container_xml = target_parent._body._element
                    ElementManipulator.insert_at_index(container_xml, paragraph._element, 0)

            # Register Paragraph
            p_id = session.register_object(paragraph, "para")
            session.update_context(p_id, action="create")

            session.cursor.element_id = p_id
            session.cursor.position = "after"

            # Context
            builder = ContextBuilder(session)
            data = builder.build_response_data(paragraph, p_id)

            return create_markdown_response(
            session=session,
            message="Image inserted successfully",
            operation="Operation",
            show_context=True,
                image_path=image_path,
                **data
            
        )
        except Exception as e:
            logger.exception(f"docx_insert_image failed (new para): {e}")
            return create_error_response(f"Failed to insert image: {str(e)}", error_type="CreationError")

    else:
        return create_error_response(f"Target object {type(target_parent).__name__} cannot contain images", error_type="InvalidParent")



def register_tools(mcp: FastMCP):
    """Register advanced document operations"""
    mcp.tool()(docx_replace_text)
    mcp.tool()(docx_batch_replace_text)
    mcp.tool()(docx_insert_image)
