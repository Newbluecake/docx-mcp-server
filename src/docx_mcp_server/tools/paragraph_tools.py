"""Paragraph manipulation tools"""
import logging
from mcp.server.fastmcp import FastMCP
from docx_mcp_server.utils.metadata_tools import MetadataTools
from docx_mcp_server.core.response import (
    create_markdown_response,
    create_error_response,
    create_context_aware_response
)
from docx_mcp_server.services.navigation import PositionResolver, ContextBuilder
from docx_mcp_server.core.xml_util import ElementManipulator

logger = logging.getLogger(__name__)


def docx_insert_paragraph(session_id: str, text: str, position: str, style: str = None) -> str:
    """
    Add a new paragraph to the document with precise positioning control.

    Creates a paragraph with the specified text content at any location in the document.
    Supports positioning relative to existing elements or at document start/end.

    **Session Context**: This tool operates on the document associated with
    session_id. Create a session first using docx_create(), which uses the
    global active file set by the Launcher GUI or --file parameter.

    Typical Use Cases:
        - Add body text to a document
        - Insert paragraphs at specific locations (after/before existing elements)
        - Add content to table cells
        - Build document structure programmatically

    Args:
        session_id (str): Active session ID returned by docx_create().
            The session maintains document state and object registry.
        text (str): Text content for the paragraph. Can be empty string for
            formatting-only paragraphs.
        position (str): Insertion position string.
            Format: "mode:target_id". Modes:
            - "after:element_id" - Insert after specified element
            - "before:element_id" - Insert before specified element
            - "inside:element_id" - Insert inside container (cell, etc.)
            - "start:document_body" - Insert at document start
            - "end:document_body" - Insert at document end
            Example: "after:para_123", "end:document_body"
        style (str, optional): Built-in style name (e.g., 'List Bullet', 'Body Text',
            'Heading 1'). Defaults to None (Normal style).

    Returns:
        str: Markdown-formatted response containing:
            - **Element ID**: Unique identifier for the created paragraph (e.g., "para_abc123")
            - **Status**: Success/Error indicator
            - **Document Context**: ASCII visualization showing paragraph position
            - **Cursor**: Updated cursor position after operation

    Raises:
        SessionNotFound: If session_id is invalid or session has expired.
        ValidationError: If position format is invalid.
        InvalidParent: If target cannot contain paragraphs.
        CreationError: If paragraph creation fails.

    Examples:
        Basic usage - add paragraph at document end:
        >>> # Step 1: Create session for active file
        >>> session_id = docx_create()
        >>> # Step 2: Insert paragraph
        >>> result = docx_insert_paragraph(session_id, "Hello World", position="end:document_body")
        >>> # Step 3: Extract element_id
        >>> import re
        >>> match = re.search(r'\*\*Element ID\*\*:\s*(\w+)', result)
        >>> para_id = match.group(1) if match else None

        Insert after existing paragraph:
        >>> result = docx_insert_paragraph(session_id, "New text", position=f"after:{para_id}")

        Insert with custom style:
        >>> result = docx_insert_paragraph(
        ...     session_id, "Bullet point",
        ...     position="end:document_body",
        ...     style="List Bullet"
        ... )

        Complete workflow:
        >>> # Set active file via Launcher or --file parameter
        >>> session_id = docx_create()
        >>> para1 = docx_insert_paragraph(session_id, "First paragraph", position="end:document_body")
        >>> para2 = docx_insert_paragraph(session_id, "Second paragraph", position="end:document_body")
        >>> docx_save(session_id, "./output.docx")
        >>> docx_close(session_id)

    Notes:
        - Empty text creates a blank paragraph (useful for spacing)
        - Position is required - no default position
        - Cursor automatically moves to the new paragraph
        - Use docx_insert_run() to add formatted text within the paragraph
        - For headings, use docx_insert_heading() instead

    See Also:
        - docx_create: Create session for active file
        - docx_insert_heading: Insert heading paragraphs
        - docx_insert_run: Add formatted text to paragraph
        - docx_update_paragraph_text: Modify existing paragraph text
        - docx_set_alignment: Set paragraph alignment
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_paragraph called: session_id={session_id}, position={position}")
    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_insert_paragraph failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    # Resolve Target Parent and Position
    target_parent = None
    ref_element = None
    mode = "append"

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    if not hasattr(target_parent, 'add_paragraph'):
         return create_error_response(f"Object {type(target_parent).__name__} cannot contain paragraphs", error_type="InvalidParent")

    try:
        # Create Paragraph (always appended first by python-docx)
        paragraph = target_parent.add_paragraph(text, style=style)

        # Move if necessary
        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, paragraph._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, paragraph._element)
            elif mode == "start":
                 # Determine correct container element
                 container_xml = target_parent._element
                 # If target is Document, we want to insert into its body
                 if hasattr(target_parent, '_body'):
                     container_xml = target_parent._body._element

                 ElementManipulator.insert_at_index(container_xml, paragraph._element, 0)
            # "inside" usually implies append (handled by default) or start (handled above)

        # Register
        p_id = session.register_object(paragraph, "para")

        # Update context
        session.update_context(p_id, action="create")

        # Update session cursor state for consistency
        session.cursor.element_id = p_id
        session.cursor.position = "after"

        return create_markdown_response(
            session=session,
            message="Paragraph created successfully",
            element_id=p_id,
            operation="Insert Paragraph",
            show_context=True,
            position=position,
            style=style or "Normal"
        )

    except Exception as e:
        logger.exception(f"docx_insert_paragraph failed: {e}")
        return create_error_response(f"Failed to create paragraph: {str(e)}", error_type="CreationError")

def docx_insert_heading(session_id: str, text: str, position: str, level: int = 1) -> str:
    """
    Add a heading paragraph to the document with specified level.

    Creates a heading with the specified level (0=Title, 1-9=Heading 1-9).
    Headings are styled paragraphs that provide document structure.

    **Session Context**: This tool operates on the document associated with
    session_id. Create a session first using docx_create(), which uses the
    global active file set by the Launcher GUI or --file parameter.

    Typical Use Cases:
        - Create document structure with hierarchical headings
        - Add section titles and subtitles
        - Build table of contents structure
        - Organize long documents

    Args:
        session_id (str): Active session ID returned by docx_create().
            The session maintains document state and object registry.
        text (str): Heading text content.
        position (str): Insertion position string.
            Format: "mode:target_id". Examples:
            - "end:document_body" - Add at document end
            - "after:para_123" - Insert after specified element
            - "before:para_456" - Insert before specified element
        level (int, optional): Heading level (0-9).
            - 0: Title style
            - 1-9: Heading 1 through Heading 9
            Defaults to 1 (Heading 1).

    Returns:
        str: Markdown-formatted response containing:
            - **Element ID**: Unique identifier for the heading (e.g., "para_abc123")
            - **Status**: Success/Error indicator
            - **Document Context**: ASCII visualization showing heading position
            - **Level**: Heading level applied

    Raises:
        SessionNotFound: If session_id is invalid or session has expired.
        ValidationError: If position format is invalid or level out of range.
        InvalidParent: If target cannot contain headings.
        CreationError: If heading creation fails.

    Examples:
        Basic usage - add heading at document end:
        >>> # Step 1: Create session
        >>> session_id = docx_create()
        >>> # Step 2: Insert heading
        >>> result = docx_insert_heading(session_id, "Chapter 1", position="end:document_body", level=1)
        >>> # Step 3: Extract element_id
        >>> import re
        >>> match = re.search(r'\*\*Element ID\*\*:\s*(\w+)', result)
        >>> heading_id = match.group(1) if match else None

        Create document structure:
        >>> session_id = docx_create()
        >>> h1 = docx_insert_heading(session_id, "Introduction", position="end:document_body", level=1)
        >>> p1 = docx_insert_paragraph(session_id, "Content...", position="end:document_body")
        >>> h2 = docx_insert_heading(session_id, "Background", position="end:document_body", level=2)
        >>> docx_save(session_id, "./document.docx")
        >>> docx_close(session_id)

        Use Title style (level 0):
        >>> result = docx_insert_heading(session_id, "Document Title", position="start:document_body", level=0)

    Notes:
        - Headings are special paragraphs with built-in styles
        - Level 0 uses "Title" style, levels 1-9 use "Heading X" styles
        - Headings automatically appear in Word's navigation pane
        - Use consistent heading levels for proper document structure
        - Cursor automatically moves to the new heading

    See Also:
        - docx_create: Create session for active file
        - docx_insert_paragraph: Insert regular paragraphs
        - docx_set_alignment: Modify heading alignment
        - docx_update_paragraph_text: Change heading text
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_heading called: session_id={session_id}, level={level}, position={position}")
    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    # Resolve Position
    # Headings are usually added to Body, but can be in cells too (though rare for main structure)
    # We assume default parent is Document Body
    target_parent = session.document
    ref_element = None
    mode = "append"

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    try:
        # Create Heading
        # Note: _Body objects (returned by resolver for main document content) do not have add_heading,
        # but they do have add_paragraph. Document objects have both.
        if hasattr(target_parent, 'add_heading'):
            heading = target_parent.add_heading(text, level=level)
        elif hasattr(target_parent, 'add_paragraph'):
            # Fallback: create paragraph with heading style
            # Level 0 is 'Title', 1-9 is 'Heading X'
            style = "Title" if level == 0 else f"Heading {level}"
            heading = target_parent.add_paragraph(text, style=style)
        else:
            return create_error_response(f"Object {type(target_parent).__name__} cannot contain headings", error_type="InvalidParent")

        # Move if necessary
        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, heading._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, heading._element)
            elif mode == "start":
                 ElementManipulator.insert_at_index(target_parent._element, heading._element, 0)

        h_id = session.register_object(heading, "para") # Headings are paragraphs

        session.update_context(h_id, action="create")

        return create_markdown_response(
            session=session,
            message=f"Heading level {level} created successfully",
            element_id=h_id,
            operation="Insert Heading",
            show_context=True,
            level=level,
            position=position
        )
    except Exception as e:
        logger.exception(f"docx_insert_heading failed: {e}")
        return create_error_response(f"Failed to create heading: {str(e)}", error_type="CreationError")

def docx_update_paragraph_text(session_id: str, paragraph_id: str, new_text: str) -> str:
    """
    Replace all text content in an existing paragraph.

    Clears all existing runs in the paragraph and replaces with a single run
    containing the new text. All previous formatting within the paragraph is lost.

    Args:
        session_id (str): Active session ID returned by docx_create().
        paragraph_id (str): ID of the paragraph to update.
        new_text (str): New text content to replace existing content.

    Returns:
        str: JSON response with success message and cursor context.
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        paragraph = session.get_object(paragraph_id)
    except ValueError as e:
        # Special ID resolution failed
        if "Special ID" in str(e) or "not available" in str(e):
            return create_error_response(str(e), error_type="SpecialIDNotAvailable")
        raise

    if not paragraph:
        return create_error_response(f"Paragraph {paragraph_id} not found", error_type="ElementNotFound")

    if not hasattr(paragraph, 'text'):
        return create_error_response(f"Object {paragraph_id} is not a paragraph", error_type="InvalidElementType")

    try:
        # Capture old content for diff
        old_text = paragraph.text

        # Clear existing runs and set new text
        paragraph.clear()
        paragraph.add_run(new_text)

        # Get the actual element ID (resolve special IDs to concrete IDs)
        actual_para_id = session._get_element_id(paragraph, auto_register=False)
        if not actual_para_id:
            # Fallback: try to find in registry
            for eid, obj in session.object_registry.items():
                if obj is paragraph:
                    actual_para_id = eid
                    break

        if actual_para_id:
            # Update context to track this as an update operation
            session.update_context(actual_para_id, action="update")

            # Update cursor to point to this paragraph
            session.cursor.element_id = actual_para_id
            session.cursor.position = "after"

            return create_markdown_response(
                session=session,
                message=f"Paragraph updated successfully",
                element_id=actual_para_id,
                operation="Update Paragraph Text",
                show_context=True,
                show_diff=True,
                old_content=old_text,
                new_content=new_text,
                changed_fields=["text"]
            )
        else:
            # Shouldn't happen, but handle gracefully
            return create_markdown_response(
                session=session,
                message=f"Paragraph updated successfully",
                element_id=paragraph_id,
                operation="Update Paragraph Text",
                show_context=False,
                show_diff=True,
                old_content=old_text,
                new_content=new_text,
                changed_fields=["text"]
            )
    except Exception as e:
        logger.exception(f"docx_update_paragraph_text failed: {e}")
        return create_error_response(f"Failed to update paragraph: {str(e)}", error_type="UpdateError")

def docx_copy_paragraph(session_id: str, paragraph_id: str, position: str) -> str:
    """
    Create a deep copy of a paragraph with all formatting and runs.

    Args:
        session_id (str): Active session ID.
        paragraph_id (str): ID of the paragraph to copy.
        position (str): Insertion position string (e.g., "after:para_123").
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        source_para = session.get_object(paragraph_id)
    except ValueError as e:
        # Special ID resolution failed
        if "Special ID" in str(e) or "not available" in str(e):
            return create_error_response(str(e), error_type="SpecialIDNotAvailable")
        raise

    if not source_para:
        return create_error_response(f"Paragraph {paragraph_id} not found", error_type="ElementNotFound")

    if not hasattr(source_para, 'runs'):
        return create_error_response(f"Object {paragraph_id} is not a paragraph", error_type="InvalidElementType")

    try:
        # Resolve target position
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)

        if not hasattr(target_parent, 'add_paragraph'):
            return create_error_response(f"Object {type(target_parent).__name__} cannot contain paragraphs", error_type="InvalidParent")

        # Create new paragraph with same style
        new_para = target_parent.add_paragraph(style=source_para.style)

        # Copy paragraph-level formatting
        if source_para.alignment is not None:
            new_para.alignment = source_para.alignment

        # Copy all runs with their formatting
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)

            # Copy font properties
            if run.font.bold is not None:
                new_run.font.bold = run.font.bold
            if run.font.italic is not None:
                new_run.font.italic = run.font.italic
            if run.font.underline is not None:
                new_run.font.underline = run.font.underline
            if run.font.size is not None:
                new_run.font.size = run.font.size
            if run.font.color.rgb is not None:
                new_run.font.color.rgb = run.font.color.rgb

        # Track lineage metadata using shared utility
        meta = MetadataTools.create_copy_metadata(
            source_id=paragraph_id,
            source_type="paragraph"
        )

        # Move if necessary
        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, new_para._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, new_para._element)
            elif mode == "start":
                container_xml = target_parent._element
                if hasattr(target_parent, '_body'):
                    container_xml = target_parent._body._element
                ElementManipulator.insert_at_index(container_xml, new_para._element, 0)

        new_para_id = session.register_object(new_para, "para", metadata=meta)

        # Update cursor to point after the new paragraph
        session.cursor.element_id = new_para_id
        session.cursor.parent_id = "document_body"
        session.cursor.position = "after"

        # Use Builder
        builder = ContextBuilder(session)
        data = builder.build_response_data(new_para, new_para_id)

        return create_markdown_response(
            session=session,
            message="Paragraph copied successfully",
            operation="Operation",
            show_context=True,
            source_id=paragraph_id,
            **data
        
        )
    except Exception as e:
        logger.exception(f"docx_copy_paragraph failed: {e}")
        return create_error_response(f"Failed to copy paragraph: {str(e)}", error_type="CopyError")

def docx_delete(session_id: str, element_id: str = None) -> str:
    """
    Delete an element from the document.
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    if not element_id:
        element_id = session.last_accessed_id

    if not element_id:
        return create_error_response("No element specified and no context available", error_type="NoContext")

    try:
        obj = session.get_object(element_id)
    except ValueError as e:
        # Special ID resolution failed
        if "Special ID" in str(e) or "not available" in str(e):
            return create_error_response(str(e), error_type="SpecialIDNotAvailable")
        raise

    if not obj:
        return create_error_response(f"Object {element_id} not found", error_type="ElementNotFound")

    # Try to delete
    try:
        if hasattr(obj, "_element") and obj._element.getparent() is not None:
            parent = obj._element.getparent()
            parent.remove(obj._element)

            # Remove from registry
            if element_id in session.object_registry:
                del session.object_registry[element_id]

            # Reset context if it pointed here
            if session.last_accessed_id == element_id:
                session.last_accessed_id = None
            if session.last_created_id == element_id:
                session.last_created_id = None

            # Update cursor to parent container since element is gone
            session.cursor.element_id = None
            session.cursor.position = "inside_end"

            # For deletion, we can't show the element anymore.
            # Maybe show the parent's context?
            # Let's keep it simple for now, deletion response doesn't strictly need the visual tree of the deleted item.

            return create_context_aware_response(
                session,
                message=f"Deleted {element_id}",
                deleted_id=element_id
            )
        else:
            return create_error_response("Element has no parent or cannot be deleted", error_type="DeletionError")
    except Exception as e:
        logger.exception(f"docx_delete failed: {e}")
        return create_error_response(f"Failed to delete {element_id}: {str(e)}", error_type="DeletionError")

def docx_insert_page_break(session_id: str, position: str) -> str:
    """
    Insert a page break at the specified position in the document.

    Args:
        session_id (str): Active session ID.
        position (str): Insertion position string (e.g., "after:para_123").
    """
    from docx_mcp_server.server import session_manager
    from docx.enum.text import WD_BREAK

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)

        if not hasattr(target_parent, 'add_paragraph'):
            return create_error_response(f"Object {type(target_parent).__name__} cannot contain paragraphs", error_type="InvalidParent")

        paragraph = target_parent.add_paragraph()
        run = paragraph.add_run()
        run.add_break(WD_BREAK.PAGE)

        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, paragraph._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, paragraph._element)
            elif mode == "start":
                container_xml = target_parent._element
                if hasattr(target_parent, '_body'):
                    container_xml = target_parent._body._element
                ElementManipulator.insert_at_index(container_xml, paragraph._element, 0)

        p_id = session.register_object(paragraph, "para")
        session.update_context(p_id, action="create")

        return create_markdown_response(
            session=session,
            message="Page break inserted",
            operation="Operation",
            show_context=True,
            element_id=p_id
        
        )
    except Exception as e:
        logger.exception(f"docx_insert_page_break failed: {e}")
        return create_error_response(f"Failed to insert page break: {str(e)}", error_type="CreationError")


def register_tools(mcp: FastMCP):
    """Register paragraph manipulation tools"""
    mcp.tool()(docx_insert_paragraph)
    mcp.tool()(docx_insert_heading)
    mcp.tool()(docx_update_paragraph_text)
    mcp.tool()(docx_copy_paragraph)
    mcp.tool()(docx_delete)
    mcp.tool()(docx_insert_page_break)
