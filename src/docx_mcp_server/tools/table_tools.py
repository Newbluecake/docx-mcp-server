"""Table manipulation tools"""
import json
import logging
from mcp.server.fastmcp import FastMCP
from docx.shared import Inches
from docx.table import _Cell, Table
from docx_mcp_server.core.finder import Finder
from docx_mcp_server.utils.copy_engine import CopyEngine
from docx_mcp_server.core.format_painter import FormatPainter
from docx_mcp_server.utils.metadata_tools import MetadataTools
from docx_mcp_server.core.table_analyzer import TableStructureAnalyzer
from docx_mcp_server.core.response import (
    create_markdown_response,
    create_error_response
)
from docx_mcp_server.services.navigation import PositionResolver, ContextBuilder
from docx_mcp_server.core.xml_util import ElementManipulator

logger = logging.getLogger(__name__)


def docx_insert_table(session_id: str, rows: int, cols: int, position: str) -> str:
    """
    Create a new table in the document.

    Adds a table with the specified dimensions. Supports precise positioning.

    Args:
        session_id (str): Active session ID returned by docx_create().
        rows (int): Number of rows to create (must be >= 1).
        cols (int): Number of columns to create (must be >= 1).
        position (str): Insertion position string (e.g., "after:para_123").

    Returns:
        str: JSON response with element_id and visual cursor context.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_table called: session_id={session_id}, rows={rows}, cols={cols}, position={position}")
    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_insert_table failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    # Resolve Target Parent and Position
    target_parent = session.document
    ref_element = None
    mode = "append"

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    if not hasattr(target_parent, 'add_table'):
         return create_error_response(f"Object {type(target_parent).__name__} cannot contain tables", error_type="InvalidParent")

    try:
        # Create Table
        # Note: Document.add_table(rows, cols) vs BlockItemContainer.add_table(rows, cols, width)
        # We need to handle both cases.
        try:
            table = target_parent.add_table(rows=rows, cols=cols)
        except TypeError as e:
            if "width" in str(e):
                # Fallback for BlockItemContainer (e.g. Body, Cell) which requires width
                # Default to 6 inches (standard page width approx)
                table = target_parent.add_table(rows=rows, cols=cols, width=Inches(6.0))
            else:
                raise e

        table.style = 'Table Grid' # Default style

        # Move if necessary
        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, table._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, table._element)
            elif mode == "start":
                 # Determine correct container element
                 container_xml = target_parent._element
                 # If target is Document, we want to insert into its body
                 if hasattr(target_parent, '_body'):
                     container_xml = target_parent._body._element

                 ElementManipulator.insert_at_index(container_xml, table._element, 0)

        t_id = session.register_object(table, "table")

        session.update_context(t_id, action="create")

        # Update cursor
        session.cursor.element_id = t_id
        session.cursor.position = "after"

        # Build Response
        builder = ContextBuilder(session)
        data = builder.build_response_data(table, t_id)

        return create_success_response(
            message=f"Table created successfully ({rows}x{cols})",
            rows=rows,
            cols=cols,
            **data
        )
    except Exception as e:
        logger.exception(f"docx_insert_table failed: {e}")
        return create_error_response(f"Failed to create table: {str(e)}", error_type="CreationError")

def docx_get_table(session_id: str, index: int) -> str:
    """
    Get a table by its position index in the document.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_get_table called: session_id={session_id}, index={index}")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        finder = Finder(session.document)
        table = finder.get_table_by_index(index)
        if not table:
            return create_error_response(f"Table at index {index} not found", error_type="ElementNotFound")

        t_id = session._get_element_id(table, auto_register=True)
        session.update_context(t_id, action="access")

        builder = ContextBuilder(session)
        data = builder.build_response_data(table, t_id)

        return create_success_response(
            message=f"Table at index {index} retrieved",
            index=index,
            **data
        )
    except Exception as e:
        logger.exception(f"docx_get_table failed: {e}")
        return create_error_response(f"Failed to get table: {str(e)}", error_type="RetrievalError")


def docx_list_tables(session_id: str, max_results: int = 50, start_element_id: str = None) -> str:
    """
    List tables in document order with basic metadata.
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        tables = list(session.document.tables)
        start_idx = 0
        if start_element_id:
            anchor = session.get_object(start_element_id)
            if anchor is not None:
                anchor_el = getattr(anchor, "_element", None)
                if anchor_el is None and hasattr(anchor, "_tc"):
                    el = anchor._tc
                    while el is not None:
                        if el.tag.split('}')[-1] == 'tbl':
                            anchor_el = el
                            break
                        el = el.getparent()
                if anchor_el is not None:
                    for idx, tbl in enumerate(tables):
                        if tbl._element is anchor_el:
                            start_idx = idx + 1
                            break

        selected = tables[start_idx: start_idx + max_results] if max_results else tables[start_idx:]

        results = []
        for idx, tbl in enumerate(selected, start=start_idx):
            t_id = session._get_element_id(tbl, auto_register=True)
            session.update_context(t_id, action="access")
            first_row_text = "\t".join([cell.text for cell in tbl.rows[0].cells]) if tbl.rows else ""
            results.append({
                "id": t_id,
                "index": idx,
                "rows": len(tbl.rows),
                "cols": len(tbl.columns) if tbl.rows else 0,
                "first_row_text": first_row_text,
            })

        return create_success_response(
            message=f"Listed {len(results)} table(s)",
            count=len(results),
            tables=results,
            start_index=start_idx,
        )
    except Exception as e:
        logger.exception(f"docx_list_tables failed: {e}")
        return create_error_response(f"Failed to list tables: {str(e)}", error_type="SearchError")

def docx_find_table(session_id: str, text: str, max_results: int = 1, start_element_id: str = None, return_structure: bool = False) -> str:
    """
    Find the first table containing specific text in any cell.
    Optionally start search after a given element and limit results.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_find_table called: session_id={session_id}, text='{text}'")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        finder = Finder(session.document)
        tables = finder.find_tables_by_text(text)
        if start_element_id:
            anchor = session.get_object(start_element_id)
            if anchor is not None:
                # compute document order of tables
                order = {tbl: idx for idx, tbl in enumerate(session.document.tables)}
                if hasattr(anchor, "_element"):
                    anchor_el = getattr(anchor, "_element")
                elif hasattr(anchor, "_tc"):
                    el = anchor._tc
                    anchor_el = None
                    while el is not None:
                        if el.tag.split('}')[-1] == 'tbl':
                            anchor_el = el
                            break
                        el = el.getparent()
                else:
                    anchor_el = None

                start_idx = None
                if anchor_el is not None:
                    for tbl, idx in order.items():
                        if tbl._element is anchor_el:
                            start_idx = idx
                            break
                if start_idx is not None:
                    tables = [t for t in tables if order.get(t, -1) > start_idx]

        if not tables:
            return create_error_response(f"No table found containing text '{text}'", error_type="NotFound")

        selected = tables[:max_results] if max_results and max_results > 0 else tables
        results = []
        for tbl in selected:
            t_id = session._get_element_id(tbl, auto_register=True)
            session.update_context(t_id, action="access")
            builder = ContextBuilder(session)
            data = builder.build_response_data(tbl, t_id)
            if return_structure:
                data["rows"] = len(tbl.rows)
                data["cols"] = len(tbl.columns) if tbl.rows else 0
            results.append(data)

        if max_results == 1:
            data = results[0]
            return create_success_response(
                message=f"Table found containing '{text}'",
                search_text=text,
                **data
            )

        return create_success_response(
            message=f"Found {len(results)} table(s) containing '{text}'",
            search_text=text,
            results=results
        )
    except Exception as e:
        logger.exception(f"docx_find_table failed: {e}")
        return create_error_response(f"Failed to find table: {str(e)}", error_type="SearchError")

def docx_get_cell(session_id: str, table_id: str, row: int, col: int) -> str:
    """
    Get a cell from a table by its row and column indices.
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    table = session.get_object(table_id)
    if not table:
        return create_error_response(f"Table {table_id} not found", error_type="ElementNotFound")

    try:
        cell = table.cell(row, col)
        c_id = session.register_object(cell, "cell")

        session.update_context(c_id, action="access")

        # For cells, the visual context is a bit tricky.
        # We probably want to show the table structure or the cell contents.
        # ElementNavigator handles getting path for cell.
        # ContextVisualizer handles visualizing siblings (other cells in row).
        builder = ContextBuilder(session)
        data = builder.build_response_data(cell, c_id)

        return create_success_response(
            message=f"Cell ({row}, {col}) retrieved",
            row=row,
            col=col,
            **data
        )
    except IndexError:
        return create_error_response(f"Cell ({row}, {col}) out of range", error_type="IndexError")
    except Exception as e:
        logger.exception(f"docx_get_cell failed: {e}")
        return create_error_response(f"Failed to get cell: {str(e)}", error_type="RetrievalError")

def docx_insert_paragraph_to_cell(session_id: str, text: str, position: str) -> str:
    """
    Add a paragraph to a table cell.

    Args:
        session_id (str): Active session ID.
        text (str): Paragraph text.
        position (str): Insertion position string targeting a cell (e.g., "inside:cell_123").
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    cell = target_parent
    if not isinstance(cell, _Cell):
        return create_error_response("Position must target a table cell", error_type="InvalidParent")

    try:
        p_id = None
        paragraph = None
        # If the cell is empty (just has the default empty paragraph), reuse it
        if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
            p = cell.paragraphs[0]
            p.text = text
            paragraph = p
            p_id = session.register_object(p, "para")
            session.update_context(p_id, action="access")
        else:
            p = cell.add_paragraph(text)
            paragraph = p
            p_id = session.register_object(p, "para")
            session.update_context(p_id, action="create")

        if mode != "append":
            if mode == "before" and ref_element:
                ElementManipulator.insert_xml_before(ref_element._element, paragraph._element)
            elif mode == "after" and ref_element:
                ElementManipulator.insert_xml_after(ref_element._element, paragraph._element)
            elif mode == "start":
                ElementManipulator.insert_at_index(cell._element, paragraph._element, 0)

        session.cursor.element_id = p_id
        session.cursor.position = "after"

        builder = ContextBuilder(session)
        data = builder.build_response_data(paragraph, p_id)

        return create_markdown_response(
            session=session,
            message="Paragraph added to cell",
            operation="Operation",
            show_context=True,
            **data
        
        )
    except Exception as e:
        logger.exception(f"docx_insert_paragraph_to_cell failed: {e}")
        return create_error_response(f"Failed to add paragraph to cell: {str(e)}", error_type="CreationError")

def docx_insert_table_row(session_id: str, position: str) -> str:
    """
    Add a new row to the end of a table.

    Args:
        session_id (str): Active session ID.
        position (str): Insertion position string (e.g., "inside:table_123").
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_table_row called: session_id={session_id}, position={position}")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    table = target_parent
    if not table or not hasattr(table, 'add_row'):
        return create_error_response("Position must target a table", error_type="InvalidElementType")

    table_id = session._get_element_id(table, auto_register=True)

    try:
        row = table.add_row()
        if mode == "start":
            table._tbl.remove(row._tr)
            table._tbl.insert(0, row._tr)
        elif mode not in ["append"]:
            return create_error_response("Table row insertion only supports inside/end/start on a table", error_type="ValidationError")
        session.update_context(table_id, action="access")

        # We don't have an ID for the Row object usually, we focus on the table or cells.
        # But we can update cursor to the table.
        session.cursor.element_id = table_id
        session.cursor.position = "inside_end"

        # Builder context on the table
        builder = ContextBuilder(session)
        data = builder.build_response_data(table, table_id)

        return create_success_response(
            message=f"Row added to table",
            table_id=table_id,
            new_row_count=len(table.rows),
            **data
        )
    except Exception as e:
        logger.exception(f"docx_insert_table_row failed: {e}")
        return create_error_response(f"Failed to add row: {str(e)}", error_type="ModificationError")

def docx_insert_table_col(session_id: str, position: str) -> str:
    """
    Add a new column to the right side of a table.

    Args:
        session_id (str): Active session ID.
        position (str): Insertion position string (e.g., "inside:table_123").
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_table_col called: session_id={session_id}, position={position}")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)
    except ValueError as e:
        return create_error_response(str(e), error_type="ValidationError")

    table = target_parent
    if not table or not hasattr(table, 'add_column'):
        return create_error_response("Position must target a table", error_type="InvalidElementType")

    table_id = session._get_element_id(table, auto_register=True)

    try:
        if mode not in ["append"]:
            return create_error_response("Table column insertion only supports inside/end on a table", error_type="ValidationError")

        table.add_column(width=Inches(1.0))
        session.update_context(table_id, action="access")

        builder = ContextBuilder(session)
        data = builder.build_response_data(table, table_id)

        return create_success_response(
            message=f"Column added to table",
            table_id=table_id,
            new_col_count=len(table.columns),
            **data
        )
    except Exception as e:
        logger.exception(f"docx_insert_table_col failed: {e}")
        return create_error_response(f"Failed to add column: {str(e)}", error_type="ModificationError")

def _set_cell_text(cell, text: str, preserve_formatting: bool, painter: FormatPainter) -> None:
    """Set cell text, optionally preserving existing run formatting."""
    safe_text = "" if text is None else str(text)

    if not preserve_formatting:
        cell.text = safe_text
        return

    # Ensure there is at least one paragraph
    paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")

    # Choose a source run to copy format from (prefer first run with text)
    source_run = None
    for run in paragraph.runs:
        if run.text:
            source_run = run
            break
    if source_run is None and paragraph.runs:
        source_run = paragraph.runs[0]

    # Clear existing runs to avoid leftover text
    for run in list(paragraph.runs):
        r_el = run._element
        r_el.getparent().remove(r_el)

    new_run = paragraph.add_run(safe_text)

    if source_run:
        painter.copy_format(source_run, new_run)
    elif paragraph.style and hasattr(paragraph.style, "font"):
        # Fallback: copy paragraph style font to the run
        painter.copy_format(paragraph, new_run)


def docx_fill_table(
    session_id: str,
    data: str,
    table_id: str = None,
    start_row: int = 0,
    preserve_formatting: bool = True
) -> str:
    """
    Batch populate table cells with data from a 2D array.

    Args:
        session_id: Active session ID.
        data: JSON string of 2D array data.
        table_id: Target table ID. Defaults to last accessed table.
        start_row: Starting row index for writing.
        preserve_formatting: If True, reuse existing cell run formatting when writing text.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_fill_table called: session_id={session_id}, data_len={len(data)}, table_id={table_id}, start_row={start_row}")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    if not table_id:
        table_id = session.last_accessed_id

    if not table_id:
        return create_error_response("No table specified and no context available", error_type="NoContext")

    table = session.get_object(table_id)
    if not table or not hasattr(table, 'rows'):
        return create_error_response(f"Valid table context not found for ID {table_id}", error_type="InvalidElementType")

    try:
        rows_data = json.loads(data)
    except json.JSONDecodeError as e:
        return create_error_response("Invalid JSON data", error_type="JSONDecodeError")

    if not isinstance(rows_data, list):
        return create_error_response("Data must be a list of lists", error_type="InvalidDataFormat")

    try:
         # Detect irregular structure
         structure_info = TableStructureAnalyzer.detect_irregular_structure(table)
         is_irregular = structure_info["is_irregular"]

         current_row_idx = start_row
         painter = FormatPainter()
         skipped_regions = []
         filled_range = {"start_row": start_row, "start_col": 0, "end_row": start_row, "end_col": 0}

         for row_data in rows_data:
             if current_row_idx >= len(table.rows):
                 table.add_row()

             row = table.rows[current_row_idx]

             for col_idx, cell_value in enumerate(row_data):
                 if col_idx < len(row.cells):
                     # Check if cell is fillable (for irregular tables only)
                     is_fillable = True
                     if is_irregular:
                         # Re-check fillability for this specific cell
                         cell = row.cells[col_idx]
                         try:
                             tc_pr = cell._element.tcPr
                             if tc_pr is not None and tc_pr.gridSpan is not None:
                                 grid_span = tc_pr.gridSpan.val
                                 if grid_span > 1:
                                     is_fillable = False
                         except:
                             pass

                     if is_fillable:
                         cell = row.cells[col_idx]
                         _set_cell_text(cell, cell_value, preserve_formatting, painter)
                         filled_range["end_row"] = current_row_idx
                         filled_range["end_col"] = max(filled_range["end_col"], col_idx)
                     else:
                         # Skip irregular cell
                         skipped_regions.append({
                             "row": current_row_idx,
                             "col": col_idx,
                             "reason": "irregular_cell"
                         })

             current_row_idx += 1

         session.update_context(table_id, action="access")

         builder = ContextBuilder(session)
         data = builder.build_response_data(table, table_id)

         return create_success_response(
             message=f"Table filled with {len(rows_data)} rows",
             rows_filled=len(rows_data),
             start_row=start_row,
             preserve_formatting=preserve_formatting,
             filled_range=filled_range,
             skipped_regions=skipped_regions,
             structure_info=structure_info,
             **data
         )
    except Exception as e:
         logger.exception(f"docx_fill_table failed: {e}")
         return create_error_response(f"Failed to fill table: {str(e)}", error_type="FillError")

def docx_copy_table(session_id: str, table_id: str, position: str) -> str:
    """
    Create a deep copy of an existing table.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_copy_table called: session_id={session_id}, table_id={table_id}")

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    table = session.get_object(table_id)
    if not table:
        return create_error_response(f"Table {table_id} not found", error_type="ElementNotFound")

    if not hasattr(table, 'rows'):
        return create_error_response(f"Object {table_id} is not a table", error_type="InvalidElementType")

    try:
        resolver = PositionResolver(session)
        target_parent, ref_element, mode = resolver.resolve(position, default_parent=session.document)

        engine = CopyEngine()
        new_xml = engine.copy_element(table)

        if mode == "after" and ref_element:
            new_table = engine.insert_element_after(target_parent, new_xml, ref_element._element)
        elif mode == "before" and ref_element:
            # Insert before by placing before ref element, then wrap
            ElementManipulator.insert_xml_before(ref_element._element, new_xml)
            new_table = Table(new_xml, target_parent)
        elif mode == "start":
            container_xml = target_parent._element
            if hasattr(target_parent, '_body'):
                container_xml = target_parent._body._element
            ElementManipulator.insert_at_index(container_xml, new_xml, 0)
            new_table = Table(new_xml, target_parent)
        else:
            # append
            new_table = engine.insert_element_after(target_parent, new_xml, None)

        meta = MetadataTools.create_copy_metadata(
            source_id=table_id,
            source_type="table"
        )

        t_id = session.register_object(new_table, "table", metadata=meta)

        session.update_context(t_id, action="create")

        session.cursor.element_id = t_id
        session.cursor.position = "after"

        builder = ContextBuilder(session)
        data = builder.build_response_data(new_table, t_id)

        return create_markdown_response(
            session=session,
            message="Table copied successfully",
            operation="Operation",
            show_context=True,
            source_id=table_id,
            **data
        
        )
    except Exception as e:
        logger.exception(f"docx_copy_table failed: {e}")
        return create_error_response(f"Failed to copy table: {str(e)}", error_type="CopyError")



def docx_get_table_structure(session_id: str, table_id: str) -> str:
    """
    Get table structure with ASCII visualization.

    Args:
        session_id: Session ID
        table_id: Table ID

    Returns:
        JSON response with ASCII visualization and metadata
    """
    from docx_mcp_server.server import session_manager

    session = session_manager.get_session(session_id)
    if not session:
        return create_error_response(
            f"Session {session_id} not found",
            error_type="SessionNotFound"
        )

    table = session.get_object(table_id)
    if not table or not isinstance(table, Table):
        return create_error_response(
            f"Table {table_id} not found",
            error_type="ElementNotFound"
        )

    try:
        # Generate visualization
        ascii_viz = TableStructureAnalyzer.generate_ascii_visualization(table)
        structure_info = TableStructureAnalyzer.detect_irregular_structure(table, session=session)

        return create_markdown_response(
            session=session,
            message="Table structure retrieved successfully",
            operation="Operation",
            show_context=True,
            element_id=table_id,
            ascii_visualization=ascii_viz,
            structure_info=structure_info,
            rows=len(table.rows
        ),
            cols=len(table.columns) if table.rows else 0
        )
    except Exception as e:
        logger.exception(f"Failed to get table structure: {e}")
        return create_error_response(
            f"Failed to get table structure: {str(e)}",
            error_type="AnalysisError"
        )


def register_tools(mcp: FastMCP):
    """Register table manipulation tools"""
    mcp.tool()(docx_insert_table)
    mcp.tool()(docx_get_table)
    mcp.tool()(docx_list_tables)
    mcp.tool()(docx_find_table)
    mcp.tool()(docx_get_cell)
    mcp.tool()(docx_insert_paragraph_to_cell)
    mcp.tool()(docx_insert_table_row)
    mcp.tool()(docx_insert_table_col)
    mcp.tool()(docx_fill_table)
    mcp.tool()(docx_copy_table)
    mcp.tool()(docx_get_table_structure)
