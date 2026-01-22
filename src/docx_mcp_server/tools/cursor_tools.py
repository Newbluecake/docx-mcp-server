"""Cursor navigation and insertion tools"""
import logging
from typing import Optional, Literal
from mcp.server.fastmcp import FastMCP
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell
from docx.oxml import OxmlElement
from docx.shared import Inches
from docx_mcp_server.core.cursor import Cursor
from docx_mcp_server.core.response import (
    create_context_aware_response,
    create_error_response,
    create_success_response,
    CursorInfo
)

logger = logging.getLogger(__name__)

def docx_cursor_get(session_id: str) -> str:
    """
    Get the current cursor position and state.

    Args:
        session_id (str): Active session ID.

    Returns:
        str: JSON response with cursor state containing:
            - parent_id: ID of the container
            - element_id: ID of the reference element (if any)
            - position: "before", "after", "inside_start", "inside_end"
            - description: Human readable description
            - context: Current cursor context
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_cursor_get called: session_id={session_id}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_cursor_get failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    cursor = session.cursor

    # Add context
    context_text = "Context unavailable"
    try:
        context_text = session.get_cursor_context()
    except Exception as e:
        logger.warning(f"Failed to get cursor context: {e}")

    logger.debug(f"docx_cursor_get success: position={cursor.position}")
    return create_success_response(
        message="Cursor position retrieved successfully",
        parent_id=cursor.parent_id,
        element_id=cursor.element_id,
        position=cursor.position,
        description=f"Cursor is {cursor.position} {cursor.element_id or cursor.parent_id}",
        context=context_text
    )

def docx_cursor_move(
    session_id: str,
    element_id: str,
    position: str = "after"
) -> str:
    """
    Move the insertion cursor to a specific location.

    Args:
        session_id (str): Active session ID.
        element_id (str): Reference element ID (paragraph, table, etc.) or container ID.
        position (str): Relation to element.
            - "before": Place cursor before the element.
            - "after": Place cursor after the element.
            - "inside_start": Place cursor at start of container.
            - "inside_end": Place cursor at end of container.

    Returns:
        str: JSON response with success message and cursor context.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_cursor_move called: session_id={session_id}, element_id={element_id}, position={position}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_cursor_move failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    # Validate position
    if position not in ["before", "after", "inside_start", "inside_end"]:
        logger.error(f"docx_cursor_move failed: Invalid position {position}")
        return create_error_response(f"Invalid position: {position}", error_type="ValidationError")

    # Special case: selecting the document body
    if element_id == "document_body":
        if position in ["before", "after"]:
            logger.error(f"docx_cursor_move failed: Cannot place cursor before/after document body")
            return create_error_response("Cannot place cursor before/after document body", error_type="ValidationError")
        session.cursor.parent_id = "document_body"
        session.cursor.element_id = None
        session.cursor.position = position
        logger.debug(f"docx_cursor_move success: moved to {position} of document")
        return create_context_aware_response(
            session,
            message=f"Cursor moved to {position} of document",
            element_id=element_id,
            position=position
        )

    # Get object to verify it exists and determine parent
    obj = session.get_object(element_id)
    if not obj:
        logger.error(f"docx_cursor_move failed: Element {element_id} not found")
        return create_error_response(f"Element {element_id} not found", error_type="ElementNotFound")

    if position in ["inside_start", "inside_end"]:
        # Element must be a container (Body is handled above, Cell is the other main one)
        # Note: We treat Table as a block, not usually a container for cursor unless we select a cell.
        # But for now let's just allow setting cursor inside container-like things.

        # If it's a paragraph, we don't usually put cursor "inside" it for block insertion,
        # unless we support run insertion (which we do).
        # But 'add_paragraph' inside a paragraph doesn't make sense.
        # For now, we assume 'inside' implies it becomes the parent_id.
        session.cursor.parent_id = element_id
        session.cursor.element_id = None
        session.cursor.position = position
    else:
        # before/after
        # We need to find the parent of this element.
        # python-docx objects don't always expose parent easily or reliably in public API
        # but we can try to find it via session context or rely on the user knowing structure?
        # Actually, for the cursor system, we mostly care about 'where do I insert?'.
        # If I insert 'after' paragraph A, I need to know A's parent.

        # We will optimistically set the cursor.
        # The insertion tool will be responsible for resolving the parent via _parent attribute or XML.

        # Try to resolve parent_id from object if possible (optional enhancement)
        # For now, we keep the previous parent_id if we are just moving around siblings?
        # No, that's risky.
        # We will store the element_id and let insertion logic figure out the parent from the element.

        session.cursor.element_id = element_id
        session.cursor.position = position

        # Attempt to resolve parent_id to ensure context generation works
        # This is critical when moving cursor between different scopes (e.g. from cell to body)
        if hasattr(obj, "_parent"):
            parent = obj._parent
            # Check if parent is the document's body
            # python-docx Paragraph._parent is usually the Body object
            if hasattr(session.document, '_body') and parent == session.document._body:
                 session.cursor.parent_id = "document_body"
            else:
                 # It's likely a cell or other container. Try to get its ID.
                 # Using protected method _get_element_id as we are inside the server package
                 p_id = session._get_element_id(parent, auto_register=True)
                 if p_id:
                     session.cursor.parent_id = p_id

    logger.debug(f"docx_cursor_move success: {position} {element_id}")
    return create_context_aware_response(
        session,
        message=f"Cursor moved {position} {element_id}",
        element_id=element_id,
        position=position
    )


def docx_insert_paragraph_at_cursor(
    session_id: str,
    text: str,
    style: str = None
) -> str:
    """
    Insert a paragraph at the current cursor position.

    Args:
        session_id (str): Active session ID.
        text (str): Paragraph text.
        style (str, optional): Style name.

    Returns:
        str: JSON response with ID of the new paragraph and cursor context.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_paragraph_at_cursor called: session_id={session_id}, text_len={len(text)}, style={style}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_insert_paragraph_at_cursor failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    cursor = session.cursor

    # 1. Determine Parent and Reference Element
    parent = None
    ref_element = None

    if cursor.element_id:
        ref_obj = session.get_object(cursor.element_id)
        if not ref_obj:
            logger.error(f"docx_insert_paragraph_at_cursor failed: Reference element {cursor.element_id} no longer exists")
            return create_error_response(f"Reference element {cursor.element_id} no longer exists", error_type="ElementNotFound")

        if cursor.position in ["before", "after"]:
            # Parent is the parent of the reference object
            # Access internal _parent
            if hasattr(ref_obj, "_parent"):
                parent = ref_obj._parent
            else:
                # Fallback: assume document body if we can't find parent
                # This might fail for nested tables
                parent = session.document
            ref_element = ref_obj
        else:
            # "inside" -> ref_obj IS the parent
            parent = ref_obj
            ref_element = None
    else:
        # No reference element, use parent_id
        if cursor.parent_id == "document_body":
            parent = session.document
        else:
            parent = session.get_object(cursor.parent_id)
            if not parent:
                logger.error(f"docx_insert_paragraph_at_cursor failed: Parent {cursor.parent_id} not found")
                return create_error_response(f"Parent {cursor.parent_id} not found", error_type="ElementNotFound")

    # 2. Create the new paragraph
    # We always use parent.add_paragraph() first to ensure proper creation
    try:
        new_para = parent.add_paragraph(text, style)
    except AttributeError as e:
        logger.error(f"docx_insert_paragraph_at_cursor failed: Target container does not support adding paragraphs")
        return create_error_response(f"Target container ({type(parent).__name__}) does not support adding paragraphs", error_type="InvalidElementType")

    # 3. Move it if necessary
    if cursor.position == "inside_end":
        # Already at end
        pass

    elif cursor.position == "inside_start":
        # Move to beginning of parent
        # parent._element is the container (e.g. w:body, w:tc)
        # We want to insert new_para._element at index 0
        parent_element = parent._element
        if len(parent_element) > 1: # >1 because we just added one
             # Remove from end
             parent_element.remove(new_para._element)
             # Insert at 0
             parent_element.insert(0, new_para._element)

    elif cursor.position == "before" and ref_element:
        # Move before ref_element
        new_para._element.getparent().remove(new_para._element)
        ref_element._element.addprevious(new_para._element)

    elif cursor.position == "after" and ref_element:
        # Move after ref_element
        new_para._element.getparent().remove(new_para._element)
        ref_element._element.addnext(new_para._element)

    para_id = session.register_object(new_para, "para")

    # Update context: creation
    session.update_context(para_id, action="create")

    # Update cursor to point after the new paragraph
    session.cursor.element_id = para_id
    # parent is resolved in logic above, we can try to set it if we know it,
    # or rely on next lookup. For safety/consistency:
    if parent == session.document:
        session.cursor.parent_id = "document_body"
    elif hasattr(parent, "_element") and id(parent._element) in session._element_id_cache:
         session.cursor.parent_id = session._element_id_cache[id(parent._element)]

    session.cursor.position = "after"

    logger.debug(f"docx_insert_paragraph_at_cursor success: {para_id}")
    return create_context_aware_response(
        session,
        message="Paragraph inserted at cursor successfully",
        element_id=para_id
    )


def docx_insert_table_at_cursor(
    session_id: str,
    rows: int,
    cols: int
) -> str:
    """
    Insert a table at the current cursor position.

    Args:
        session_id (str): Active session ID.
        rows (int): Number of rows.
        cols (int): Number of columns.

    Returns:
        str: JSON response with ID of the new table and cursor context.
    """
    from docx_mcp_server.server import session_manager

    logger.debug(f"docx_insert_table_at_cursor called: session_id={session_id}, rows={rows}, cols={cols}")

    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_insert_table_at_cursor failed: Session {session_id} not found")
        return create_error_response(f"Session {session_id} not found", error_type="SessionNotFound")

    cursor = session.cursor

    # Logic similar to insert_paragraph
    parent = None
    ref_element = None

    if cursor.element_id:
        ref_obj = session.get_object(cursor.element_id)
        if not ref_obj:
            logger.error(f"docx_insert_table_at_cursor failed: Reference element {cursor.element_id} no longer exists")
            return create_error_response(f"Reference element {cursor.element_id} no longer exists", error_type="ElementNotFound")

        if cursor.position in ["before", "after"]:
            if hasattr(ref_obj, "_parent"):
                parent = ref_obj._parent
            else:
                parent = session.document
            ref_element = ref_obj
        else:
            parent = ref_obj
            ref_element = None
    else:
        if cursor.parent_id == "document_body":
            parent = session.document
        else:
            parent = session.get_object(cursor.parent_id)
            if not parent:
                logger.error(f"docx_insert_table_at_cursor failed: Parent {cursor.parent_id} not found")
                return create_error_response(f"Parent {cursor.parent_id} not found", error_type="ElementNotFound")

    # Create table
    try:
        new_table = parent.add_table(rows, cols)
    except TypeError as e:
        if "width" in str(e):
             # BlockItemContainer (like Body) requires width
             new_table = parent.add_table(rows, cols, Inches(6.0))
        else:
            logger.error(f"docx_insert_table_at_cursor failed: {e}")
            raise e
    except AttributeError as e:
        logger.error(f"docx_insert_table_at_cursor failed: Target container does not support adding tables")
        return create_error_response(f"Target container ({type(parent).__name__}) does not support adding tables", error_type="InvalidElementType")

    # Move it
    if cursor.position == "inside_start":
        parent_element = parent._element
        if len(parent_element) > 1:
             parent_element.remove(new_table._element)
             parent_element.insert(0, new_table._element)

    elif cursor.position == "before" and ref_element:
        new_table._element.getparent().remove(new_table._element)
        ref_element._element.addprevious(new_table._element)

    elif cursor.position == "after" and ref_element:
        new_table._element.getparent().remove(new_table._element)
        ref_element._element.addnext(new_table._element)

    table_id = session.register_object(new_table, "table")

    # Update context
    session.update_context(table_id, action="create")

    # Update cursor
    session.cursor.element_id = table_id
    session.cursor.position = "after"
    if parent == session.document:
        session.cursor.parent_id = "document_body"

    logger.debug(f"docx_insert_table_at_cursor success: {table_id}")
    return create_context_aware_response(
        session,
        message="Table inserted at cursor successfully",
        element_id=table_id
    )

def register_tools(mcp: FastMCP):
    """Register cursor tools"""
    mcp.tool()(docx_cursor_get)
    mcp.tool()(docx_cursor_move)
    mcp.tool()(docx_insert_paragraph_at_cursor)
    mcp.tool()(docx_insert_table_at_cursor)
