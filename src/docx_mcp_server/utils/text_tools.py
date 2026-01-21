from typing import Dict, List, Any
from docx.text.paragraph import Paragraph
from docx.table import Table

class TextTools:
    """
    Utilities for text manipulation in docx elements.
    """

    @staticmethod
    def collect_paragraphs_from_scope(scope_obj: Any, document: Any = None) -> List[Paragraph]:
        """
        Collect all paragraphs from a scope object (document, table, cell, or paragraph).

        This is a shared utility to avoid duplicating scope resolution logic across
        different text replacement tools.

        Args:
            scope_obj: The scope object to collect paragraphs from. Can be:
                - Document: Collects all paragraphs from document body
                - Table: Collects paragraphs from all cells
                - Cell: Collects paragraphs from the cell
                - Paragraph: Returns the paragraph itself
                - None: Returns all paragraphs from document (requires document param)
            document: Optional document object used when scope_obj is None

        Returns:
            List of Paragraph objects found in the scope

        Examples:
            >>> tool = TextTools()
            >>> # From document
            >>> paragraphs = tool.collect_paragraphs_from_scope(None, document)
            >>> # From table
            >>> paragraphs = tool.collect_paragraphs_from_scope(table_obj)
            >>> # From single paragraph
            >>> paragraphs = tool.collect_paragraphs_from_scope(para_obj)
        """
        targets = []

        if scope_obj is None:
            # Full document scope
            if document is None:
                raise ValueError("Document must be provided when scope_obj is None")
            targets.extend(document.paragraphs)
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        targets.extend(cell.paragraphs)
        else:
            # Identify scope object type and collect paragraphs
            if hasattr(scope_obj, "paragraphs"):
                # Document-like, Cell, or other container
                targets.extend(scope_obj.paragraphs)

            if hasattr(scope_obj, "rows"):
                # Table - collect from all cells
                for row in scope_obj.rows:
                    for cell in row.cells:
                        targets.extend(cell.paragraphs)

            if hasattr(scope_obj, "runs") and scope_obj not in targets:
                # Single Paragraph (and not already added)
                targets.append(scope_obj)

        return targets

    def batch_replace_text(self, elements: List[Any], replacements: Dict[str, str]) -> int:
        """
        Perform batch replacement of text in the provided elements.
        Currently strictly operates on Run level to preserve formatting.
        Does NOT handle text spanning across multiple runs (will skip those matches).

        Args:
            elements: List of docx objects (Paragraphs, Tables, or Cells) to process.
            replacements: Dictionary mapping {old_text: new_text}.

        Returns:
            Total count of replacements made.
        """
        count = 0
        for element in elements:
            if isinstance(element, Paragraph):
                count += self._process_paragraph(element, replacements)
            elif isinstance(element, Table):
                for row in element.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs:
                            count += self._process_paragraph(p, replacements)
            # Handle other types if necessary (e.g. Cell directly if passed)

        return count

    def _process_paragraph(self, paragraph: Paragraph, replacements: Dict[str, str]) -> int:
        count = 0
        for run in paragraph.runs:
            if not run.text:
                continue

            original_text = run.text
            modified_text = original_text

            for old, new in replacements.items():
                if old in modified_text:
                    matches = modified_text.count(old)
                    modified_text = modified_text.replace(old, new)
                    count += matches

            if modified_text != original_text:
                run.text = modified_text

        return count
