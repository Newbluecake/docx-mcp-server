import os
import json
import sys
import time
import platform
import logging
from typing import Optional
from mcp.server.fastmcp import FastMCP
from docx_mcp_server.core.session import SessionManager
from docx_mcp_server.core.properties import set_properties
from docx_mcp_server.core.finder import Finder, list_docx_files
from docx_mcp_server.core.copier import clone_table
from docx_mcp_server.core.replacer import replace_text_in_paragraph
from docx_mcp_server.core.format_painter import FormatPainter
from docx_mcp_server.core.template_parser import TemplateParser
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

SERVER_START_TIME = time.time()
VERSION = "0.1.3"

logger = logging.getLogger(__name__)

# Initialize the MCP server
mcp = FastMCP("docx-mcp-server")

# Global session manager
session_manager = SessionManager()

# Helper for alignment
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

@mcp.tool()
def docx_server_status() -> str:
    """
    Get the current status and environment information of the server.
    Useful for clients to understand the server's running environment (OS, paths).

    Returns:
        str: JSON string containing server status info.
    """
    info = {
        "status": "running",
        "version": VERSION,
        "cwd": os.getcwd(),
        "os_name": os.name,
        "os_system": platform.system(),
        "path_sep": os.sep,
        "python_version": sys.version,
        "start_time": SERVER_START_TIME,
        "uptime_seconds": time.time() - SERVER_START_TIME,
        "active_sessions": len(session_manager.sessions)
    }
    return json.dumps(info, indent=2)

@mcp.tool()
def docx_create(file_path: str = None, auto_save: bool = False) -> str:
    """
    Create a new Word document session or load an existing document.

    This is the entry point for all document operations. Creates an isolated session
    with a unique session_id that maintains document state and object registry.

    Typical Use Cases:
        - Create a new blank document for content generation
        - Load an existing template for modification
        - Enable auto-save for real-time document updates

    Args:
        file_path (str, optional): Path to an existing .docx file to load.
            If None, creates a new blank document. Use relative paths (e.g., "./doc.docx")
            for cross-platform compatibility. Absolute paths must match the server's OS.
        auto_save (bool, optional): Enable automatic saving after each modification.
            Requires file_path to be set. Defaults to False.

    Returns:
        str: Unique session_id (UUID format) for subsequent operations.

    Raises:
        FileNotFoundError: If file_path is specified but file does not exist.
        ValueError: If auto_save is True but file_path is None.

    Examples:
        Create a new blank document:
        >>> session_id = docx_create()
        >>> print(session_id)
        'abc-123-def-456'

        Load an existing document:
        >>> session_id = docx_create(file_path="./template.docx")

        Enable auto-save mode:
        >>> session_id = docx_create(file_path="./output.docx", auto_save=True)

    Notes:
        - Each session is independent and isolated from others
        - Sessions auto-expire after 1 hour of inactivity
        - Always call docx_close() when done to free resources

    See Also:
        - docx_save: Save document to disk
        - docx_close: Close session and free resources
    """
    logger.info(f"docx_create called: file_path={file_path}, auto_save={auto_save}")
    try:
        session_id = session_manager.create_session(file_path, auto_save=auto_save)
        logger.info(f"docx_create success: session_id={session_id}")
        return session_id
    except Exception as e:
        logger.error(f"docx_create failed: {e}")
        raise

@mcp.tool()
def docx_read_content(session_id: str) -> str:
    """
    Read and extract all text content from the document.

    Extracts text from all paragraphs in the document body, preserving order
    but not formatting. Useful for content analysis, search, or preview.

    Typical Use Cases:
        - Preview document content before modification
        - Extract text for analysis or indexing
        - Verify document content after generation

    Args:
        session_id (str): Active session ID returned by docx_create().

    Returns:
        str: Newline-separated text content of all paragraphs.
            Returns "[Empty Document]" if document has no content.

    Raises:
        ValueError: If session_id is invalid or session has expired.

    Examples:
        Read content from a document:
        >>> session_id = docx_create(file_path="./report.docx")
        >>> content = docx_read_content(session_id)
        >>> print(content)
        'Chapter 1\nIntroduction\nThis is the first paragraph...'

    Notes:
        - Only extracts text, formatting information is not included
        - Empty paragraphs are skipped
        - Does not extract text from tables or headers/footers

    See Also:
        - docx_find_paragraphs: Search for specific text in paragraphs
        - docx_extract_template_structure: Get full document structure
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    paragraphs = [p.text for p in session.document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs) if paragraphs else "[Empty Document]"

@mcp.tool()
def docx_find_paragraphs(session_id: str, query: str) -> str:
    """
    Find all paragraphs containing specific text and return their IDs.

    Searches through all paragraphs in the document and returns those containing
    the query text. Useful for locating and modifying specific content.

    Typical Use Cases:
        - Find placeholders to replace (e.g., "{{NAME}}")
        - Locate paragraphs for modification
        - Search document content programmatically

    Args:
        session_id (str): Active session ID returned by docx_create().
        query (str): Text to search for (case-insensitive substring match).

    Returns:
        str: JSON array of objects with paragraph IDs and text:
            [{"id": "para_xxx", "text": "paragraph content"}, ...]

    Raises:
        ValueError: If session_id is invalid or session has expired.

    Examples:
        Find placeholders:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> matches = docx_find_paragraphs(session_id, "{{NAME}}")
        >>> import json
        >>> results = json.loads(matches)
        >>> for match in results:
        ...     docx_update_paragraph_text(session_id, match["id"], "John Doe")

        Search for keyword:
        >>> matches = docx_find_paragraphs(session_id, "important")
        >>> results = json.loads(matches)
        >>> print(f"Found {len(results)} paragraphs with 'important'")

    Notes:
        - Search is case-insensitive
        - Returns all matching paragraphs
        - Paragraphs are automatically registered for subsequent operations
        - Empty paragraphs are not searched

    See Also:
        - docx_update_paragraph_text: Modify found paragraphs
        - docx_replace_text: Replace text globally
        - docx_find_table: Find tables by content
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    matches = []
    for p in session.document.paragraphs:
        if query.lower() in p.text.lower():
            # Only register if we found a match to keep registry clean?
            # The design says yes.
            p_id = session.register_object(p, "para")
            matches.append({"id": p_id, "text": p.text})

    return json.dumps(matches)

@mcp.tool()
def docx_get_table(session_id: str, index: int) -> str:
    """
    Get a table by its position index in the document.

    Retrieves a table by its sequential position (0-based index) in the document.
    Useful when you know the table's position.

    Typical Use Cases:
        - Access tables by position in structured documents
        - Iterate through all tables in a document
        - Access specific table when structure is known

    Args:
        session_id (str): Active session ID returned by docx_create().
        index (int): Zero-based index of the table (0 for first table).

    Returns:
        str: Element ID of the table (format: "table_xxxxx").

    Raises:
        ValueError: If session_id is invalid.
        ValueError: If index is out of range (no table at that position).

    Examples:
        Get first table:
        >>> session_id = docx_create(file_path="./report.docx")
        >>> table_id = docx_get_table(session_id, 0)
        >>> print(table_id)
        'table_abc123'

        Iterate through all tables:
        >>> for i in range(3):
        ...     try:
        ...         table_id = docx_get_table(session_id, i)
        ...         print(f"Table {i}: {table_id}")
        ...     except ValueError:
        ...         break

    Notes:
        - Index is 0-based (first table is index 0)
        - Sets table as current context
        - Raises ValueError if index out of range

    See Also:
        - docx_find_table: Find table by content
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    finder = Finder(session.document)
    table = finder.get_table_by_index(index)
    if not table:
        raise ValueError(f"Table at index {index} not found")

    t_id = session.register_object(table, "table")
    session.update_context(t_id, action="access")
    return t_id

@mcp.tool()
def docx_find_table(session_id: str, text: str) -> str:
    """
    Find the first table containing specific text in any cell.

    Searches through all tables in the document and returns the first table
    that contains the specified text in any of its cells.

    Typical Use Cases:
        - Locate tables by header text
        - Find tables containing specific keywords
        - Navigate to tables in large documents

    Args:
        session_id (str): Active session ID returned by docx_create().
        text (str): Text to search for within table cells (case-insensitive).

    Returns:
        str: Element ID of the found table (format: "table_xxxxx").

    Raises:
        ValueError: If session_id is invalid.
        ValueError: If no table found containing the specified text.

    Examples:
        Find table by header:
        >>> session_id = docx_create(file_path="./report.docx")
        >>> table_id = docx_find_table(session_id, "Sales Data")
        >>> print(table_id)
        'table_abc123'

        Find and modify table:
        >>> table_id = docx_find_table(session_id, "Employee")
        >>> docx_add_table_row(session_id, table_id)

    Notes:
        - Search is case-insensitive
        - Returns first matching table only
        - Searches all cells in all tables
        - Sets table as current context

    See Also:
        - docx_get_table: Get table by index
        - docx_find_paragraphs: Find paragraphs by text
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    finder = Finder(session.document)
    tables = finder.find_tables_by_text(text)
    if not tables:
        raise ValueError(f"No table found containing text '{text}'")

    # Return the first match
    t_id = session.register_object(tables[0], "table")
    session.update_context(t_id, action="access")
    return t_id

@mcp.tool()
def docx_list_files(directory: str = ".") -> str:
    """
    List all .docx files in the specified directory.

    Scans a directory and returns a list of Word document files. Useful for
    discovering available documents before opening them.

    Typical Use Cases:
        - Discover available templates
        - List documents in a folder
        - Find documents to process

    Args:
        directory (str, optional): Directory path to scan. Defaults to current directory (".").
            Can be absolute or relative path.

    Returns:
        str: JSON array of filenames (strings): ["file1.docx", "file2.docx", ...]

    Raises:
        ValueError: If directory does not exist or is not accessible.

    Examples:
        List files in current directory:
        >>> files = docx_list_files()
        >>> import json
        >>> file_list = json.loads(files)
        >>> print(f"Found {len(file_list)} documents")

        List files in specific directory:
        >>> files = docx_list_files("./templates")
        >>> for filename in json.loads(files):
        ...     print(f"Found: {filename}")

    Notes:
        - Only returns .docx files (not .doc or other formats)
        - Returns filenames only, not full paths
        - Does not search subdirectories recursively
        - Hidden files (starting with .) are excluded

    See Also:
        - docx_create: Open discovered files
    """
    try:
        files = list_docx_files(directory)
        return json.dumps(files)
    except Exception as e:
        raise ValueError(str(e))

@mcp.tool()
def docx_copy_table(session_id: str, table_id: str) -> str:
    """
    Create a deep copy of an existing table.

    Duplicates a table with all its content, formatting, and structure. The new
    table is appended to the document.

    Typical Use Cases:
        - Duplicate table templates
        - Create multiple similar tables
        - Replicate table structure with formatting

    Args:
        session_id (str): Active session ID returned by docx_create().
        table_id (str): ID of the source table to copy.

    Returns:
        str: Element ID of the new copied table (format: "table_xxxxx").

    Raises:
        ValueError: If session_id or table_id is invalid.
        ValueError: If specified object is not a table.

    Examples:
        Copy a table:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 2, 3)
        >>> docx_fill_table(session_id, '[["A", "B", "C"]]', table_id)
        >>> new_table_id = docx_copy_table(session_id, table_id)
        >>> print(new_table_id)
        'table_xyz789'

    Notes:
        - Creates complete deep copy (content + formatting)
        - New table is appended to document end
        - Original table is unchanged

    See Also:
        - docx_add_table: Create new table
        - docx_copy_paragraph: Copy paragraphs
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    table = session.get_object(table_id)
    if not table:
        raise ValueError(f"Table {table_id} not found")

    # Check if it is a table
    if not hasattr(table, 'rows'):
        raise ValueError(f"Object {table_id} is not a table")

    new_table = clone_table(table)
    t_id = session.register_object(new_table, "table")

    session.update_context(t_id, action="create")
    return t_id

@mcp.tool()
def docx_add_table_row(session_id: str, table_id: str = None) -> str:
    """
    Add a new row to the end of a table.

    Appends a row with the same number of columns as existing rows. Useful for
    dynamically expanding tables as data is added.

    Typical Use Cases:
        - Add data rows to existing tables
        - Expand tables dynamically
        - Build tables incrementally

    Args:
        session_id (str): Active session ID returned by docx_create().
        table_id (str, optional): ID of the table. If None, uses last accessed table
            from session context.

    Returns:
        str: Success message confirming row addition.

    Raises:
        ValueError: If session_id or table_id is invalid.
        ValueError: If no table context available when table_id is None.

    Examples:
        Add row to specific table:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 2, 3)
        >>> docx_add_table_row(session_id, table_id)
        'Row added to table_xxx'

        Use implicit context:
        >>> table_id = docx_add_table(session_id, 2, 3)
        >>> docx_add_table_row(session_id)

    Notes:
        - New row has same column count as existing rows
        - Cells in new row are empty by default
        - Use docx_get_cell() to populate new row

    See Also:
        - docx_add_table_col: Add column to table
        - docx_get_cell: Access cells in new row
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not table_id:
        table_id = session.last_accessed_id

    table = session.get_object(table_id)
    if not table or not hasattr(table, 'add_row'):
        raise ValueError(f"Valid table context not found for ID {table_id}")

    table.add_row()
    # We could return the new row's cells IDs but usually just adding row is structural.
    # Let's keep context on the table.
    session.update_context(table_id, action="access")
    return f"Row added to {table_id}"

@mcp.tool()
def docx_add_table_col(session_id: str, table_id: str = None) -> str:
    """
    Add a new column to the right side of a table.

    Appends a column to all rows in the table. Default width is 1 inch.

    Typical Use Cases:
        - Expand table structure dynamically
        - Add columns for additional data fields
        - Adjust table layout after creation

    Args:
        session_id (str): Active session ID returned by docx_create().
        table_id (str, optional): ID of the table. If None, uses last accessed table
            from session context.

    Returns:
        str: Success message confirming column addition.

    Raises:
        ValueError: If session_id or table_id is invalid.
        ValueError: If no table context available when table_id is None.

    Examples:
        Add column to table:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 3, 2)
        >>> docx_add_table_col(session_id, table_id)
        'Column added to table_xxx'

    Notes:
        - New column is added to the right (last position)
        - Default width is 1 inch
        - All rows receive the new column

    See Also:
        - docx_add_table_row: Add row to table
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not table_id:
        table_id = session.last_accessed_id

    table = session.get_object(table_id)
    if not table or not hasattr(table, 'add_column'):
        raise ValueError(f"Valid table context not found for ID {table_id}")

    # width is mandatory for add_column in python-docx usually, or defaults?
    # It requires width. Let's default to 1 inch.
    table.add_column(width=Inches(1.0))
    session.update_context(table_id, action="access")
    return f"Column added to {table_id}"

@mcp.tool()
def docx_fill_table(session_id: str, data: str, table_id: str = None, start_row: int = 0) -> str:
    """
    Batch populate table cells with data from a 2D array.

    Efficiently fills multiple rows and columns in one operation. Data is provided
    as a JSON array of arrays, where each inner array represents a row.

    Typical Use Cases:
        - Populate tables from database query results
        - Fill tables with CSV or JSON data
        - Batch insert data efficiently

    Args:
        session_id (str): Active session ID returned by docx_create().
        data (str): JSON string representing 2D array: [["col1", "col2"], ["val1", "val2"]].
            Each inner array is a row, each element is a cell value.
        table_id (str, optional): ID of the table. If None, uses last accessed table.
        start_row (int, optional): Row index to start filling from (0-based). Defaults to 0.

    Returns:
        str: Success message with number of rows filled.

    Raises:
        ValueError: If session_id or table_id is invalid.
        ValueError: If data is not valid JSON or not a 2D array.

    Examples:
        Fill table with data:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 1, 3)
        >>> data = '[["Name", "Age", "City"], ["Alice", "30", "NYC"], ["Bob", "25", "LA"]]'
        >>> result = docx_fill_table(session_id, data, table_id)
        >>> print(result)
        'Table filled with 3 rows starting at 0'

        Fill starting from row 1 (skip header):
        >>> data = '[["Alice", "30", "NYC"], ["Bob", "25", "LA"]]'
        >>> docx_fill_table(session_id, data, table_id, start_row=1)

    Notes:
        - Automatically adds rows if data exceeds table size
        - If data columns > table columns, extra data is truncated
        - Null values are converted to empty strings
        - Overwrites existing cell content

    See Also:
        - docx_add_table: Create table to fill
        - docx_add_paragraph_to_cell: Fill individual cells
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not table_id:
        table_id = session.last_accessed_id

    table = session.get_object(table_id)
    if not table or not hasattr(table, 'rows'):
        raise ValueError(f"Valid table context not found for ID {table_id}")

    try:
        rows_data = json.loads(data)
    except json.JSONDecodeError:
        raise ValueError("Invalid JSON data")

    if not isinstance(rows_data, list):
         raise ValueError("Data must be a list of lists")

    current_row_idx = start_row

    for row_data in rows_data:
        # Ensure we have enough rows
        if current_row_idx >= len(table.rows):
            table.add_row()

        row = table.rows[current_row_idx]

        # Fill cells
        for col_idx, cell_value in enumerate(row_data):
            # Ensure we have enough columns?
            # python-docx doesn't auto-expand cols on row access usually.
            # We skip if out of bounds or expand?
            # Requirement says: "If data cols > table cols, auto truncate or ignore" (from requirements doc analysis implicit)
            # Actually design said: "If data cols > table cols, auto truncate or ignore."
            if col_idx < len(row.cells):
                cell = row.cells[col_idx]
                # We simply set text. If complex formatting needed, use atomic tools.
                cell.text = str(cell_value) if cell_value is not None else ""

        current_row_idx += 1

    session.update_context(table_id, action="access")
    return f"Table filled with {len(rows_data)} rows starting at {start_row}"

@mcp.tool()
def docx_insert_image(session_id: str, image_path: str, width: float = None, height: float = None, parent_id: str = None) -> str:
    """
    Insert an image into the document.

    Adds an image file to the document with optional size constraints. Images can be
    added to the document body or within specific paragraphs.

    Typical Use Cases:
        - Add logos or diagrams to documents
        - Insert charts or screenshots
        - Embed visual content in reports

    Args:
        session_id (str): Active session ID returned by docx_create().
        image_path (str): Absolute or relative path to the image file (PNG, JPG, etc.).
        width (float, optional): Image width in inches. If None, uses original size.
        height (float, optional): Image height in inches. If None, uses original size.
        parent_id (str, optional): ID of parent paragraph. If None, creates new paragraph.

    Returns:
        str: Element ID of the container paragraph (format: "para_xxxxx").

    Raises:
        ValueError: If session_id is invalid or image_path not found.
        FileNotFoundError: If image file does not exist.

    Examples:
        Insert image with default size:
        >>> session_id = docx_create()
        >>> para_id = docx_insert_image(session_id, "./logo.png")

        Insert with specific dimensions:
        >>> para_id = docx_insert_image(session_id, "./chart.png", width=4.0, height=3.0)

        Insert into existing paragraph:
        >>> para_id = docx_add_paragraph(session_id, "See image: ")
        >>> docx_insert_image(session_id, "./diagram.png", width=2.0, parent_id=para_id)

    Notes:
        - Supported formats: PNG, JPG, GIF, BMP
        - Size is in inches (1 inch = 2.54 cm)
        - If only width or height specified, aspect ratio is preserved
        - Image is embedded in document (increases file size)

    See Also:
        - docx_add_paragraph: Create container paragraph
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not os.path.exists(image_path):
        raise ValueError(f"Image file not found: {image_path}")

    # Determine parent
    parent = session.document
    if parent_id:
        parent_obj = session.get_object(parent_id)
        if parent_obj:
            # If parent is a paragraph, we can add a run with picture?
            # python-docx: run.add_picture()
            if hasattr(parent_obj, "add_run"):
                 # It's a paragraph
                 run = parent_obj.add_run()
                 run.add_picture(image_path, width=Inches(width) if width else None, height=Inches(height) if height else None)
                 r_id = session.register_object(run, "run")
                 session.update_context(r_id, action="create")
                 return r_id
            elif hasattr(parent_obj, "add_picture"):
                 # Maybe it's a run already? No, runs don't have add_picture, they ARE the container
                 # Actually run.add_picture exists.
                 pass

    # Default: Add new paragraph then run with picture
    paragraph = session.document.add_paragraph()
    run = paragraph.add_run()
    run.add_picture(image_path, width=Inches(width) if width else None, height=Inches(height) if height else None)

    # Register the run as the object of interest? Or the paragraph?
    # Let's register the paragraph as the structural element.
    p_id = session.register_object(paragraph, "para")
    session.update_context(p_id, action="create")
    return p_id

@mcp.tool()
def docx_replace_text(session_id: str, old_text: str, new_text: str, scope_id: str = None) -> str:
    """
    Replace all occurrences of text in the document or a specific scope.

    Searches and replaces text while preserving formatting. Can operate on the entire
    document or be limited to a specific element (paragraph, table, etc.).

    Typical Use Cases:
        - Replace placeholders in templates
        - Update repeated text globally
        - Batch text modifications

    Args:
        session_id (str): Active session ID returned by docx_create().
        old_text (str): Text to find and replace (exact match).
        new_text (str): Replacement text.
        scope_id (str, optional): ID of element to limit scope (paragraph, table, cell).
            If None, searches entire document body.

    Returns:
        str: Summary message with count of replacements made.

    Raises:
        ValueError: If session_id or scope_id is invalid.

    Examples:
        Replace globally:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> result = docx_replace_text(session_id, "{{COMPANY}}", "Acme Corp")
        >>> print(result)
        'Replaced 5 occurrences of {{COMPANY}}'

        Replace in specific paragraph:
        >>> para_id = docx_find_paragraphs(session_id, "{{NAME}}")[0]["id"]
        >>> docx_replace_text(session_id, "{{NAME}}", "John", scope_id=para_id)

        Replace in table:
        >>> table_id = docx_find_table(session_id, "Invoice")
        >>> docx_replace_text(session_id, "TBD", "Completed", scope_id=table_id)

    Notes:
        - Preserves text formatting (bold, italic, etc.)
        - Exact text match (case-sensitive)
        - Searches document body, tables, and cells
        - Does not search headers/footers

    See Also:
        - docx_find_paragraphs: Find specific paragraphs
        - docx_update_paragraph_text: Replace entire paragraph
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    count = 0
    targets = []

    if scope_id:
        obj = session.get_object(scope_id)
        if not obj:
            raise ValueError(f"Scope object {scope_id} not found")
        # Identify what kind of object
        if hasattr(obj, "paragraphs"): # Document or Cell or Table (no, table has rows)
            targets.extend(obj.paragraphs)
        elif hasattr(obj, "rows"): # Table
            for row in obj.rows:
                for cell in row.cells:
                    targets.extend(cell.paragraphs)
        elif hasattr(obj, "text"): # Paragraph or Run
            # If it's a paragraph, we treat it as a target
            if hasattr(obj, "runs"):
                targets.append(obj)
    else:
        # Full document
        targets.extend(session.document.paragraphs)
        for table in session.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    targets.extend(cell.paragraphs)

    for p in targets:
        if replace_text_in_paragraph(p, old_text, new_text):
            count += 1
            # If successful, we updated the paragraph in place

    return f"Replaced {count} occurrences of '{old_text}'"

@mcp.tool()
def docx_get_context(session_id: str) -> str:
    """
    Get the current context state of the session.

    Returns information about the session's current state, including the last created
    and accessed elements, file path, and auto-save status.

    Typical Use Cases:
        - Debug session state
        - Verify current context before operations
        - Check session configuration

    Args:
        session_id (str): Active session ID returned by docx_create().

    Returns:
        str: JSON string with session context information:
            {
                "session_id": "...",
                "last_created_id": "para_xxx",
                "last_accessed_id": "table_xxx",
                "file_path": "./doc.docx",
                "auto_save": false
            }

    Raises:
        ValueError: If session_id is invalid or session has expired.

    Examples:
        Check session context:
        >>> session_id = docx_create(file_path="./report.docx")
        >>> para_id = docx_add_paragraph(session_id, "Text")
        >>> context = docx_get_context(session_id)
        >>> import json
        >>> data = json.loads(context)
        >>> print(f"Last created: {data['last_created_id']}")

    Notes:
        - Useful for debugging implicit context operations
        - Shows which element will be used for context-based operations
        - Auto-save status indicates if changes are automatically persisted

    See Also:
        - docx_create: Create session with configuration
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    info = {
        "session_id": session.session_id,
        "last_created_id": session.last_created_id,
        "last_accessed_id": session.last_accessed_id,
        "file_path": session.file_path,
        "auto_save": session.auto_save
    }
    return json.dumps(info, indent=2)

@mcp.tool()
def docx_delete(session_id: str, element_id: str = None) -> str:
    """
    Delete an element from the document.

    Removes a paragraph or table from the document. Note that deletion in python-docx
    is complex due to XML structure; primarily supports paragraphs and tables.

    Typical Use Cases:
        - Remove unwanted content
        - Clean up template placeholders
        - Delete empty or obsolete elements

    Args:
        session_id (str): Active session ID returned by docx_create().
        element_id (str, optional): ID of element to delete (paragraph or table).
            If None, uses last accessed element from context.

    Returns:
        str: Success message confirming deletion.

    Raises:
        ValueError: If session_id or element_id is invalid.
        ValueError: If no element context available when element_id is None.
        ValueError: If element cannot be deleted (no parent or unsupported type).

    Examples:
        Delete specific paragraph:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "Temporary")
        >>> docx_delete(session_id, para_id)
        'Deleted para_xxx'

        Delete using context:
        >>> para_id = docx_add_paragraph(session_id, "Remove me")
        >>> docx_delete(session_id)

        Delete found paragraph:
        >>> matches = docx_find_paragraphs(session_id, "{{REMOVE}}")
        >>> for match in json.loads(matches):
        ...     docx_delete(session_id, match["id"])

    Notes:
        - Primarily supports paragraphs and tables
        - Element is removed from object registry
        - Context is reset if deleted element was in context
        - Cannot delete runs independently (delete parent paragraph instead)

    See Also:
        - docx_find_paragraphs: Find paragraphs to delete
        - docx_replace_text: Alternative to deletion for text changes
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not element_id:
        element_id = session.last_accessed_id

    obj = session.get_object(element_id)
    if not obj:
         raise ValueError(f"Object {element_id} not found")

    # Try to delete
    # Strategy: get the XML element and remove it from its parent
    try:
        if hasattr(obj, "_element") and obj._element.getparent() is not None:
            obj._element.getparent().remove(obj._element)
            # Remove from registry?
            if element_id in session.object_registry:
                del session.object_registry[element_id]

            # Reset context if it pointed here
            if session.last_accessed_id == element_id:
                session.last_accessed_id = None
            if session.last_created_id == element_id:
                session.last_created_id = None

            return f"Deleted {element_id}"
        else:
             raise ValueError("Element has no parent or cannot be deleted")
    except Exception as e:
        raise ValueError(f"Failed to delete {element_id}: {str(e)}")

@mcp.tool()
def docx_save(session_id: str, file_path: str) -> str:
    """
    Save the document to disk at the specified path.

    Persists all modifications made during the session to a .docx file.
    Supports saving to a new location or overwriting the original file.

    Typical Use Cases:
        - Save generated document to output location
        - Create a modified copy of a template
        - Checkpoint document state during long operations

    Args:
        session_id (str): Active session ID returned by docx_create().
        file_path (str): Absolute or relative path where the file should be saved.
            Parent directory must exist. Use relative paths for portability.

    Returns:
        str: Success message with the saved file path.

    Raises:
        ValueError: If session_id is invalid or session has expired.
        RuntimeError: If file cannot be saved (permission denied, disk full, etc.).
        FileNotFoundError: If parent directory does not exist.

    Examples:
        Save a new document:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "Hello World")
        >>> result = docx_save(session_id, "./output.docx")
        >>> print(result)
        'Document saved successfully to ./output.docx'

        Save modified template:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> # ... make modifications ...
        >>> docx_save(session_id, "./filled_template.docx")

    Notes:
        - If file exists, it will be overwritten without warning
        - In auto_save mode, this is called automatically after each modification
        - Live preview feature will refresh Word if file is open

    See Also:
        - docx_create: Create session with auto_save option
        - docx_close: Close session after saving
    """
    logger.info(f"docx_save called: session_id={session_id}, file_path={file_path}")
    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_save failed: Session {session_id} not found")
        raise ValueError(f"Session {session_id} not found or expired")

    # Security check: Ensure we are writing to an allowed location if needed.
    # For now, we assume the user (Claude) acts with the user's permissions.

    # LIVE PREVIEW: Prepare for save (release locks if file is open in Word)
    try:
        session.preview_controller.prepare_for_save(file_path)
    except Exception as e:
        # Log but don't block save
        logger.warning(f"Preview prepare failed: {e}")

    try:
        session.document.save(file_path)
        logger.info(f"docx_save success: {file_path}")
    except Exception as e:
        logger.error(f"docx_save failed: {e}")
        raise RuntimeError(f"Failed to save document: {str(e)}")

    # LIVE PREVIEW: Refresh (reload file in Word)
    try:
        session.preview_controller.refresh(file_path)
    except Exception as e:
         logger.warning(f"Preview refresh failed: {e}")

    return f"Document saved successfully to {file_path}"

@mcp.tool()
def docx_close(session_id: str) -> str:
    """
    Close the session and free all associated resources.

    Terminates the document session, releasing memory and clearing the object registry.
    Should be called when document operations are complete to prevent resource leaks.

    Typical Use Cases:
        - Clean up after document generation is complete
        - Free resources in long-running processes
        - Ensure proper session lifecycle management

    Args:
        session_id (str): Active session ID to close.

    Returns:
        str: Success message confirming session closure.

    Examples:
        Complete document workflow:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "Content")
        >>> docx_save(session_id, "./output.docx")
        >>> result = docx_close(session_id)
        >>> print(result)
        'Session abc-123 closed successfully'

    Notes:
        - Unsaved changes will be lost - always call docx_save() first
        - Closed sessions cannot be reopened - create a new session instead
        - Sessions auto-expire after 1 hour, but explicit closure is recommended
        - Calling close on an already closed session returns a not-found message

    See Also:
        - docx_create: Create a new session
        - docx_save: Save before closing
    """
    logger.info(f"docx_close called: session_id={session_id}")
    success = session_manager.close_session(session_id)
    if success:
        return f"Session {session_id} closed successfully"
    else:
        logger.warning(f"docx_close: Session {session_id} not found")
        return f"Session {session_id} not found"

@mcp.tool()
def docx_add_paragraph(session_id: str, text: str, style: str = None, parent_id: str = None) -> str:
    """
    Add a new paragraph to the document or a specific parent container.

    Creates a paragraph with the specified text content. Paragraphs are the fundamental
    building blocks for text content in Word documents.

    Typical Use Cases:
        - Add body text to a document
        - Add content to table cells
        - Create structured content with specific styles

    Args:
        session_id (str): Active session ID returned by docx_create().
        text (str): Text content for the paragraph. Can be empty string.
        style (str, optional): Built-in style name (e.g., 'List Bullet', 'Body Text').
            Defaults to None (Normal style).
        parent_id (str, optional): ID of parent container (e.g., cell_id from docx_get_cell).
            If None, adds to document body.

    Returns:
        str: Element ID of the new paragraph (format: "para_xxxxx").

    Raises:
        ValueError: If session_id is invalid or parent_id not found.
        ValueError: If parent object cannot contain paragraphs.

    Examples:
        Add paragraph to document:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "Hello World")
        >>> print(para_id)
        'para_a1b2c3d4'

        Add styled paragraph:
        >>> para_id = docx_add_paragraph(session_id, "Item 1", style="List Bullet")

        Add paragraph to table cell:
        >>> table_id = docx_add_table(session_id, 2, 2)
        >>> cell_id = docx_get_cell(session_id, table_id, 0, 0)
        >>> para_id = docx_add_paragraph(session_id, "Cell content", parent_id=cell_id)

    Notes:
        - Returns element_id for use in subsequent operations (e.g., adding runs)
        - Sets session context to this paragraph for implicit operations
        - Empty text is valid - use for formatting-only paragraphs

    See Also:
        - docx_add_run: Add formatted text to paragraph
        - docx_add_heading: Add heading instead of paragraph
        - docx_update_paragraph_text: Modify existing paragraph
    """
    logger.debug(f"docx_add_paragraph called: session_id={session_id}, parent_id={parent_id}")
    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_add_paragraph failed: Session {session_id} not found")
        raise ValueError(f"Session {session_id} not found")

    parent = session.document
    if parent_id:
        parent = session.get_object(parent_id)
        if not parent:
            logger.error(f"docx_add_paragraph failed: Parent {parent_id} not found")
            raise ValueError(f"Parent object {parent_id} not found")
        if not hasattr(parent, 'add_paragraph'):
            logger.error(f"docx_add_paragraph failed: Parent {parent_id} cannot contain paragraphs")
            raise ValueError(f"Object {parent_id} cannot contain paragraphs")

    paragraph = parent.add_paragraph(text, style=style)
    p_id = session.register_object(paragraph, "para")

    # Update context: this is a creation action
    session.update_context(p_id, action="create")

    logger.debug(f"docx_add_paragraph success: {p_id}")
    return p_id

@mcp.tool()
def docx_add_heading(session_id: str, text: str, level: int = 1) -> str:
    """
    Add a heading to the document.

    Creates a heading paragraph with the specified level. Headings provide document
    structure and are used for navigation and table of contents generation.

    Typical Use Cases:
        - Create document sections and chapters
        - Structure reports and articles
        - Generate navigable document outlines

    Args:
        session_id (str): Active session ID returned by docx_create().
        text (str): Heading text content.
        level (int, optional): Heading level from 0-9. Level 0 is Title style,
            levels 1-9 are Heading 1-9 styles. Defaults to 1.

    Returns:
        str: Element ID of the new heading paragraph (format: "para_xxxxx").

    Raises:
        ValueError: If session_id is invalid or level is out of range.

    Examples:
        Add main heading:
        >>> session_id = docx_create()
        >>> h1_id = docx_add_heading(session_id, "Chapter 1", level=1)

        Add title and subheadings:
        >>> title_id = docx_add_heading(session_id, "Report Title", level=0)
        >>> h2_id = docx_add_heading(session_id, "Introduction", level=2)

    Notes:
        - Level 0 applies Title style (larger, centered)
        - Levels 1-9 apply Heading 1-9 styles
        - Headings are paragraphs with special styling

    See Also:
        - docx_add_paragraph: Add regular paragraph
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    heading = session.document.add_heading(text, level=level)
    return session.register_object(heading, "para")

@mcp.tool()
def docx_add_run(session_id: str, text: str, paragraph_id: str = None) -> str:
    """
    Add a text run to a paragraph with independent formatting.

    Runs are the atomic text units within paragraphs. Each run can have its own
    formatting (bold, italic, color, size) independent of other runs in the paragraph.

    Typical Use Cases:
        - Add formatted text within a paragraph
        - Mix different text styles in one paragraph
        - Build complex formatted content incrementally

    Args:
        session_id (str): Active session ID returned by docx_create().
        text (str): Text content for the run.
        paragraph_id (str, optional): ID of the parent paragraph. If None, uses
            the last created paragraph from session context.

    Returns:
        str: Element ID of the new run (format: "run_xxxxx").

    Raises:
        ValueError: If session_id is invalid or paragraph_id not found.
        ValueError: If no paragraph context available when paragraph_id is None.
        ValueError: If specified object is not a paragraph.

    Examples:
        Add run to specific paragraph:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Hello ", paragraph_id=para_id)
        >>> run_id2 = docx_add_run(session_id, "World", paragraph_id=para_id)

        Use implicit context:
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Implicit context")

        Build formatted paragraph:
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run1 = docx_add_run(session_id, "Normal ", paragraph_id=para_id)
        >>> run2 = docx_add_run(session_id, "Bold", paragraph_id=para_id)
        >>> docx_set_font(session_id, run2, bold=True)

    Notes:
        - Supports legacy signature: docx_add_run(sid, para_id, text) for compatibility
        - Context mechanism allows omitting paragraph_id for sequential operations
        - Use docx_set_font() to apply formatting after creation

    See Also:
        - docx_add_paragraph: Create parent paragraph
        - docx_set_font: Format the run
        - docx_update_run_text: Modify existing run
    """
    # Compatibility shim for legacy calls: docx_add_run(sid, para_id, text)
    # If 'text' looks like a para_id and 'paragraph_id' is present (and doesn't look like a para_id)
    if text and text.startswith("para_") and paragraph_id and not paragraph_id.startswith("para_"):
        # Swap arguments
        real_para_id = text
        real_text = paragraph_id
        text = real_text
        paragraph_id = real_para_id

    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if paragraph_id is None:
        # Implicit context
        if not session.last_created_id:
            raise ValueError("No paragraph context available. Please specify paragraph_id.")
        paragraph_id = session.last_created_id

    paragraph = session.get_object(paragraph_id)
    if not paragraph:
        raise ValueError(f"Paragraph {paragraph_id} not found")

    # Simple type check
    if not hasattr(paragraph, 'add_run'):
         # If the last_created_id was a table or something else, we might have an issue
         raise ValueError(f"Object {paragraph_id} is not a paragraph (cannot add run)")

    run = paragraph.add_run(text)
    r_id = session.register_object(run, "run")

    # Update context: access action (we modified the paragraph, but the 'current' object of interest is now the run)
    session.update_context(r_id, action="access")

    return r_id

@mcp.tool()
def docx_set_properties(session_id: str, properties: str, element_id: str = None) -> str:
    """
    Set advanced properties on a document element using JSON configuration.

    Provides flexible property setting for runs, paragraphs, tables, and cells.
    Supports font, paragraph format, and other advanced formatting options.

    Typical Use Cases:
        - Apply complex formatting in one call
        - Set properties not covered by dedicated tools
        - Batch apply multiple formatting options

    Args:
        session_id (str): Active session ID returned by docx_create().
        properties (str): JSON string defining properties to set. Structure:
            {
                "font": {"bold": true, "size": 12, "name": "Arial"},
                "paragraph_format": {"alignment": "center", "line_spacing": 1.5}
            }
        element_id (str, optional): ID of element to modify (run, paragraph, table, cell).
            If None, uses last accessed object from session context.

    Returns:
        str: Success message confirming property update.

    Raises:
        ValueError: If session_id or element_id is invalid.
        ValueError: If properties is not valid JSON.
        ValueError: If no element context available when element_id is None.

    Examples:
        Set font properties:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Text", paragraph_id=para_id)
        >>> props = '{"font": {"bold": true, "size": 14, "name": "Arial"}}'
        >>> docx_set_properties(session_id, props, element_id=run_id)

        Set paragraph format:
        >>> props = '{"paragraph_format": {"alignment": "center", "line_spacing": 1.5}}'
        >>> docx_set_properties(session_id, props, element_id=para_id)

    Notes:
        - JSON must be valid and properly escaped
        - Available properties depend on element type
        - Use dedicated tools (docx_set_font, docx_set_alignment) for simpler cases

    See Also:
        - docx_set_font: Simpler font formatting
        - docx_set_alignment: Simpler alignment setting
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    target_id = element_id
    if not target_id:
        target_id = session.last_accessed_id
        if not target_id:
            raise ValueError("No element context available. Please specify element_id.")

    obj = session.get_object(target_id)
    if not obj:
        raise ValueError(f"Object {target_id} not found")

    try:
        props_dict = json.loads(properties)
    except json.JSONDecodeError:
        raise ValueError("Invalid JSON string in properties")

    set_properties(obj, props_dict)

    # Update context: we accessed/modified this object
    session.update_context(target_id, action="access")

    return f"Properties updated for {target_id}"

@mcp.tool()
def docx_format_copy(session_id: str, source_id: str, target_id: str) -> str:
    """
    Copy formatting from source element to target element.

    Transfers formatting properties between elements of the same type. Supports
    Run-to-Run, Paragraph-to-Paragraph, and Table-to-Table copying.

    Typical Use Cases:
        - Apply consistent formatting across elements
        - Replicate styling from templates
        - Maintain formatting consistency

    Args:
        session_id (str): Active session ID returned by docx_create().
        source_id (str): ID of source element to copy formatting from.
        target_id (str): ID of target element to apply formatting to.

    Returns:
        str: Success message confirming format copy.

    Raises:
        ValueError: If session_id, source_id, or target_id is invalid.
        ValueError: If source and target are not the same element type.

    Examples:
        Copy run formatting:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run1 = docx_add_run(session_id, "Styled", paragraph_id=para_id)
        >>> docx_set_font(session_id, run1, bold=True, size=14, color_hex="FF0000")
        >>> run2 = docx_add_run(session_id, "Copy style", paragraph_id=para_id)
        >>> docx_format_copy(session_id, run1, run2)

        Copy paragraph formatting:
        >>> para1 = docx_add_paragraph(session_id, "Source")
        >>> docx_set_alignment(session_id, para1, "center")
        >>> para2 = docx_add_paragraph(session_id, "Target")
        >>> docx_format_copy(session_id, para1, para2)

    Notes:
        - Source and target must be same type (run/paragraph/table)
        - Only formatting is copied, not content
        - Target content is preserved
        - Supports: font properties, alignment, spacing, borders

    See Also:
        - docx_set_font: Set specific font properties
        - docx_set_alignment: Set alignment directly
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    source = session.get_object(source_id)
    if not source:
        raise ValueError(f"Source object {source_id} not found")

    target = session.get_object(target_id)
    if not target:
        raise ValueError(f"Target object {target_id} not found")

    painter = FormatPainter()
    painter.copy_format(source, target)

    # Update context: we modified the target
    session.update_context(target_id, action="access")

    return f"Format copied from {source_id} to {target_id}"

@mcp.tool()
def docx_set_font(
    session_id: str,
    run_id: str,
    size: float = None,
    bold: bool = None,
    italic: bool = None,
    color_hex: str = None
) -> str:
    """
    Set font properties for a specific text run.

    Applies formatting to a run, allowing fine-grained control over text appearance.
    Multiple properties can be set in a single call.

    Typical Use Cases:
        - Format important text (bold, larger size)
        - Apply color coding to text
        - Create styled headings and emphasis

    Args:
        session_id (str): Active session ID returned by docx_create().
        run_id (str): ID of the text run to format.
        size (float, optional): Font size in points (e.g., 12.0, 14.5).
        bold (bool, optional): True to make bold, False to remove bold.
        italic (bool, optional): True to make italic, False to remove italic.
        color_hex (str, optional): Hex color code without '#' (e.g., "FF0000" for red,
            "0000FF" for blue). Must be 6 characters (RRGGBB format).

    Returns:
        str: Success message confirming font update.

    Raises:
        ValueError: If session_id or run_id is invalid.
        ValueError: If color_hex is not valid 6-character hex string.

    Examples:
        Basic formatting:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Important", paragraph_id=para_id)
        >>> docx_set_font(session_id, run_id, size=14, bold=True)

        Apply color:
        >>> run_id = docx_add_run(session_id, "Red text", paragraph_id=para_id)
        >>> docx_set_font(session_id, run_id, color_hex="FF0000")

        Multiple properties:
        >>> run_id = docx_add_run(session_id, "Styled", paragraph_id=para_id)
        >>> docx_set_font(session_id, run_id, size=16, bold=True, italic=True, color_hex="0000FF")

    Notes:
        - Size is in points (1 point = 1/72 inch)
        - Color format is RRGGBB without '#' prefix
        - Omitted parameters are not changed
        - Setting bold/italic to False explicitly removes the formatting

    See Also:
        - docx_add_run: Create run to format
        - docx_set_properties: Advanced formatting options
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    run = session.get_object(run_id)
    if not run:
        raise ValueError(f"Run {run_id} not found")

    font = run.font
    if size is not None:
        font.size = Pt(size)
    if bold is not None:
        font.bold = bold
    if italic is not None:
        font.italic = italic
    if color_hex:
        # Parse hex string "RRGGBB"
        try:
            r = int(color_hex[0:2], 16)
            g = int(color_hex[2:4], 16)
            b = int(color_hex[4:6], 16)
            font.color.rgb = RGBColor(r, g, b)
        except ValueError:
            raise ValueError(f"Invalid hex color: {color_hex}")

    return f"Font updated for {run_id}"

@mcp.tool()
def docx_set_alignment(session_id: str, paragraph_id: str, alignment: str) -> str:
    """
    Set horizontal alignment for a paragraph.

    Controls how text is aligned within the paragraph margins. Affects the entire
    paragraph, not individual runs.

    Typical Use Cases:
        - Center headings and titles
        - Right-align dates or signatures
        - Justify body text for formal documents

    Args:
        session_id (str): Active session ID returned by docx_create().
        paragraph_id (str): ID of the paragraph to align.
        alignment (str): Alignment value. Must be one of:
            - "left": Align text to left margin (default)
            - "center": Center text between margins
            - "right": Align text to right margin
            - "justify": Distribute text evenly between margins

    Returns:
        str: Success message confirming alignment change.

    Raises:
        ValueError: If session_id or paragraph_id is invalid.
        ValueError: If alignment is not one of the valid values.

    Examples:
        Center a heading:
        >>> session_id = docx_create()
        >>> para_id = docx_add_heading(session_id, "Title", level=0)
        >>> docx_set_alignment(session_id, para_id, "center")

        Right-align a date:
        >>> para_id = docx_add_paragraph(session_id, "January 21, 2026")
        >>> docx_set_alignment(session_id, para_id, "right")

        Justify body text:
        >>> para_id = docx_add_paragraph(session_id, "Long paragraph text...")
        >>> docx_set_alignment(session_id, para_id, "justify")

    Notes:
        - Alignment is case-insensitive
        - Applies to entire paragraph, not individual runs
        - Default alignment is "left"

    See Also:
        - docx_add_paragraph: Create paragraph to align
        - docx_set_properties: Advanced paragraph formatting
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    paragraph = session.get_object(paragraph_id)
    if not paragraph:
        raise ValueError(f"Paragraph {paragraph_id} not found")

    if alignment.lower() not in ALIGNMENT_MAP:
        raise ValueError(f"Invalid alignment: {alignment}. Must be one of {list(ALIGNMENT_MAP.keys())}")

    paragraph.alignment = ALIGNMENT_MAP[alignment.lower()]
    return f"Alignment set to {alignment} for {paragraph_id}"

@mcp.tool()
def docx_add_page_break(session_id: str) -> str:
    """
    Insert a page break at the current position in the document.

    Forces content after the break to start on a new page. Useful for separating
    sections or chapters in a document.

    Typical Use Cases:
        - Separate chapters or major sections
        - Start new content on a fresh page
        - Control pagination in reports

    Args:
        session_id (str): Active session ID returned by docx_create().

    Returns:
        str: Success message confirming page break insertion.

    Raises:
        ValueError: If session_id is invalid or session has expired.

    Examples:
        Add page break between sections:
        >>> session_id = docx_create()
        >>> docx_add_heading(session_id, "Chapter 1", level=1)
        >>> docx_add_paragraph(session_id, "Content of chapter 1...")
        >>> docx_add_page_break(session_id)
        >>> docx_add_heading(session_id, "Chapter 2", level=1)

    Notes:
        - Page break is inserted at the end of the document
        - Content added after this call will appear on the new page
        - Does not return an element_id (structural operation)

    See Also:
        - docx_add_heading: Add section headings
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    session.document.add_page_break()
    return "Page break added"

@mcp.tool()
def docx_set_margins(
    session_id: str,
    top: float = None,
    bottom: float = None,
    left: float = None,
    right: float = None
) -> str:
    """
    Set page margins for the document section.

    Adjusts the white space around the content area on each page. Applies to the
    last section of the document (typically the entire document if single-section).

    Typical Use Cases:
        - Adjust margins for printing requirements
        - Create narrow margins for more content
        - Set wide margins for annotations

    Args:
        session_id (str): Active session ID returned by docx_create().
        top (float, optional): Top margin in inches.
        bottom (float, optional): Bottom margin in inches.
        left (float, optional): Left margin in inches.
        right (float, optional): Right margin in inches.

    Returns:
        str: Success message confirming margin update.

    Raises:
        ValueError: If session_id is invalid or session has expired.

    Examples:
        Set all margins:
        >>> session_id = docx_create()
        >>> docx_set_margins(session_id, top=1.0, bottom=1.0, left=1.0, right=1.0)

        Set narrow margins:
        >>> docx_set_margins(session_id, top=0.5, bottom=0.5, left=0.5, right=0.5)

        Set only top and bottom:
        >>> docx_set_margins(session_id, top=1.5, bottom=1.5)

    Notes:
        - All measurements are in inches (1 inch = 2.54 cm)
        - Omitted parameters leave existing margins unchanged
        - Applies to the last section (usually entire document)
        - Default Word margins are typically 1 inch on all sides

    See Also:
        - docx_add_page_break: Control page breaks
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    # For simplicity, apply to the last section (usually where we are working)
    # or all sections? Let's apply to the last section for now as a default "current" context.
    section = session.document.sections[-1]

    if top is not None:
        section.top_margin = Inches(top)
    if bottom is not None:
        section.bottom_margin = Inches(bottom)
    if left is not None:
        section.left_margin = Inches(left)
    if right is not None:
        section.right_margin = Inches(right)

    return "Margins updated"

@mcp.tool()
def docx_add_table(session_id: str, rows: int, cols: int) -> str:
    """
    Create a new table in the document.

    Adds a table with the specified dimensions. Tables are created with default
    'Table Grid' style and can be populated using cell operations.

    Typical Use Cases:
        - Create data tables for reports
        - Build structured layouts
        - Display tabular information

    Args:
        session_id (str): Active session ID returned by docx_create().
        rows (int): Number of rows to create (must be >= 1).
        cols (int): Number of columns to create (must be >= 1).

    Returns:
        str: Element ID of the new table (format: "table_xxxxx").

    Raises:
        ValueError: If session_id is invalid or rows/cols are less than 1.

    Examples:
        Create a simple table:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, rows=3, cols=2)
        >>> print(table_id)
        'table_a1b2c3d4'

        Create and populate table:
        >>> table_id = docx_add_table(session_id, 2, 3)
        >>> cell_id = docx_get_cell(session_id, table_id, 0, 0)
        >>> docx_add_paragraph_to_cell(session_id, cell_id, "Header 1")

    Notes:
        - Default style is 'Table Grid' with borders
        - Cells are not pre-registered - use docx_get_cell() to access them
        - Rows and columns can be added later with docx_add_table_row/col()

    See Also:
        - docx_get_cell: Access table cells
        - docx_fill_table: Batch populate table data
        - docx_add_table_row: Add rows dynamically
    """
    logger.debug(f"docx_add_table called: session_id={session_id}, rows={rows}, cols={cols}")
    session = session_manager.get_session(session_id)
    if not session:
        logger.error(f"docx_add_table failed: Session {session_id} not found")
        raise ValueError(f"Session {session_id} not found")

    table = session.document.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid' # Default style
    t_id = session.register_object(table, "table")

    session.update_context(t_id, action="create")

    logger.debug(f"docx_add_table success: {t_id}")
    return t_id

@mcp.tool()
def docx_get_cell(session_id: str, table_id: str, row: int, col: int) -> str:
    """
    Get a cell from a table by its row and column indices.

    Retrieves and registers a specific cell for subsequent operations. Cells are
    not automatically registered when tables are created to conserve memory.

    Typical Use Cases:
        - Access cells to add content
        - Modify specific cell properties
        - Navigate table structure

    Args:
        session_id (str): Active session ID returned by docx_create().
        table_id (str): ID of the table containing the cell.
        row (int): Row index (0-based, 0 is first row).
        col (int): Column index (0-based, 0 is first column).

    Returns:
        str: Element ID of the cell (format: "cell_xxxxx").

    Raises:
        ValueError: If session_id or table_id is invalid.
        ValueError: If row or col indices are out of range.

    Examples:
        Access table cells:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 3, 2)
        >>> cell_id = docx_get_cell(session_id, table_id, 0, 0)
        >>> print(cell_id)
        'cell_a1b2c3d4'

        Populate first row:
        >>> for col in range(2):
        ...     cell_id = docx_get_cell(session_id, table_id, 0, col)
        ...     docx_add_paragraph_to_cell(session_id, cell_id, f"Header {col+1}")

    Notes:
        - Indices are 0-based (first cell is row=0, col=0)
        - Cell is registered and can be reused in subsequent calls
        - Out-of-range indices raise ValueError

    See Also:
        - docx_add_table: Create table
        - docx_add_paragraph_to_cell: Add content to cell
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    table = session.get_object(table_id)
    if not table:
        raise ValueError(f"Table {table_id} not found")

    try:
        cell = table.cell(row, col)
    except IndexError:
        raise ValueError(f"Cell ({row}, {col}) out of range")

    return session.register_object(cell, "cell")

@mcp.tool()
def docx_add_paragraph_to_cell(session_id: str, cell_id: str, text: str) -> str:
    """
    Add a paragraph to a table cell.

    Adds text content to a cell. If the cell contains only an empty default paragraph,
    it reuses that paragraph; otherwise, it adds a new paragraph.

    Typical Use Cases:
        - Populate table cells with content
        - Add multiple paragraphs to a single cell
        - Fill table data programmatically

    Args:
        session_id (str): Active session ID returned by docx_create().
        cell_id (str): ID of the cell (from docx_get_cell()).
        text (str): Text content for the paragraph.

    Returns:
        str: Element ID of the paragraph (format: "para_xxxxx").

    Raises:
        ValueError: If session_id or cell_id is invalid.

    Examples:
        Add content to cell:
        >>> session_id = docx_create()
        >>> table_id = docx_add_table(session_id, 2, 2)
        >>> cell_id = docx_get_cell(session_id, table_id, 0, 0)
        >>> para_id = docx_add_paragraph_to_cell(session_id, cell_id, "Header")

        Add multiple paragraphs to cell:
        >>> para1 = docx_add_paragraph_to_cell(session_id, cell_id, "Line 1")
        >>> para2 = docx_add_paragraph_to_cell(session_id, cell_id, "Line 2")

    Notes:
        - Cells have one empty paragraph by default
        - First call reuses default paragraph if empty
        - Subsequent calls add new paragraphs

    See Also:
        - docx_get_cell: Get cell to populate
        - docx_fill_table: Batch populate multiple cells
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    cell = session.get_object(cell_id)
    if not cell:
        raise ValueError(f"Cell {cell_id} not found")

    # If the cell is empty (just has the default empty paragraph), we might want to use it?
    # For simplicity/consistency, let's just use the add_paragraph method of the cell.
    # But usually, user wants to fill the cell.
    # Check if cell.paragraphs[0] is empty?
    if len(cell.paragraphs) == 1 and not cell.paragraphs[0].text:
        p = cell.paragraphs[0]
        p.text = text
        return session.register_object(p, "para")
    else:
        p = cell.add_paragraph(text)
        return session.register_object(p, "para")

@mcp.tool()
def docx_copy_paragraph(session_id: str, paragraph_id: str) -> str:
    """
    Create a deep copy of a paragraph with all formatting and runs.

    Duplicates a paragraph including its style, alignment, and all text runs with
    their individual formatting. The new paragraph is appended to the document.

    Typical Use Cases:
        - Duplicate formatted content
        - Replicate paragraph templates
        - Copy complex formatted text

    Args:
        session_id (str): Active session ID returned by docx_create().
        paragraph_id (str): ID of the paragraph to copy.

    Returns:
        str: Element ID of the new copied paragraph (format: "para_xxxxx").

    Raises:
        ValueError: If session_id or paragraph_id is invalid.
        ValueError: If specified object is not a paragraph.

    Examples:
        Copy a paragraph:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Formatted", paragraph_id=para_id)
        >>> docx_set_font(session_id, run_id, bold=True, size=14)
        >>> new_para_id = docx_copy_paragraph(session_id, para_id)

    Notes:
        - Preserves paragraph style and alignment
        - Copies all runs with their formatting (bold, italic, size, color)
        - New paragraph is appended to document end
        - Original paragraph is unchanged

    See Also:
        - docx_copy_table: Copy tables
        - docx_add_paragraph: Create new paragraph
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    source_para = session.get_object(paragraph_id)
    if not source_para:
        raise ValueError(f"Paragraph {paragraph_id} not found")

    if not hasattr(source_para, 'runs'):
        raise ValueError(f"Object {paragraph_id} is not a paragraph")

    # Create new paragraph with same style
    new_para = session.document.add_paragraph(style=source_para.style)

    # Copy paragraph-level formatting
    if source_para.alignment is not None:
        new_para.alignment = source_para.alignment

    # Copy all runs with their formatting
    for run in source_para.runs:
        new_run = new_para.add_run(run.text)

        # Copy font properties
        if run.font.bold is not None:
            new_run.font.bold = run.font.bold
        if run.font.italic is not None:
            new_run.font.italic = run.font.italic
        if run.font.underline is not None:
            new_run.font.underline = run.font.underline
        if run.font.size is not None:
            new_run.font.size = run.font.size
        if run.font.color.rgb is not None:
            new_run.font.color.rgb = run.font.color.rgb

    return session.register_object(new_para, "para")

@mcp.tool()
def docx_update_paragraph_text(session_id: str, paragraph_id: str, new_text: str) -> str:
    """
    Replace all text content in an existing paragraph.

    Clears all existing runs in the paragraph and replaces with a single run
    containing the new text. All previous formatting within the paragraph is lost.

    Typical Use Cases:
        - Update placeholder text in templates
        - Replace entire paragraph content
        - Reset paragraph to plain text

    Args:
        session_id (str): Active session ID returned by docx_create().
        paragraph_id (str): ID of the paragraph to update.
        new_text (str): New text content to replace existing content.

    Returns:
        str: Success message confirming the update.

    Raises:
        ValueError: If session_id or paragraph_id is invalid.
        ValueError: If specified object is not a paragraph.

    Examples:
        Update paragraph text:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "Old text")
        >>> result = docx_update_paragraph_text(session_id, para_id, "New text")
        >>> print(result)
        'Paragraph para_xxx updated successfully'

        Update template placeholder:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> matches = docx_find_paragraphs(session_id, "{{NAME}}")
        >>> para_id = json.loads(matches)[0]["id"]
        >>> docx_update_paragraph_text(session_id, para_id, "John Doe")

    Notes:
        - Removes all existing runs and their formatting
        - Creates a single new run with the new text
        - To preserve formatting, use docx_update_run_text() on individual runs

    See Also:
        - docx_update_run_text: Update run while preserving formatting
        - docx_find_paragraphs: Find paragraphs to update
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    paragraph = session.get_object(paragraph_id)
    if not paragraph:
        raise ValueError(f"Paragraph {paragraph_id} not found")

    if not hasattr(paragraph, 'text'):
        raise ValueError(f"Object {paragraph_id} is not a paragraph")

    # Clear existing runs and set new text
    paragraph.clear()
    paragraph.add_run(new_text)

    return f"Paragraph {paragraph_id} updated successfully"

@mcp.tool()
def docx_update_run_text(session_id: str, run_id: str, new_text: str) -> str:
    """
    Update the text content of an existing run while preserving formatting.

    Replaces the text in a run without affecting its formatting properties
    (bold, italic, color, size, etc.). Useful for updating content while maintaining style.

    Typical Use Cases:
        - Update formatted text without losing styling
        - Replace placeholders in styled content
        - Modify specific text portions in a paragraph

    Args:
        session_id (str): Active session ID returned by docx_create().
        run_id (str): ID of the run to update.
        new_text (str): New text content to replace existing text.

    Returns:
        str: Success message confirming the update.

    Raises:
        ValueError: If session_id or run_id is invalid.
        ValueError: If specified object is not a run.

    Examples:
        Update run text:
        >>> session_id = docx_create()
        >>> para_id = docx_add_paragraph(session_id, "")
        >>> run_id = docx_add_run(session_id, "Old", paragraph_id=para_id)
        >>> docx_set_font(session_id, run_id, bold=True, size=14)
        >>> result = docx_update_run_text(session_id, run_id, "New")
        >>> print(result)
        'Run run_xxx updated successfully'

    Notes:
        - Preserves all formatting: bold, italic, size, color, etc.
        - Only the text content is changed
        - More precise than docx_update_paragraph_text()

    See Also:
        - docx_update_paragraph_text: Replace entire paragraph content
        - docx_set_font: Modify run formatting
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    run = session.get_object(run_id)
    if not run:
        raise ValueError(f"Run {run_id} not found")

    if not hasattr(run, 'text'):
        raise ValueError(f"Object {run_id} is not a run")

    # Update text while preserving formatting
    run.text = new_text

    return f"Run {run_id} updated successfully"


@mcp.tool()
def docx_extract_template_structure(session_id: str) -> str:
    """
    Extract the complete structure of a Word document template.

    Analyzes and returns a comprehensive JSON representation of the document structure,
    including headings, tables, paragraphs, and their properties. Automatically detects
    table headers based on formatting (bold text or background color).

    Typical Use Cases:
        - Analyze template structure before filling
        - Understand document layout programmatically
        - Generate documentation from templates
        - Validate template format

    Args:
        session_id (str): Active session ID returned by docx_create().

    Returns:
        str: JSON string containing document structure with metadata. Format:
            {
                "metadata": {"extracted_at": "...", "docx_version": "..."},
                "document_structure": [
                    {"type": "heading", "level": 1, "text": "...", "style": {...}},
                    {"type": "table", "rows": 5, "cols": 3, "headers": [...], ...},
                    {"type": "paragraph", "text": "...", "style": {...}}
                ]
            }

    Raises:
        ValueError: If session_id is invalid or document is empty.
        ValueError: If table header cannot be detected in strict mode.

    Examples:
        Extract template structure:
        >>> session_id = docx_create(file_path="./template.docx")
        >>> structure = docx_extract_template_structure(session_id)
        >>> import json
        >>> data = json.loads(structure)
        >>> print(f"Document has {len(data['document_structure'])} elements")

        Find all tables in structure:
        >>> data = json.loads(structure)
        >>> tables = [e for e in data['document_structure'] if e['type'] == 'table']
        >>> print(f"Found {len(tables)} tables")

    Notes:
        - Preserves element order as in original document
        - Automatically detects table headers (bold or colored background)
        - Returns complete styling information
        - Does not extract images or embedded objects

    See Also:
        - docx_read_content: Simple text extraction
        - docx_find_table: Find specific tables
    """
    session = session_manager.get_session(session_id)
    if not session:
        raise ValueError(f"Session {session_id} not found")

    if not session.document:
        raise ValueError("Document not found in session")

    parser = TemplateParser()
    structure = parser.extract_structure(session.document)

    return json.dumps(structure, indent=2, ensure_ascii=False)


def main():
    import argparse
    parser = argparse.ArgumentParser(description="DOCX MCP Server")
    parser.add_argument("--transport", default="stdio", choices=["stdio", "sse"], help="Transport protocol")
    parser.add_argument("--host", default="127.0.0.1", help="Host to bind to (SSE only)")
    parser.add_argument("--port", type=int, default=8000, help="Port to listen on (SSE only)")

    # Parse known args to avoid conflict if FastMCP parses its own (though we call run explicitly)
    args, unknown = parser.parse_known_args()

    if args.transport == "sse":
        print(f"Starting SSE server on {args.host}:{args.port}...", flush=True)

        # FastMCP.run() doesn't accept host/port args, we must update settings directly
        mcp.settings.host = args.host
        mcp.settings.port = args.port

        # Disable DNS rebinding protection for non-localhost addresses (LAN access)
        # This is needed for 0.0.0.0 or any specific non-localhost IP (192.168.x.x, 10.x.x.x, etc.)
        # Otherwise, strict checks on Host header will fail for external IPs
        if args.host not in ("127.0.0.1", "localhost"):
            mcp.settings.transport_security = None

        mcp.run(transport="sse")
    else:
        mcp.run(transport="stdio")

if __name__ == "__main__":
    main()
